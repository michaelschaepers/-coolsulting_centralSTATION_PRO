# -*- coding: utf-8 -*-
# ==========================================
# DATEI: coolMATH_PRO.py
# VERSION: 44.10 (UnboundLocalError Fix)
# ZEITSTEMPEL: 20.02.2026 23:15 Uhr
# AUTOR: Michael Schäpers, °coolsulting
# ==========================================
# ÄNDERUNGEN v44.10 (gegenüber v44.9):
# - FIX: partner_firma UnboundLocalError - wird initial gesetzt, später überschrieben
# ==========================================
# ÄNDERUNGEN v44.5 (gegenüber v44.4):
# - Monday.com Fix: save_to_monday Wrapper-Methode hinzugefügt
# - requirements.txt: python-docx==1.1.2 hinzugefügt
# ==========================================
# ÄNDERUNGEN v44.4 (gegenüber v44.3):
# - Word-Export: Prüft ob python-docx installiert ist, zeigt Warnung wenn nicht
# - Excel-Anfrage: Neue Funktion zum Export der Geräteauswahl als Excel
#   (4 Sheets: Projektinfo, Innengeräte, Außengeräte, Bestellung)
# ==========================================
# ÄNDERUNGEN v44.3 (gegenüber v44.2):
# - Liefertermin-Feld im UI (optionaler Kalender-Picker)
# - Liefertermin wird im Technikübergabe-PDF angezeigt
# ==========================================
# ÄNDERUNGEN v44.2 (gegenüber v44.1):
# - Gerätetabelle aufgeteilt: Separate Tabellen für Innengeräte und Außengeräte
# - Bessere Übersichtlichkeit im Übergabe-PDF
# ==========================================
# ÄNDERUNGEN v44.1 (gegenüber v44.0):
# - Zahlenformat: Deutsche Schreibweise mit PUNKT als Tausendertrennzeichen
#   (1.775 statt 1,775) in allen PDFs und Berichten
# - Neue Funktion fmt_number() für konsistente Formatierung
# ==========================================
# ÄNDERUNGEN v44.0 (gegenüber v42):
# - SAMSUNG_SERIEN Datenbank: 9 echte Serien aus Excel 
#   Airise Living, WF Standard, WF Exklusiv, WF Exklusiv Black,
#   WF Exklusiv-Premiere, WF Exklusiv-Premiere Black,
#   WF Elite, WF Elite-Premiere Plus, WF Elite-Premiere Plus Black
# - Serienauswahl im UI (Dropdown vor Vergleichstabelle)
# - find_samsung_device() unterstützt optionalen serie-Parameter
# - device_label() nutzt gewählte Serie
# - hw_map aktualisiert mit echten Artikelnummern
# - Preise aus Excel-Import 
# - METHODEN-Physik: alle 6 Formeln korrigiert (Praktiker, Recknagel,
#   VDI Neu RC-Modell 96h, Kaltluftsee /1.3, KI-Hybrid Phase/Dämpfung)
# - matplotlib Agg Backend gesetzt (PDF kein Crash)
# - Max. Gerät 5.0kW Standard (6.5kW verfügbar in Exklusiv-Premiere)
# ==========================================

import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import json
import os
import io
import tempfile
# PDF: reportlab + python-docx
from datetime import datetime
from typing import Dict, Optional, Tuple

# --- BRANDING KONSTANTEN ---
APP_VERSION = "4.81.0"
CI_BLUE = "#36A9E1"
CI_GRAY = "#3C3C3B"
CI_WHITE = "#FFFFFF"

# --- PREISLISTE ---
PREISLISTE_PATH = "S_Klima_Artikel_Import_2026-02-02-APP.xlsx"

def load_samsung_prices():
    """Lädt Samsung-Preise aus Excel-Datei"""
    import os
    import re
    
    # Suche Preisliste in verschiedenen Pfaden
    search_paths = [
        PREISLISTE_PATH,
        os.path.join(os.path.dirname(__file__), PREISLISTE_PATH),
        os.path.join('/mnt/user-data/outputs', PREISLISTE_PATH),
    ]
    
    xlsx_path = None
    for path in search_paths:
        if os.path.exists(path):
            xlsx_path = path
            break
    
    if not xlsx_path:
        # Fallback auf hardcoded Preise
        return {
            4.0:  {"preis": 2347},
            5.0:  {"preis": 2706},
            5.2:  {"preis": 3061},
            6.8:  {"preis": 3548},
            8.0:  {"preis": 4494},
            10.0: {"preis": 5533},
        }
    
    try:
        import pandas as pd
        df = pd.read_excel(xlsx_path, sheet_name=' Kima 2026-02-02')
        
        # FJM AG Außengeräte extrahieren
        fjm_ag = df[(df['Artikelgruppe'] == 'S_FJM') & 
                    (df['Bezeichnung'].str.contains('AG', case=False)) &
                    (df['Artikelnummer'].str.startswith('AJ0'))]
        
        prices = {}
        for _, row in fjm_ag.iterrows():
            langtext = row['Langtext']
            preis = row['Listenpreis']
            
            # kW aus Langtext extrahieren
            kw_match = re.search(r'Kühlen\s+(\d+\.?\d*)\s*kW', langtext, re.IGNORECASE)
            if kw_match:
                kw = float(kw_match.group(1))
                prices[kw] = {"preis": preis}
        
        return prices if prices else {
            4.0:  {"preis": 2347},
            5.0:  {"preis": 2706},
            5.2:  {"preis": 3061},
            6.8:  {"preis": 3548},
            8.0:  {"preis": 4494},
            10.0: {"preis": 5533},
        }
    except Exception as e:
        print(f"⚠️ Preisliste konnte nicht geladen werden: {e}")
        # Fallback
        return {
            4.0:  {"preis": 2347},
            5.0:  {"preis": 2706},
            5.2:  {"preis": 3061},
            6.8:  {"preis": 3548},
            8.0:  {"preis": 4494},
            10.0: {"preis": 5533},
        }

# --- GERÄTE-PREISLISTEN (für PDF-Berichte) ---
FJM_AG_PRICES = load_samsung_prices()
RAC_AG_PRICES = {
    # Wird aus main() RAC_AG_BY_SERIE extrahiert falls benötigt
}


def pdf_safe(text):
    """Sanitize text for FPDF (latin-1 only). Replace common non-latin-1 chars."""
    if not isinstance(text, str):
        text = str(text)
    replacements = {
        '°': 'deg',   # °
        '²': '2',     # ²
        '³': '3',     # ³
        'ä': 'ae',    # ä
        'ö': 'oe',    # ö
        'ü': 'ue',    # ü
        'Ä': 'Ae',    # Ä
        'Ö': 'Oe',    # Ö
        'Ü': 'Ue',    # Ü
        'ß': 'ss',    # ß
        '×': 'x',     # ×
        '→': '->',    # →
        '–': '-',     # –
        '—': '--',    # —
        '≥': '>=',    # ≥
        '≤': '<=',    # ≤
        'τ': 'tau',   # τ
        'φ': 'phi',   # φ
        'ε': 'eps',   # ε
        'Δ': 'D',     # Δ
        '★': '*',     # ★
        'é': 'e',     # é
        '€': 'EUR',   # €
        '‘': "'",     # '
        '’': "'",     # '
        '“': '"',     # "
        '”': '"',     # "
    }
    for uni, asc in replacements.items():
        text = text.replace(uni, asc)
    # Final fallback: encode to latin-1, replace remaining unknowns
    return text.encode('latin-1', errors='replace').decode('latin-1')


def is_docx_available():
    """Prüft ob python-docx installiert ist"""
    try:
        import docx
        return True
    except ImportError:
        return False


def fmt_number(num, decimals=0):
    """
    Formatiert Zahlen in deutscher Schreibweise MIT PUNKT als Tausendertrennzeichen.
    Beispiel: 1775 -> "1.775" (nicht "1,775")
    """
    if decimals == 0:
        # Ganzzahl
        formatted = f"{int(num):,}".replace(',', '.')
    else:
        # Mit Dezimalstellen
        formatted = f"{num:,.{decimals}f}".replace(',', '.')
    return formatted


# ==========================================
# 1. SETUP & CSS (Kein weißer Balken Bug)
# ==========================================
try:
    st.set_page_config(page_title="coolMATH Pro Simulation", layout="wide", initial_sidebar_state="collapsed")
except Exception:
    pass

def setup_page():
    st.markdown(f"""
        <style>
        /* Basis Reset */
        .stApp {{
            background-color: {CI_BLUE};
            color: white;
        }}
        /* Entfernung des weißen Balken-Bugs */
        [data-testid="stAppViewContainer"] > .main {{
            background-color: {CI_BLUE};
        }}
        [data-testid="stHeader"] {{
            background-color: transparent !important;
            display: none !important;
        }}
        [data-testid="stToolbar"] {{
            display: none !important;
        }}
        [data-testid="stDecoration"] {{
            display: none !important;
        }}
        /* Sidebar verstecken */
        [data-testid="stSidebar"] {{
            display: none;
        }}
        /* Streiche den Deploy-Button */
        .stDeployButton {{
            display: none !important;
        }}
        /* Main container padding */
        .block-container {{
            padding-top: 2rem;
            padding-bottom: 2rem;
        }}
        /* Typografie */
        .cool-text {{
            color: {CI_WHITE};
            font-weight: 700;
            font-size: 68px;
            line-height: 1;
            display: inline-block;
            vertical-align: bottom;
            letter-spacing: -1px;
        }}
        .math-text {{
            color: {CI_GRAY};
            font-weight: 700;
            font-size: 68px;
            line-height: 1;
            display: inline-block;
            vertical-align: bottom;
            letter-spacing: -1px;
        }}
        .version-badge {{
            background: rgba(255,255,255,0.2);
            color: white;
            padding: 4px 12px;
            border-radius: 20px;
            font-size: 11px;
            font-weight: 700;
            letter-spacing: 1px;
            display: inline-block;
            margin-top: 8px;
        }}
        /* Karten */
        .card {{
            background-color: white;
            border-radius: 15px;
            padding: 25px;
            margin-bottom: 20px;
            color: {CI_GRAY};
            box-shadow: 0 4px 20px rgba(0,0,0,0.1);
        }}
        .card-blue {{
            background: linear-gradient(135deg, rgba(255,255,255,0.15), rgba(255,255,255,0.05));
            border: 1px solid rgba(255,255,255,0.3);
            border-radius: 15px;
            padding: 20px;
            margin-bottom: 15px;
            color: white;
            backdrop-filter: blur(10px);
        }}
        /* Ergebnis Matrix */
        .matrix-wrapper {{
            background: white;
            padding: 30px;
            border-radius: 15px;
            margin-top: 20px;
            box-shadow: 0 4px 30px rgba(0,0,0,0.15);
        }}
        .matrix-title {{
            color: {CI_GRAY};
            font-size: 24px;
            font-weight: 700;
            text-align: center;
            margin-bottom: 20px;
            text-transform: uppercase;
            letter-spacing: 2px;
        }}
        /* Streamlit Labels - nicht fett */
        label, .stTextInput label, .stNumberInput label, .stSelectbox label, .stSlider label {{
            font-weight: 400 !important;
        }}
        /* Tabellen */
        .styled-table {{
            width: 100%;
            border-collapse: collapse;
            background-color: white;
        }}
        .styled-table th {{
            background: {CI_GRAY};
            color: white;
            padding: 12px 15px;
            text-align: center !important;
            font-weight: 700;
            text-transform: uppercase;
            font-size: 11px;
            letter-spacing: 1px;
        }}
        .styled-table td {{
            padding: 10px 15px;
            border-bottom: 1px solid #eee;
            text-align: center !important;
            font-weight: 700;
            color: {CI_GRAY};
            font-size: 13px;
        }}
        .total-row td {{
            background-color: #f0f8ff;
            font-size: 14px !important;
            border-top: 3px solid {CI_BLUE} !important;
            color: {CI_BLUE} !important;
            font-weight: 700 !important;
        }}
        .samsung-row td {{
            background-color: #e8f5e9;
            color: #2e7d32 !important;
            font-weight: 700;
        }}
        /* ALL labels: uniform white uppercase style */
        label {{
            color: white !important;
            font-weight: 800 !important;
            font-size: 11px !important;
            letter-spacing: 0.8px !important;
            text-transform: uppercase !important;
        }}
        /* Number inputs: preserve m² superscript — no uppercase transform */
        [data-testid="stNumberInput"] label {{
            text-transform: none !important;
            font-weight: 800 !important;
            color: white !important;
            font-size: 11px !important;
            letter-spacing: 0.8px !important;
        }}
        /* Slider label same style */
        [data-testid="stSlider"] label {{
            text-transform: uppercase !important;
        }}
        /* Divider */
        hr {{
            border-color: rgba(255,255,255,0.3) !important;
            margin: 20px 0;
        }}
        /* Expander */
        [data-testid="stExpander"] {{
            background: rgba(255,255,255,0.15);
            border-radius: 10px;
            border: 1px solid rgba(255,255,255,0.3);
        }}
        /* Expander header button - readable text */
        [data-testid="stExpander"] summary,
        [data-testid="stExpander"] summary p,
        [data-testid="stExpander"] summary span,
        [data-testid="stExpander"] > div > div > div > button,
        [data-testid="stExpander"] > div > div > div > button p,
        [data-testid="stExpander"] > div > div > div > button span {{
            color: white !important;
            font-weight: 700 !important;
            font-size: 13px !important;
            letter-spacing: 1px !important;
            text-transform: uppercase !important;
        }}
        /* Expander arrow icon */
        [data-testid="stExpander"] svg {{
            stroke: white !important;
            fill: white !important;
        }}
        /* Inputs inside expander: white bg + dark text for readability */
        [data-testid="stExpander"] .stTextInput > div > div > input {{
            background: rgba(255,255,255,0.92) !important;
            color: {CI_GRAY} !important;
            font-weight: 700 !important;
            border: 1px solid rgba(255,255,255,0.5) !important;
        }}
        /* Labels inside expander */
        [data-testid="stExpander"] label {{
            color: white !important;
            font-weight: 800 !important;
            font-size: 11px !important;
            letter-spacing: 1px !important;
        }}
        /* Buttons */
        .stButton > button {{
            background: {CI_GRAY};
            color: white;
            border: none;
            font-weight: 700;
            text-transform: uppercase;
            letter-spacing: 1px;
            border-radius: 8px;
            padding: 12px 24px;
        }}
        .stButton > button:hover {{
            background: {CI_BLUE};
            color: white;
            transform: translateY(-1px);
        }}
        /* Download Button */
        .stDownloadButton > button {{
            background: {CI_BLUE};
            color: white;
            border: none;
            font-weight: 700;
            text-transform: uppercase;
            letter-spacing: 1px;
            border-radius: 8px;
        }}
        /* Tabs */
        .stTabs [data-baseweb="tab-list"] {{
            background: rgba(255,255,255,0.1);
            border-radius: 10px;
            padding: 4px;
        }}
        .stTabs [data-baseweb="tab"] {{
            color: rgba(255,255,255,0.7);
            font-weight: 700;
            font-size: 12px;
            letter-spacing: 1px;
        }}
        .stTabs [aria-selected="true"] {{
            background: white !important;
            color: {CI_GRAY} !important;
            border-radius: 8px;
        }}
        /* Selectbox, Input — dark text for readability */
        .stNumberInput > div > div > input,
        .stTextInput > div > div > input,
        .stTextArea > div > div > textarea {{
            background: rgba(255,255,255,0.92) !important;
            color: {CI_GRAY} !important;
            font-weight: 700 !important;
            border: 1px solid rgba(255,255,255,0.5) !important;
            border-radius: 8px !important;
        }}
        /* Selectbox dropdown — dark text */
        .stSelectbox > div > div {{
            background: rgba(255,255,255,0.92) !important;
            color: {CI_GRAY} !important;
            font-weight: 700 !important;
            border: 1px solid rgba(255,255,255,0.5) !important;
            border-radius: 8px !important;
        }}
        /* Selectbox selected value text */
        .stSelectbox [data-baseweb="select"] span,
        .stSelectbox [data-baseweb="select"] div {{
            color: {CI_GRAY} !important;
            font-weight: 700 !important;
        }}
        /* Number input +/- buttons */
        .stNumberInput button {{
            color: {CI_GRAY} !important;
            background: rgba(255,255,255,0.7) !important;
        }}
        /* Samsung Gerät Karte */
        .samsung-card {{
            background: linear-gradient(135deg, #1a5276, #2e86ab);
            border-radius: 12px;
            padding: 16px;
            color: white;
            margin: 8px 0;
        }}
        .samsung-card-match {{
            background: linear-gradient(135deg, #1e8449, #27ae60);
            border-radius: 12px;
            padding: 16px;
            color: white;
            margin: 8px 0;
        }}
        /* Method Badge */
        .method-badge {{
            display: inline-block;
            padding: 3px 10px;
            border-radius: 20px;
            font-size: 10px;
            font-weight: 800;
            letter-spacing: 1px;
            text-transform: uppercase;
        }}
        .badge-vdi-n {{ background: #2ecc71; color: white; }}
        .badge-vdi-a {{ background: #f39c12; color: white; }}
        .badge-reck {{ background: {CI_GRAY}; color: white; }}
        .badge-prak {{ background: #e74c3c; color: white; }}
        .badge-klts {{ background: #9b59b6; color: white; }}
        .badge-ki {{ background: #1abc9c; color: white; }}
        /* Section Header */
        .section-header {{
            color: white;
            font-size: 20px;
            font-weight: 700;
            text-transform: uppercase;
            letter-spacing: 3px;
            margin: 30px 0 15px 0;
            padding-bottom: 8px;
            border-bottom: 2px solid rgba(255,255,255,0.3);
        }}
        </style>
    """, unsafe_allow_html=True)


# ==========================================
# 2. PHYSIK ENGINE — 6 METHODEN
# ==========================================
HOURS = np.arange(24)

# Solare Einstrahlungsdaten [W/m²] pro Stunde und Ausrichtung
SOLAR_DB = {
    "NORD":      [20,20,20,20,20,40,60,80,100,110,120,120,120,120,110,100,80,60,40,20,20,20,20,20],
    "OST":       [0,0,0,0,100,350,550,650,680,600,450,250,150,100,80,60,40,20,0,0,0,0,0,0],
    "SUED":       [0,0,0,0,0,50,150,300,450,550,620,650,650,620,550,450,300,150,50,0,0,0,0,0],
    "WEST":      [0,0,0,0,0,0,20,50,70,90,110,150,250,450,600,680,650,550,350,150,0,0,0,0],
    "SUED-OST":   [0,0,0,50,200,450,580,650,620,500,350,200,120,90,70,40,20,0,0,0,0,0,0,0],
    "SUED-WEST":  [0,0,0,0,0,0,20,50,80,120,200,350,500,620,650,580,450,250,80,20,0,0,0,0],
}

def get_phys_constants(standard, glass, shade):
    """Physikalische Konstanten je Gebäudestandard"""
    u_vals  = {"Altbau": 1.7, "Bestand": 0.8, "Neubau (GEG)": 0.28, "Passivhaus": 0.15}
    g_vals  = {"Einfach": 0.85, "Doppel": 0.65, "Dreifach": 0.50, "Sonnenschutz": 0.32}
    fc_vals = {"Keine": 1.0, "Vorhang (Innen)": 0.6, "Raffstore (Aussen)": 0.25, "Rollladen": 0.15}
    return u_vals.get(standard, 0.8), g_vals.get(glass, 0.65), fc_vals.get(shade, 1.0)


def calc_praktiker(area, orient, standard, glass, shade, pers, tech):
    """METHODE 1: Praktiker (Heuristik) — q-Wert je Standard, Orientierung, Sonnenschutz REDUZIERT"""
    _, _, fc = get_phys_constants(standard, glass, shade)
    q_std  = {"Altbau": 90.0, "Bestand": 75.0, "Neubau (GEG)": 55.0, "Passivhaus": 35.0}
    q_ori  = {"SUED": 15, "SUED-OST": 12, "SUED-WEST": 12, "WEST": 8, "OST": 5, "NORD": 0}
    q_base = (q_std.get(standard, 75.0) + q_ori.get(orient, 0)) * fc
    q_int  = (pers * 100 + tech) / max(area, 1)
    return np.full(24, area * (q_base + q_int))

def calc_recknagel(area, orient, standard, glass, shade, pers, tech, win_area):
    """METHODE 4: Recknagel — Q_tr (dT standard-abh.) + Q_solar + Q_int"""
    u, g, fc = get_phys_constants(standard, glass, shade)
    sol_raw  = np.array(SOLAR_DB[orient], dtype=float)
    dT_map   = {"Altbau": 9.0, "Bestand": 7.0, "Neubau (GEG)": 5.0, "Passivhaus": 3.0}
    delta_T  = dT_map.get(standard, 6.0)
    q_tr     = area * u * delta_T
    q_st     = sol_raw * win_area * g * fc
    q_int    = np.array([
        (pers * 100 + tech + area * 8) if 8 <= h <= 18 else (pers * 50 + tech * 0.1)
        for h in HOURS])
    return q_tr + q_st + q_int

def calc_vdi_alt(reck_curve):
    """
    METHODE 2: VDI 2078 Alt (1996)
    Pauschaler Zuschlag +15-25% auf Recknagel-Basis
    """
    return reck_curve * 1.20


def calc_vdi_neu(area, orient, standard, glass, shade, pers, tech, win_area, bau_m):
    """METHODE 3: VDI 6007 — RC-Tiefpass, gleiche Eingangslast wie Recknagel"""
    u, g, fc = get_phys_constants(standard, glass, shade)
    sol_raw  = np.array(SOLAR_DB[orient], dtype=float)
    tau      = {"Schwer (Beton/Stein)": 18.0, "Mittel (Ziegel/Holz-Beton)": 10.0, "Leicht (Holz/Trockenbau)": 4.0}.get(bau_m, 10.0)
    delta_T  = {"Altbau": 9.0, "Bestand": 7.0, "Neubau (GEG)": 5.0, "Passivhaus": 3.0}.get(standard, 6.0)
    q_ext    = (area * u * delta_T + sol_raw * win_area * g * fc
                + np.array([(pers*100+tech+area*8) if 8<=h<=18 else (pers*50+tech*0.1) for h in HOURS]))
    q_vdi    = np.zeros(24)
    for _ in range(4):  # 96h Einschwingen
        for h in range(24):
            q_prev   = q_vdi[(h-1)%24]
            q_vdi[h] = q_prev + (q_ext[h] - q_prev) / (tau + 1)
    return q_vdi


def calc_kaltluftsee(area, orient, standard, glass, shade, pers, tech, win_area, bau_m, raumhoehe=2.5):
    """METHODE 5: Kaltluftsee — Recknagel / epsilon (eps=1.3 → 23% Reduktion)"""
    return calc_recknagel(area, orient, standard, glass, shade, pers, tech, win_area) / 1.3


def calc_ki_hybrid(area, orient, standard, glass, shade, pers, tech, win_area, bau_m):
    """METHODE 6: KI-Hybrid — Phasenverschiebung + Daempfung + Pre-Cooling"""
    u, g, fc  = get_phys_constants(standard, glass, shade)
    sol_raw   = np.array(SOLAR_DB[orient], dtype=float)
    delta_T   = {"Altbau": 9.0, "Bestand": 7.0, "Neubau (GEG)": 5.0, "Passivhaus": 3.0}.get(standard, 6.0)
    phi       = {"Schwer (Beton/Stein)": 10, "Mittel (Ziegel/Holz-Beton)": 6, "Leicht (Holz/Trockenbau)": 2}.get(bau_m, 6)
    f         = {"Schwer (Beton/Stein)": 0.55, "Mittel (Ziegel/Holz-Beton)": 0.70, "Leicht (Holz/Trockenbau)": 0.88}.get(bau_m, 0.70)
    q_sol     = np.roll(sol_raw, phi) * f * win_area * g * fc
    q_tr      = area * u * delta_T
    q_int     = np.array([(pers*100+tech+area*6) if 8<=h<=18 else (pers*50+tech*0.05) for h in HOURS])
    pre_cool  = np.array([0.75 if 0<=h<=6 else 1.0 for h in HOURS])
    return (q_sol + q_tr + q_int) * pre_cool


# ==========================================
# 3. SAMSUNG DATENBANK — Wind-Free Wandgeräte
# ==========================================

# Samsung Wind-Free Standard Wandgeräte (AR-Serie)
# ==========================================
# SAMSUNG GERÄTEDATENBANK
# Quelle: MTF-Samsung-Klima_Artikel_Import-WaWi_2026-02-02-APP.xlsx
# Gruppen: S_FJM + S_RAC — nur Wandgeräte (Inneneinheiten)
# ==========================================

# Serien-Datenbank: {serie_name: {kw: {art_nr, bez, preis}}}
SAMSUNG_SERIEN = {
    "Wind-Free Standard": {
        2.0: {"art_nr": "AR60F07C1AWN/EU", "bez": "Wind-Free Standard 2.0kW",  "preis": 693.0},
        2.5: {"art_nr": "AR60F09C1AWN/EU", "bez": "Wind-Free Standard 2.5kW",  "preis": 768.0},
        3.5: {"art_nr": "AR60F12C1AWN/EU", "bez": "Wind-Free Standard 3.5kW",  "preis": 843.0},
        5.0: {"art_nr": "AR60F18C1AWN/EU", "bez": "Wind-Free Standard 5.0kW",  "preis": 1219.0},
    },
    "Airise Living": {
        2.0: {"art_nr": "AR50F07C1BHN/EU", "bez": "Airise Living 2.0kW",       "preis": 554.0},
        2.5: {"art_nr": "AR50F09C1BHN/EU", "bez": "Airise Living 2.5kW",       "preis": 616.0},
        3.5: {"art_nr": "AR50F12C1BHN/EU", "bez": "Airise Living 3.5kW",       "preis": 673.0},
        5.0: {"art_nr": "AR50F18C1BHN/EU", "bez": "Airise Living 5.0kW",       "preis": 1017.0},
        6.5: {"art_nr": "AR50F24C1BHN/EU", "bez": "Airise Living 6.5kW",       "preis": 1335.0},
    },
    "Wind-Free Exklusiv": {
        2.0: {"art_nr": "AR70F07C1AWN/EU", "bez": "Wind-Free Exklusiv 2.0kW",  "preis": 1048.0},
        2.5: {"art_nr": "AR70F09C1AWN/EU", "bez": "Wind-Free Exklusiv 2.5kW",  "preis": 1185.0},
        3.5: {"art_nr": "AR70F12C1AWN/EU", "bez": "Wind-Free Exklusiv 3.5kW",  "preis": 1280.0},
        5.0: {"art_nr": "AR70F18C1AWN/EU", "bez": "Wind-Free Exklusiv 5.0kW",  "preis": 1726.0},
        6.5: {"art_nr": "AR70F24C1AWN/EU", "bez": "Wind-Free Exklusiv 6.5kW",  "preis": 2260.0},
    },
    "Wind-Free Exklusiv Black": {
        2.0: {"art_nr": "AR70F07C1ABN/EU", "bez": "WF Exklusiv Black 2.0kW",   "preis": 1178.0},
        2.5: {"art_nr": "AR70F09C1ABN/EU", "bez": "WF Exklusiv Black 2.5kW",   "preis": 1298.0},
        3.5: {"art_nr": "AR70F12C1ABN/EU", "bez": "WF Exklusiv Black 3.5kW",   "preis": 1440.0},
    },
    "Wind-Free Exklusiv-Premiere": {
        2.0: {"art_nr": "AR70H07C1AWN/EU", "bez": "WF Exklusiv-Premiere 2.0kW","preis": 1072.0},
        2.5: {"art_nr": "AR70H09C1AWN/EU", "bez": "WF Exklusiv-Premiere 2.5kW","preis": 1180.0},
        3.5: {"art_nr": "AR70H12C1AWN/EU", "bez": "WF Exklusiv-Premiere 3.5kW","preis": 1312.0},
        4.3: {"art_nr": "AR70H15C1AWN/EU", "bez": "WF Exklusiv-Premiere 4.3kW","preis": 1656.0},
        5.0: {"art_nr": "AR70H18C1AWN/EU", "bez": "WF Exklusiv-Premiere 5.0kW","preis": 1908.0},
        6.5: {"art_nr": "AR70H24C1AWN/EU", "bez": "WF Exklusiv-Premiere 6.5kW","preis": 2492.0},
    },
    "Wind-Free Exklusiv-Premiere Black": {
        2.0: {"art_nr": "AR70H07C1ABN/EU", "bez": "WF Exkl.-Premiere Blk 2.0kW","preis": 1112.0},
        2.5: {"art_nr": "AR70H09C1ABN/EU", "bez": "WF Exkl.-Premiere Blk 2.5kW","preis": 1228.0},
        3.5: {"art_nr": "AR70H12C1ABN/EU", "bez": "WF Exkl.-Premiere Blk 3.5kW","preis": 1364.0},
    },
    "Wind-Free Elite": {
        2.0: {"art_nr": "AR70F07CAAWKN/EU","bez": "Wind-Free Elite 2.0kW",     "preis": 1347.0},
        2.5: {"art_nr": "AR70F09CAAWKN/EU","bez": "Wind-Free Elite 2.5kW",     "preis": 1487.0},
        3.5: {"art_nr": "AR70F12CAAWKN/EU","bez": "Wind-Free Elite 3.5kW",     "preis": 1637.0},
    },
    "Wind-Free Elite-Premiere Plus": {
        2.0: {"art_nr": "AR70H07CAAWN/EU", "bez": "WF Elite-Premiere Plus 2.0kW","preis": 1296.0},
        2.5: {"art_nr": "AR70H09CAAWN/EU", "bez": "WF Elite-Premiere Plus 2.5kW","preis": 1432.0},
        3.5: {"art_nr": "AR70H12CAAWN/EU", "bez": "WF Elite-Premiere Plus 3.5kW","preis": 1576.0},
    },
    "Wind-Free Elite-Premiere Plus Black": {
        2.0: {"art_nr": "AR70H07CAABN/EU", "bez": "WF Elite-Prem.Plus Blk 2.0kW","preis": 1348.0},
        2.5: {"art_nr": "AR70H09CAABN/EU", "bez": "WF Elite-Prem.Plus Blk 2.5kW","preis": 1488.0},
        3.5: {"art_nr": "AR70H12CAABN/EU", "bez": "WF Elite-Prem.Plus Blk 3.5kW","preis": 1640.0},
    },
    # --- Kassetten, Kanaleinbau, Standtruhe (FJM IG) ---
    "Mini-Kassette 620x620": {
        1.6: {"art_nr": "AJ016TNNDKG/EU", "bez": "FJM Mini-Kassette 1.6kW",  "preis": 902.0},
        2.0: {"art_nr": "AJ020TNNDKG/EU", "bez": "FJM Mini-Kassette 2.0kW",  "preis": 946.0},
        2.6: {"art_nr": "AJ026TNNDKG/EU", "bez": "FJM Mini-Kassette 2.6kW",  "preis": 1050.0},
        3.5: {"art_nr": "AJ035TNNDKG/EU", "bez": "FJM Mini-Kassette 3.5kW",  "preis": 1176.0},
        5.2: {"art_nr": "AJ052TNNDKG/EU", "bez": "FJM Mini-Kassette 5.2kW",  "preis": 1388.0},
    },
    "1-Weg-Kassette": {
        2.6: {"art_nr": "AJ026TN1DKG/EU", "bez": "FJM 1-Weg-Kassette 2.6kW", "preis": 946.0},
        3.5: {"art_nr": "AJ035TN1DKG/EU", "bez": "FJM 1-Weg-Kassette 3.5kW", "preis": 1160.0},
    },
    "Kanaleinbau": {
        2.6: {"art_nr": "AJ026TNLPEG/EU", "bez": "FJM Kanaleinbau 2.6kW",    "preis": 1368.0},
        3.5: {"art_nr": "AJ035TNLPEG/EU", "bez": "FJM Kanaleinbau 3.5kW",    "preis": 1422.0},
        5.2: {"art_nr": "AJ052BNMDEG/EU", "bez": "FJM Kanaleinbau 5.2kW",    "preis": 1391.0},
    },
    "Standtruhe": {
        2.6: {"art_nr": "AJ026TNJDKG/EU", "bez": "FJM Standtruhe 2.6kW",     "preis": 1074.0},
        3.5: {"art_nr": "AJ035TNJDKG/EU", "bez": "FJM Standtruhe 3.5kW",     "preis": 1198.0},
        5.2: {"art_nr": "AJ052TNJDKG/EU", "bez": "FJM Standtruhe 5.2kW",     "preis": 1430.0},
    },
}

# Standard-Serie für Erstempfehlung
SAMSUNG_DEFAULT_SERIE = "Wind-Free Standard"

# Kurznamen für Dropdowns (damit der Text nicht abgeschnitten wird)
SERIE_SHORT = {
    "Wind-Free Standard":               "WF Standard",
    "Airise Living":                    "Airise Living",
    "Wind-Free Exklusiv":               "WF Exklusiv",
    "Wind-Free Exklusiv Black":         "WF Exklusiv Blk",
    "Wind-Free Exklusiv-Premiere":      "WF Exkl.-Prem.",
    "Wind-Free Exklusiv-Premiere Black":"WF Exkl.-Prem. Blk",
    "Wind-Free Elite":                  "WF Elite",
    "Wind-Free Elite-Premiere Plus":    "WF Elite-Prem.+",
    "Wind-Free Elite-Premiere Plus Black":"WF Elite-Prem.+ Blk",
    "Mini-Kassette 620x620":            "Mini-Kassette",
    "1-Weg-Kassette":                   "1-Weg-Kassette",
    "Kanaleinbau":                      "Kanaleinbau",
    "Standtruhe":                       "Standtruhe",
}


# Rückwärtskompatible Flach-DB für find_samsung_device (Standard-Serie)
SAMSUNG_WINDFREE_WALL = {}
for kw, d in SAMSUNG_SERIEN[SAMSUNG_DEFAULT_SERIE].items():
    SAMSUNG_WINDFREE_WALL[kw] = {
        "model":    d["bez"],
        "art_nr":   d["art_nr"],
        "cool_kw":  kw,
        "heat_kw":  round(kw * 1.2, 1),
        "seer":     6.2,
        "scop":     4.6,
        "eer":      3.50,
        "preis":    d["preis"],
        "btus":     str(int(kw * 3412)),
        "serie":    SAMSUNG_DEFAULT_SERIE,
    }

SAMSUNG_SIZES_KW = sorted(SAMSUNG_WINDFREE_WALL.keys())

def find_samsung_device(peak_watt, safety_factor=1.10, serie=None):
    """
    Findet passendes Samsung Wandgerät für gegebene Spitzenlast.
    safety_factor: 1.10 = 10% Norm-Zuschlag
    serie: optionaler Serienname (default = Wind-Free Standard)
    Gibt primäres und alternatives Gerät zurück.
    """
    if serie is None:
        serie = SAMSUNG_DEFAULT_SERIE
    # Fallback auf Standard wenn Serie keine passende Größe hat
    db = SAMSUNG_SERIEN.get(serie, SAMSUNG_SERIEN[SAMSUNG_DEFAULT_SERIE])
    sizes = sorted(db.keys())
    required_kw = (peak_watt * safety_factor) / 1000.0

    def make_entry(kw):
        d = db[kw]
        return {
            "model":       d["bez"],
            "art_nr":      d["art_nr"],
            "cool_kw":     kw,
            "heat_kw":     round(kw * 1.2, 1),
            "preis":       d["preis"],
            "kw_class":    kw,
            "required_kw": required_kw,
            "peak_w":      peak_watt,
            "serie":       serie,
        }

    # Primär: kleinstes Gerät >= required_kw
    primary = None
    for kw in sizes:
        if kw >= required_kw:
            primary = make_entry(kw)
            break

    # Fallback: größtes verfügbares
    if primary is None:
        primary = make_entry(sizes[-1])
        primary["oversized"] = True

    # Alternativ: nächstkleineres
    alt = None
    idx_p = sizes.index(primary["kw_class"])
    if idx_p > 0:
        alt = make_entry(sizes[idx_p - 1])

    return primary, alt


def load_samsung_from_file():
    """Versucht Samsung-Daten aus Excel-Datei zu laden (falls vorhanden)"""
    try:
        samsung_files = [f for f in os.listdir('.') 
                        if any(kw in f.lower() for kw in ['samsung', 'mtf', 'klima', 'artikel']) 
                        and f.endswith('.xlsx')]
        if samsung_files:
            df = pd.read_excel(samsung_files[0], engine='openpyxl')
            # Nur Wandgeräte / Wind-Free filtern
            if 'Bezeichnung' in df.columns:
                mask = df['Bezeichnung'].astype(str).str.contains('Wind.?Free|Wandgerät|AR[0-9]', 
                                                                    case=False, regex=True)
                df_wf = df[mask].copy()
                if len(df_wf) > 0:
                    return df_wf
    except Exception:
        pass
    return None


# ==========================================


# ==========================================
# 4. AUTHENTICATION
# ==========================================
USERS = {
    "coolsulting": {"pw": "coolsulting2026!", "firma": "°coolsulting", "role": "admin"},
    "demo":        {"pw": "demo2026",          "firma": "Demo GmbH",     "role": "partner"},
}

def check_login():
    """Zeigt Login-Maske wenn nicht eingeloggt. Gibt (ok, user_info) zurück."""
    if "auth_ok" not in st.session_state:
        st.session_state.auth_ok = False
        st.session_state.auth_user = {}

    # Versuche aus st.secrets zu laden (Streamlit Cloud)
    try:
        for uname, udata in st.secrets.get("users", {}).items():
            USERS[uname] = udata
    except Exception:
        pass

    if st.session_state.auth_ok:
        return True, st.session_state.auth_user

    # Login-Screen
    st.markdown("""
    <style>
    .login-box{max-width:420px;margin:80px auto;background:white;border-radius:16px;
               padding:40px;box-shadow:0 8px 32px rgba(54,169,225,0.18);}
    .login-title{font-size:26px;font-weight:700;color:#36A9E1;margin-bottom:4px;}
    .login-sub{font-size:13px;color:#888;margin-bottom:28px;}
    </style>
    <div class="login-box">
      <div class="login-title">°coolMATH Pro</div>
      <div class="login-sub">Kühllast-Simulation | Bitte einloggen</div>
    </div>""", unsafe_allow_html=True)

    col_l, col_m, col_r = st.columns([1, 2, 1])
    with col_m:
        st.markdown("### 🔐 Login")
        username = st.text_input("Benutzername", key="login_user")
        password = st.text_input("Passwort", type="password", key="login_pw")
        if st.button("Einloggen", type="primary"):
            u = USERS.get(username)
            if u and u["pw"] == password:
                st.session_state.auth_ok   = True
                st.session_state.auth_user = {
                    "username": username,
                    "firma":    u["firma"],
                    "role":     u["role"],
                }
                st.rerun()
            else:
                st.error("❌ Benutzername oder Passwort falsch")
        st.caption("Partner-Zugänge über °coolsulting beantragen.")

    return False, {}


# ==========================================
# 5. DATENBANK (SQLite lokal + Turso-ready)
# ==========================================
import sqlite3, json as _json, hashlib

DB_PATH = "coolmath_projects.db"

def _get_db():
    """Gibt DB-Verbindung zurück. SQLite lokal (Turso deaktiviert)."""
    # Turso deaktiviert - nutze lokale SQLite
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    return conn, "sqlite"

def db_init():
    """Erstellt Tabellen falls noch nicht vorhanden."""
    try:
        conn, _ = _get_db()
        conn.execute("""
            CREATE TABLE IF NOT EXISTS coolmath_projects (
                id          INTEGER PRIMARY KEY AUTOINCREMENT,
                projekt_id  TEXT UNIQUE,
                firma       TEXT,
                username    TEXT,
                projekt     TEXT,
                kunde       TEXT,
                bearbeiter  TEXT,
                datum       TEXT,
                room_data   TEXT,
                results     TEXT,
                devices     TEXT,
                monday_id   TEXT,
                created_at  TEXT
            )""")
        conn.commit()
    except Exception as e:
        pass  # silent fail – app läuft auch ohne DB

def db_save_project(firma, username, proj, kunde, bearbeiter,
                    room_inputs, room_results, g_sums, selected_hw, selected_hw_ag):
    """Speichert Projekt in DB. Gibt projekt_id zurück."""
    try:
        conn, _ = _get_db()
        pid = hashlib.md5(f"{firma}{proj}{kunde}{datetime.now().isoformat()}".encode()).hexdigest()[:12]
        conn.execute("""
            INSERT OR REPLACE INTO coolmath_projects
            (projekt_id, firma, username, projekt, kunde, bearbeiter, datum,
             room_data, results, devices, created_at)
            VALUES (?,?,?,?,?,?,?,?,?,?,?)""", (
            pid, firma, username, proj, kunde, bearbeiter,
            datetime.now().strftime("%d.%m.%Y %H:%M"),
            _json.dumps({"room_inputs": room_inputs}, ensure_ascii=False),
            _json.dumps({
                "room_results": room_results,
                "peaks": {k: float(v.max()) for k, v in g_sums.items()}
            }, ensure_ascii=False),
            _json.dumps({
                "selected_hw": selected_hw,
                "selected_hw_ag": [list(x) if isinstance(x, tuple) else x for x in selected_hw_ag],
            }, ensure_ascii=False),
            datetime.now().isoformat()
        ))
        conn.commit()
        return pid
    except Exception as e:
        st.warning(f"⚠️ DB-Speicherung: {e}")
        return None

def db_load_projects(firma, role="partner"):
    """Lädt Projektliste. Admin sieht alle, Partner nur eigene Firma."""
    try:
        conn, _ = _get_db()
        if role == "admin":
            rows = conn.execute(
                "SELECT projekt_id,firma,projekt,kunde,bearbeiter,datum FROM coolmath_projects ORDER BY created_at DESC"
            ).fetchall()
        else:
            rows = conn.execute(
                "SELECT projekt_id,firma,projekt,kunde,bearbeiter,datum FROM coolmath_projects WHERE firma=? ORDER BY created_at DESC",
                (firma,)
            ).fetchall()
        return rows
    except Exception:
        return []

def db_load_project(projekt_id):
    """Lädt ein Projekt vollständig."""
    try:
        conn, _ = _get_db()
        row = conn.execute(
            "SELECT * FROM coolmath_projects WHERE projekt_id=?", (projekt_id,)
        ).fetchone()
        return row
    except Exception:
        return None

def db_update_monday_id(projekt_id, monday_id):
    try:
        conn, _ = _get_db()
        conn.execute("UPDATE coolmath_projects SET monday_id=? WHERE projekt_id=?",
                     (monday_id, projekt_id))
        conn.commit()
    except Exception:
        pass


# ==========================================
# 6. MONDAY.COM INTEGRATION
# ==========================================
import requests as _requests


def get_monday_secrets():
    """Lädt Monday Secrets mit Fallback"""
    try:
        api_token = st.secrets.get("MONDAY_API_TOKEN", "")
        board_id = st.secrets.get("MONDAY_BOARD_ID", "")
        if not api_token:
            api_token = st.secrets.get("monday_key", "")
        if not board_id:
            board_id = st.secrets.get("monday_board_id", "")
        return api_token, board_id
    except Exception:
        return "", ""

class MondayIntegration:
    """Verwaltet die Kommunikation mit Monday.com"""

    def __init__(self, api_token: str = None, board_id: str = None):
        self.api_url = "https://api.monday.com/v2"
        self.file_api_url = "https://api.monday.com/v2/file"

        if api_token is None or board_id is None:
            default_token, default_board = get_monday_secrets()
            self.api_token = api_token or default_token
            self.board_id = board_id or default_board
        else:
            self.api_token = api_token
            self.board_id = board_id

        self.headers = {
            "Authorization": self.api_token,
            "Content-Type": "application/json"
        }

    def is_configured(self) -> bool:
        return bool(self.api_token and self.board_id)

    def create_item(self, item_name: str, column_values: Dict) -> Optional[str]:
        """
        Erstellt ein neues Item in Monday.com.
        Bei ColumnValueException (z.B. unbekannter Dropdown-Wert) →
        automatischer Retry ohne die problematische Spalte.
        """
        if not self.is_configured():
            return None

        def _try_create(cv: Dict) -> Optional[str]:
            cv_json = json.dumps(cv)
            cv_escaped = cv_json.replace("\\", "\\\\").replace('"', '\\"')
            item_name_escaped = item_name.replace("\\", "\\\\").replace('"', '\\"')
            query = f'''
            mutation {{
                create_item (
                    board_id: {self.board_id},
                    item_name: "{item_name_escaped}",
                    column_values: "{cv_escaped}"
                ) {{
                    id
                }}
            }}
            '''
            try:
                response = requests.post(
                    self.api_url,
                    headers=self.headers,
                    json={"query": query},
                    timeout=10
                )
                if response.status_code == 200:
                    data = response.json()
                    # Prüfe auf ColumnValueException
                    if 'errors' in data:
                        for err in data['errors']:
                            code = err.get('extensions', {}).get('code', '')
                            if code == 'ColumnValueException':
                                return 'COLUMN_ERROR'
                        print(f"Monday.com GraphQL Error: {data['errors']}")
                        return None
                    if 'data' in data and data['data'] and 'create_item' in data['data']:
                        item = data['data']['create_item']
                        if item:
                            return item['id']
                else:
                    print(f"Monday.com HTTP Error: {response.status_code}")
                return None
            except Exception as e:
                print(f"Monday.com API Error: {e}")
                return None

        # Versuch 1: Mit allen Feldern
        result = _try_create(column_values)

        # Versuch 2: Bei Dropdown-Fehler → ohne Dropdown-Spalten wiederholen
        if result == 'COLUMN_ERROR':
            print("⚠️ ColumnValueException → Retry ohne Dropdown-Spalten")
            cv_fallback = {k: v for k, v in column_values.items()
                          if not k.startswith('dropdown_') and not k.startswith('color_')}
            result = _try_create(cv_fallback)

        return result if result and result != 'COLUMN_ERROR' else None

    def upload_file_to_item(self, item_id: str, file_bytes: bytes, filename: str,
                            column_id: str = "file_mkngj4yq") -> bool:
        """
        Lädt eine Datei zu einem Monday Item hoch
        """
        if not self.is_configured():
            return False

        try:
            query = '''
            mutation ($file: File!, $itemId: ID!, $columnId: String!) {
                add_file_to_column (
                    file: $file,
                    item_id: $itemId,
                    column_id: $columnId
                ) {
                    id
                    name
                }
            }
            '''

            variables = {
                "itemId": int(item_id),
                "columnId": column_id
            }

            map_data = {"image": ["variables.file"]}

            files = {
                'query': (None, query),
                'variables': (None, json.dumps(variables)),
                'map': (None, json.dumps(map_data)),
                'image': (filename, file_bytes, 'application/pdf')
            }

            upload_headers = {"Authorization": self.api_token}

            response = requests.post(
                self.file_api_url,
                headers=upload_headers,
                files=files,
                timeout=30
            )

            if response.status_code == 200:
                data = response.json()
                if 'data' in data and data['data'] and 'add_file_to_column' in data['data']:
                    return True
                elif 'errors' in data:
                    print(f"Monday.com File Upload Error: {data['errors']}")
            else:
                print(f"Monday.com File Upload HTTP Error: {response.status_code} - {response.text}")

            return False

        except Exception as e:
            print(f"Monday.com File Upload Exception: {e}")
            return False

    def save_quote_to_monday(self, quote_data: Dict, pdf_bytes: bytes = None,
                             filename: str = None) -> tuple:
        """
        Speichert ein Angebot in Monday.com mit PDF.

        FIX: Korrekte Column-Formate für alle Spaltentypen:
          - date   → {"date": "YYYY-MM-DD"}
          - number → plain String "123.45"
          - dropdown → {"labels": ["Wert"]}   (⚠ war bisher falscher Plain-String)
          - status → {"label": "Wert"}         (⚠ war bisher 'status' als Column-ID)
          - text   → plain String
        """
        if not self.is_configured():
            return False, ""

        column_values = {}

        # ── Datum (date) ──
        if 'datum' in quote_data:
            date_obj = quote_data['datum']
            if isinstance(date_obj, str):
                try:
                    date_obj = datetime.strptime(date_obj, "%Y-%m-%d")
                except Exception:
                    date_obj = datetime.now()
            column_values['date_mknqdvj8'] = {"date": date_obj.strftime("%Y-%m-%d")}
        else:
            column_values['date_mknqdvj8'] = {"date": datetime.now().strftime("%Y-%m-%d")}

        # ── Angebotswert (numeric) ── plain String
        if 'angebotswert' in quote_data:
            column_values['numeric_mknst7mm'] = str(round(float(quote_data['angebotswert']), 2))

        # ── Partner (dropdown) ── nur setzen wenn Label bekannt, sonst weglassen
        # Verhindert ColumnValueException wenn Label nicht in Monday existiert
        if 'partner' in quote_data:
            partner_raw = str(quote_data['partner']).lstrip('°').strip()
            if partner_raw:
                partner_clean = partner_raw[0].upper() + partner_raw[1:]
                column_values['dropdown_mknagc5a'] = {"labels": [partner_clean]}

        # ── PLZ (text) ── plain String
        if 'plz' in quote_data:
            column_values['text_mkn9v26m'] = str(quote_data['plz'])

        # ── Status (status-Spalte) ── korrekte Column-ID aus Board: color_mkncgyk5
        # FIX: früher stand hier 'status' als Column-ID → existiert nicht → Fehler
        column_values['color_mkncgyk5'] = {"label": "Angebot"}

        # Item-Name
        item_name = quote_data.get('angebots_nr', quote_data.get('kunde', 'Neues Angebot'))

        # Item erstellen
        item_id = self.create_item(item_name, column_values)

        if not item_id:
            return False, ""

        # PDF hochladen (wenn vorhanden)
        if pdf_bytes and filename:
            pdf_success = self.upload_file_to_item(item_id, pdf_bytes, filename)
            if not pdf_success:
                print(f"⚠️ Warning: Item created ({item_id}) but PDF upload failed")

        return True, item_id

    def get_board_data(self) -> Optional[Dict]:
        if not self.is_configured():
            return None

        query = f"""
        query {{
            boards (ids: {self.board_id}) {{
                name
                items_page {{
                    items {{
                        name
                        column_values {{
                            id
                            text
                            value
                        }}
                    }}
                }}
            }}
        }}
        """

        try:
            response = requests.post(
                self.api_url,
                headers=self.headers,
                json={"query": query},
                timeout=10
            )
            if response.status_code == 200:
                return response.json()
            return None
        except Exception as e:
            print(f"Monday.com API Error: {e}")
            return None

    def test_connection(self) -> tuple:
        if not self.is_configured():
            return False, "API Token oder Board ID fehlt"

        query = """
        query {
            me {
                name
                email
            }
        }
        """

        try:
            response = requests.post(
                self.api_url,
                headers=self.headers,
                json={"query": query},
                timeout=5
            )

            if response.status_code == 200:
                data = response.json()
                if 'data' in data and data['data'] and 'me' in data['data']:
                    user = data['data']['me']
                    return True, f"Verbunden als {user['name']} ({user['email']})"
                elif 'errors' in data:
                    return False, f"GraphQL Error: {data['errors']}"

            return False, f"HTTP {response.status_code}: {response.text[:100]}"

        except Exception as e:
            return False, f"Exception: {str(e)}"
    
    def save_to_monday(self, data: Dict, pdf_bytes: bytes = None, filename: str = None) -> tuple:
        """
        Wrapper für save_quote_to_monday - kompatibel mit altem Aufruf
        """
        return self.save_quote_to_monday(data, pdf_bytes, filename)


# ── Streamlit Helper ──

def init_monday_integration() -> MondayIntegration:
    if 'monday_client' not in st.session_state:
        st.session_state.monday_client = MondayIntegration()
    return st.session_state.monday_client


def save_quote_to_monday_ui(quote_data: Dict, pdf_bytes: bytes = None,
                             filename: str = None) -> bool:
    monday = init_monday_integration()

    if not monday.is_configured():
        st.warning("⚠️ Monday.com nicht konfiguriert.")
        return False

    with st.spinner("📤 Speichere in Monday.com..."):
        success, item_id = monday.save_quote_to_monday(quote_data, pdf_bytes, filename)

    if success:
        st.success(f"✅ Angebot in Monday.com gespeichert! (Item ID: {item_id})")
        return True
    else:
        st.error("❌ Fehler beim Speichern in Monday.com")
        return False


def render_monday_status():
    monday = init_monday_integration()

    st.markdown("### 🔗 Monday.com Status")

    if monday.is_configured():
        connected, message = monday.test_connection()
        if connected:
            st.success(f"✅ {message}")
        else:
            st.error(f"❌ {message}")
            with st.expander("🔧 Troubleshooting"):
                st.code(f"""
Token: {'✔' if monday.api_token else '✗'}
Board ID: {'✔' if monday.board_id else '✗'}
API URL: {monday.api_url}
                """)
    else:
        st.info("⚙️ Nicht konfiguriert")
        with st.expander("ℹ️ Konfiguration"):
            st.markdown("""
            **Erforderliche Secrets:**
            ```toml
            # .streamlit/secrets.toml
            MONDAY_API_TOKEN = "eyJhbG..."
            MONDAY_BOARD_ID  = "1234567890"
            ```

            **Korrekte Column IDs (aus Board):**
            | Feld          | Column-ID           | Typ      |
            |---------------|---------------------|----------|
            | Datum         | date_mknqdvj8       | date     |
            | Datei         | file_mkngj4yq       | file     |
            | Angebotswert  | numeric_mknst7mm    | number   |
            | Partner       | dropdown_mknagc5a   | dropdown |
            | PLZ           | text_mkn9v26m       | text     |
            | Status        | color_mkncgyk5      | status   |
            """)

# ==========================================
# 7. PDF ENGINE — reportlab
# ==========================================
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import mm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                 TableStyle, PageBreak, Image as RLImage, HRFlowable)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
import io as _io

_BLUE  = colors.HexColor('#36A9E1')
_DARK  = colors.HexColor('#3C3C3B')
_GREEN = colors.HexColor('#1b5e20')
_LGRAY = colors.HexColor('#F4F4F4')
_WHITE = colors.white
_A4W, _A4H = A4

def _make_styles():
    return {
        'h1':    ParagraphStyle('h1',    fontName='Helvetica-Bold', fontSize=13,
                                 textColor=_BLUE, leading=17, spaceBefore=10, spaceAfter=4),
        'h2':    ParagraphStyle('h2',    fontName='Helvetica-Bold', fontSize=10,
                                 textColor=_DARK, leading=14, spaceBefore=6, spaceAfter=2),
        'body':  ParagraphStyle('body',  fontName='Helvetica', fontSize=9,
                                 textColor=_DARK, leading=13, spaceAfter=4),
        'small': ParagraphStyle('small', fontName='Helvetica', fontSize=7.5,
                                 textColor=colors.HexColor('#888888'), leading=11),
        'cover_title': ParagraphStyle('ct', fontName='Helvetica-Bold', fontSize=22,
                                       textColor=_WHITE, leading=28, alignment=TA_LEFT),
        'cover_sub':   ParagraphStyle('cs', fontName='Helvetica', fontSize=11,
                                       textColor=colors.HexColor('#e0f4fc'), leading=16),
        'cover_body':  ParagraphStyle('cb', fontName='Helvetica', fontSize=10,
                                       textColor=_DARK, leading=15),
        'disclaimer':  ParagraphStyle('disc', fontName='Helvetica', fontSize=7,
                                       textColor=colors.HexColor('#999'), leading=10,
                                       alignment=TA_CENTER),
    }

_S = _make_styles()


_COPYRIGHT = "© 2026 °coolsulting — Michael Schäpers | coolMATH Pro 4.76.5"

_COPYRIGHT_LONG = (
    "© 2026 °coolsulting — Michael Schäpers | coolMATH Pro {v} | Alle Rechte vorbehalten. "
    "Die in diesem Bericht enthaltenen Berechnungen und Empfehlungen wurden nach bestem Wissen "
    "und anerkannten Normen (VDI 6007, VDI 2078, Recknagel) erstellt. °coolsulting und der "
    "jeweilige Partner/Nutzer übernehmen keine Haftung für etwaige Fehler, Auslassungen oder "
    "Schäden, die durch die Verwendung dieser Daten entstehen. Die Ergebnisse ersetzen keine "
    "planerische Einzelfallprüfung durch einen zugelassenen Fachplaner. "
    "Energieausweise und behördliche Genehmigungen sind gesondert zu erstellen."
).format(v=APP_VERSION)

def _hf_cover(canvas, doc):
    """Deckblatt: nur Footer"""
    canvas.saveState()
    canvas.setFillColor(_DARK)
    canvas.setFont('Helvetica', 7)
    canvas.drawCentredString(_A4W/2, 12*mm, _COPYRIGHT[:200])
    canvas.restoreState()

def _hf_normal(canvas, doc, partner_firma="", user_firma=""):
    """Normale Seiten: Header-Balken + Footer"""
    canvas.saveState()
    # Header
    canvas.setFillColor(_BLUE)
    canvas.rect(0, _A4H - 28*mm, _A4W, 28*mm, fill=1, stroke=0)
    canvas.setFillColor(_WHITE)
    canvas.setFont('Helvetica-Bold', 13)
    canvas.drawString(15*mm, _A4H - 14*mm, 'coolMATH Pro — Kühllastanalyse')
    canvas.setFont('Helvetica', 8)
    canvas.drawString(15*mm, _A4H - 21*mm, f'Version {APP_VERSION}  |  {datetime.now().strftime("%d.%m.%Y")}')
    if partner_firma:
        canvas.drawRightString(_A4W - 15*mm, _A4H - 21*mm, partner_firma)
    canvas.setStrokeColor(_WHITE)
    canvas.setLineWidth(0.4)
    canvas.line(15*mm, _A4H - 25*mm, _A4W - 15*mm, _A4H - 25*mm)
    # Footer
    canvas.setFillColor(colors.HexColor('#555'))
    canvas.setFont('Helvetica', 7)
    canvas.drawString(15*mm, 10*mm,
        f'coolMATH Pro {APP_VERSION}  |  © 2026 °coolsulting  |  Seite {doc.page}')
    canvas.drawRightString(_A4W - 15*mm, 10*mm, '°coolsulting — KI-gestützte Kühllastsimulation')
    canvas.restoreState()

def _tbl_style_fn(total_row=True):
    ts = TableStyle([
        ('FONTNAME',      (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTNAME',      (0,1), (-1,-1), 'Helvetica'),
        ('FONTSIZE',      (0,0), (-1,-1), 8),
        ('TEXTCOLOR',     (0,0), (-1,-1), _DARK),
        ('BACKGROUND',    (0,0), (-1,0), _DARK),
        ('TEXTCOLOR',     (0,0), (-1,0), _WHITE),
        ('ROWBACKGROUNDS', (0,1), (-1,-2 if total_row else -1), [_WHITE, _LGRAY]),
        ('GRID',          (0,0), (-1,-1), 0.3, colors.HexColor('#CCCCCC')),
        ('ALIGN',         (0,0), (-1,-1), 'CENTER'),
        ('ALIGN',         (0,0), (0,-1), 'LEFT'),
        ('VALIGN',        (0,0), (-1,-1), 'MIDDLE'),
        ('TOPPADDING',    (0,0), (-1,-1), 4),
        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
        ('LEFTPADDING',   (0,0), (-1,-1), 5),
        ('RIGHTPADDING',  (0,0), (-1,-1), 5),
    ])
    if total_row:
        ts.add('BACKGROUND', (0,-1), (-1,-1), _BLUE)
        ts.add('TEXTCOLOR',  (0,-1), (-1,-1), _WHITE)
        ts.add('FONTNAME',   (0,-1), (-1,-1), 'Helvetica-Bold')
    return ts

def _section_hdr(title, subtitle=''):
    items = [HRFlowable(width='100%', thickness=0.5, color=_BLUE, spaceAfter=1),
             Paragraph(title.upper(), _S['h1'])]
    if subtitle:
        items.append(Paragraph(subtitle, _S['small']))
    items.append(Spacer(1, 3*mm))
    return items

def _chart(img_bytes, width=165*mm):
    if not img_bytes:
        return []
    try:
        img = RLImage(_io.BytesIO(img_bytes), width=width, height=width*0.44)
        return [img, Spacer(1, 4*mm)]
    except Exception as e:
        return [Paragraph(f'[Diagramm: {e}]', _S['small'])]

def _thermal_cover_graphic(canvas, x, y, w, h):
    """Thermodynamische Heatmap-Grafik fürs Deckblatt"""
    import math, random
    random.seed(42)
    
    # Gesamte Grafik mit 75% Transparenz
    canvas.saveState()
    canvas.setFillAlpha(0.75)
    canvas.setStrokeAlpha(0.75)
    
    # Gebäude-Silhouette Hintergrund
    canvas.setFillColor(colors.HexColor('#1a3a5c'))
    canvas.rect(x, y, w, h, fill=1, stroke=0)
    # Thermische Heatmap-Streifen (Farbverlauf)
    heat_colors = ['#0d47a1','#1565c0','#1976d2','#0288d1','#0097a7',
                   '#00897b','#43a047','#c0ca33','#fdd835','#fb8c00',
                   '#e53935','#b71c1c']
    stripe_w = w / len(heat_colors)
    for ci, hc in enumerate(heat_colors):
        canvas.setFillColor(colors.HexColor(hc))
        # Wellenform-Streifen
        for row in range(8):
            alpha_mult = 0.3 + 0.7*(row/8)
            ry = y + row*(h/8)
            rh = h/8 * 0.85
            rx = x + ci*stripe_w + math.sin(row*0.8 + ci*0.3)*stripe_w*0.15
            canvas.setFillAlpha(0.75 * alpha_mult)  # 75% base
            canvas.roundRect(rx, ry, stripe_w*0.85, rh, 2, fill=1, stroke=0)
    canvas.setFillAlpha(0.75)
    # Isothermen-Linien
    canvas.setStrokeColor(colors.HexColor('#ffffff'))
    canvas.setLineWidth(0.4)
    canvas.setDash([3,4])
    for line_y in range(3, 9):
        pts_y = y + line_y*(h/10)
        canvas.line(x, pts_y, x+w, pts_y + math.sin(line_y)*8)
    canvas.setDash([])
    # Gebäude-Umriss
    canvas.setStrokeColor(colors.HexColor('#ffffff'))
    canvas.setFillColor(colors.HexColor('#0d2137'))
    canvas.setLineWidth(1.2)
    bw = w*0.22; bh = h*0.55
    bx = x + w*0.12; by = y + h*0.02
    canvas.rect(bx, by, bw, bh, fill=1, stroke=1)
    canvas.rect(bx+bw+6, by, bw*1.3, bh*0.75, fill=1, stroke=1)
    canvas.rect(bx-10, by, bw*0.6, bh*1.15, fill=1, stroke=1)
    # Fenster (gelb = Wärmequellen)
    canvas.setFillColor(colors.HexColor('#ffd600'))
    for wrow in range(3):
        for wcol in range(2):
            canvas.rect(bx+8+wcol*18, by+bh*0.15+wrow*bh*0.22, 10, 8, fill=1, stroke=0)
    # Temperatur-Label
    canvas.setFillColor(colors.HexColor('#ffffff'))
    canvas.setFont('Helvetica-Bold', 7)
    canvas.drawString(x+5, y+h-8, 'THERMISCHE GEBÄUDESIMULATION  •  KÜHLLASTSTROMANALYSE')
    
    canvas.restoreState()

def _make_cover(story, proj, kunde, bearbeiter, firma,
                partner_firma, report_type, g_sums, selected_hw, liefertermin="—"):
    """Deckblatt - EINFACHER STIL WIE SEITE 2"""
    
    # 1. BLAUER HEADER
    badge_text = '👔 KUNDENBERICHT' if report_type == 'kunde' else '🔧 TECHNIKÜBERGABE'
    
    header_title = Paragraph('°coolMATH Pro — Kühllastanalyse', 
        ParagraphStyle('ht', fontName='Helvetica-Bold', fontSize=22, 
                      textColor=_WHITE, alignment=TA_LEFT))
    header_sub = Paragraph(f'Version {APP_VERSION}  |  {datetime.now().strftime("%d.%m.%Y")}', 
        ParagraphStyle('hs', fontName='Helvetica', fontSize=10, 
                      textColor=colors.HexColor('#e0f2f9'), alignment=TA_LEFT))
    header_firma = Paragraph(partner_firma or firma, 
        ParagraphStyle('hf', fontName='Helvetica', fontSize=10, 
                      textColor=_WHITE, alignment=TA_RIGHT))
    
    header_tbl = Table([
        [header_title, header_firma],
        [header_sub, ''],
    ], colWidths=[140*mm, 35*mm])
    header_tbl.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), _BLUE),
        ('VALIGN', (0,0), (0,0), 'TOP'),
        ('VALIGN', (1,0), (1,0), 'TOP'),
        ('SPAN', (0,1), (1,1)),
        ('TOPPADDING', (0,0), (-1,-1), 15),
        ('BOTTOMPADDING', (0,0), (-1,-1), 15),
        ('LEFTPADDING', (0,0), (-1,-1), 15),
        ('RIGHTPADDING', (0,0), (-1,-1), 15),
    ]))
    story.append(header_tbl)
    story.append(Spacer(1, 20*mm))
    
    # 2. TYP-BADGE
    badge_para = Paragraph(badge_text, ParagraphStyle('badge', 
        fontName='Helvetica-Bold', fontSize=13, textColor=_WHITE, alignment=TA_CENTER))
    badge_color = _BLUE if report_type == 'kunde' else colors.HexColor('#546e7a')
    badge_tbl = Table([[badge_para]], colWidths=[175*mm])
    badge_tbl.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), badge_color),
        ('TOPPADDING', (0,0), (-1,-1), 12),
        ('BOTTOMPADDING', (0,0), (-1,-1), 12),
    ]))
    story.append(badge_tbl)
    story.append(Spacer(1, 15*mm))
    
    # 3. PROJEKT-INFO TABELLE
    peak_vdi = int(np.max(g_sums['VDI_N']))
    total_kw = sum(selected_hw)
    
    info_rows = [
        ['Projekt',         proj],
        ['Kunde',           kunde],
        ['Bearbeiter',      bearbeiter],
        ['Firma',           partner_firma or firma],
        ['Datum',           datetime.now().strftime('%d.%m.%Y')],
        ['Installation',    f'{total_kw:.1f} kW gesamt (Samsung Wind-Free)'],
    ]
    
    # Liefertermin nur bei Technikübergabe
    if report_type == 'uebergabe' and liefertermin != "—":
        info_rows.append(['📅 Liefertermin', liefertermin])
    
    info_tbl = Table([[
        Paragraph(r[0], ParagraphStyle('lbl', fontName='Helvetica-Bold',
            fontSize=10, textColor=_DARK)),
        Paragraph(r[1], ParagraphStyle('val', fontName='Helvetica',
            fontSize=10, textColor=_DARK))]
        for r in info_rows],
        colWidths=[50*mm, 125*mm])
    
    info_tbl.setStyle(TableStyle([
        ('ROWBACKGROUNDS', (0,0), (-1,-1), [_WHITE, _LGRAY]),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#ddd')),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('TOPPADDING', (0,0), (-1,-1), 8),
        ('BOTTOMPADDING', (0,0), (-1,-1), 8),
        ('LEFTPADDING', (0,0), (-1,-1), 12),
        ('RIGHTPADDING', (0,0), (-1,-1), 12),
    ]))
    story.append(info_tbl)
    story.append(PageBreak())  # Kein Copyright hier - kommt am Ende


def _eingabe_tabelle(story, room_inputs, zone_names):
    """Eingabedaten pro Raum als Tabelle"""
    story += _section_hdr('Eingabedaten', 'Raumparameter je Zone')
    hdr = ['Parameter', 'Zone 1', 'Zone 2', 'Zone 3', 'Zone 4', 'Zone 5']
    params = [
        ('Bezeichnung',    'name',        lambda v: str(v)),
        ('Fläche [m²]',    'flaeche',     lambda v: f'{v:.1f}' if isinstance(v,(int,float)) else str(v)),
        ('Höhe [m]',       'hoehe',       lambda v: f'{v:.1f}' if isinstance(v,(int,float)) else str(v)),
        ('Personen',       'personen',    lambda v: str(v)),
        ('Fensterfläche',  'fenster',     lambda v: f'{v:.1f}' if isinstance(v,(int,float)) else str(v)),
        ('Orientierung',   'orientierung',lambda v: str(v)),
        ('Nutzung',        'nutzung',     lambda v: str(v)),
        ('U-Wert [W/m²K]', 'u_wert',      lambda v: f'{v:.2f}' if isinstance(v,(int,float)) else str(v)),
    ]
    rows = [hdr]
    for param_label, key, fmt in params:
        row = [param_label]
        for zi in range(5):
            ri = room_inputs[zi] if isinstance(room_inputs, list) and zi < len(room_inputs) else {}
            val = ri.get(key, '—') if isinstance(ri, dict) else '—'
            try:
                row.append(fmt(val))
            except Exception:
                row.append(str(val))
        rows.append(row)
    widths = [38*mm, 25*mm, 25*mm, 25*mm, 25*mm, 25*mm]
    t = Table(rows, colWidths=widths, repeatRows=1)
    t.setStyle(_tbl_style_fn(total_row=False))
    story += [t, Spacer(1, 5*mm)]


def _geraete_tabelle(story, room_results, selected_hw, selected_hw_ag, zone_names, 
                     show_prices=True, show_artnr=True, selected_ig_artnr=None):
    """IG + AG Gerätetabelle - AUFGETEILT IN ZWEI SEPARATE TABELLEN"""
    if selected_ig_artnr is None:
        selected_ig_artnr = ['—'] * 5
    
    # ===== TABELLE 1: INNENGERÄTE =====
    story += _section_hdr('Innengeräte', 'Übersicht Innengeräte je Zone')
    
    if show_prices:
        hdr_ig = ['Zone', 'Leistung', 'Artikelnummer', 'Listenpreis']
        widths_ig = [35*mm, 25*mm, 55*mm, 35*mm]
    else:
        hdr_ig = ['Zone', 'Leistung', 'Artikelnummer']
        widths_ig = [50*mm, 35*mm, 65*mm]
    
    rows_ig = [hdr_ig]
    for zi in range(5):
        ig_kw = selected_hw[zi] if zi < len(selected_hw) else 0
        ig_artnr = selected_ig_artnr[zi] if zi < len(selected_ig_artnr) else '—'
        zone_n = zone_names[zi] if zi < len(zone_names) else f'Zone {zi+1}'
        
        # IG-Preis (TODO: aus Daten holen wenn verfügbar)
        ig_preis_str = '—'
        
        if show_prices:
            rows_ig.append([zone_n, f'{ig_kw:.1f} kW' if ig_kw else 'N.V.', 
                           ig_artnr, ig_preis_str])
        else:
            rows_ig.append([zone_n, f'{ig_kw:.1f} kW' if ig_kw else 'N.V.', ig_artnr])
    
    t_ig = Table(rows_ig, colWidths=widths_ig, repeatRows=1)
    t_ig.setStyle(_tbl_style_fn(total_row=False))
    story += [t_ig, Spacer(1, 8*mm)]
    
    # ===== TABELLE 2: AUSSENGERÄTE =====
    story += _section_hdr('Außengeräte', 'Übersicht Außengeräte je Zone')
    
    if show_prices:
        hdr_ag = ['Zone', 'Typ', 'Leistung', 'Artikelnummer', 'Listenpreis']
        widths_ag = [30*mm, 22*mm, 25*mm, 50*mm, 35*mm]
    else:
        hdr_ag = ['Zone', 'Typ', 'Artikelnummer']
        widths_ag = [50*mm, 35*mm, 65*mm]
    
    rows_ag = [hdr_ag]
    for zi in range(5):
        ag_inf = selected_hw_ag[zi] if zi < len(selected_hw_ag) else ('—', 0, 'N.V.')
        ag_typ = ag_inf[0] if isinstance(ag_inf,(list,tuple)) and len(ag_inf)>0 else '—'
        ag_kw = ag_inf[1] if isinstance(ag_inf,(list,tuple)) and len(ag_inf)>1 else 0
        ag_artnr = ag_inf[2] if isinstance(ag_inf,(list,tuple)) and len(ag_inf)>2 else 'N.V.'
        zone_n = zone_names[zi] if zi < len(zone_names) else f'Zone {zi+1}'
        
        # AG-Preis ermitteln
        ag_preis_str = '—'
        if ag_kw and ag_kw > 0 and ag_typ == 'FJM':
            try:
                if ag_kw in FJM_AG_PRICES:
                    ag_preis_str = f"{fmt_number(FJM_AG_PRICES[ag_kw]['preis'])} EUR"
            except Exception:
                pass
        
        if show_prices:
            rows_ag.append([zone_n, ag_typ, f'{ag_kw:.1f} kW' if ag_kw else 'N.V.', 
                           ag_artnr, ag_preis_str])
        else:
            rows_ag.append([zone_n, ag_typ, ag_artnr])
    
    t_ag = Table(rows_ag, colWidths=widths_ag, repeatRows=1)
    t_ag.setStyle(_tbl_style_fn(total_row=False))
    story += [t_ag, Spacer(1, 5*mm)]



def make_pdf_chart(profiles, total, title, mode_key, hours=HOURS):
    """Erstellt Matplotlib-Chart für PDF-Export"""
    fig, ax = plt.subplots(figsize=(10, 4.5))
    fig.patch.set_facecolor('white')
    ax.set_facecolor('#fafafa')
    
    colors = ['#36A9E1', '#E74C3C', '#2ECC71', '#F39C12', '#9B59B6']
    for idx, p in enumerate(profiles):
        ax.plot(hours, p[mode_key], alpha=0.6, linewidth=1.5, 
                label=p["name"], color=colors[idx % len(colors)], linestyle='--')
    
    ax.plot(hours, total, color='#3C3C3B', linewidth=3.5, label='GESAMT SIMULTAN', zorder=5)
    
    ax.set_title(title, fontweight='bold', fontsize=12, color='#3C3C3B')
    ax.set_xlabel('Stunde', fontsize=9)
    ax.set_ylabel('Kuhllast [W]', fontsize=9)
    ax.grid(True, alpha=0.3, linestyle=':')
    ax.legend(loc='upper left', fontsize=8, ncol=3)
    ax.set_xlim(0, 23)
    
    buf = io.BytesIO()
    plt.savefig(buf, format='png', bbox_inches='tight', dpi=150)
    plt.close(fig)
    return buf.getvalue()


def make_comparison_chart(g_sums, hours=HOURS):
    """Erstellt Vergleichs-Chart aller Methoden für PDF"""
    fig, ax = plt.subplots(figsize=(10, 5))
    fig.patch.set_facecolor('white')
    ax.set_facecolor('#fafafa')
    
    method_colors = {
        "VDI NEU (VDI 6007)": ('#36A9E1', 3.5, '-'),
        "VDI ALT (2078-1996)": ('#F39C12', 2.0, '--'),
        "Recknagel":           ('#3C3C3B', 2.0, ':'),
        "Praktiker":           ('#E74C3C', 2.5, '-.'),
        "Kaltluftsee":         ('#9B59B6', 2.0, '--'),
        "KI-Hybrid":           ('#1ABC9C', 2.5, '-'),
    }
    key_map = {
        "VDI NEU (VDI 6007)": "VDI_N",
        "VDI ALT (2078-1996)": "VDI_A",
        "Recknagel":           "RECK",
        "Praktiker":           "PRAK",
        "Kaltluftsee":         "KLTS",
        "KI-Hybrid":           "KI",
    }
    
    for name, (color, lw, ls) in method_colors.items():
        key = key_map[name]
        ax.plot(hours, g_sums[key], color=color, linewidth=lw, linestyle=ls, label=name)
    
    ax.set_title('METHODENVERGLEICH - SIMULTAN-TRENDKURVEN', fontweight='bold', 
                 fontsize=12, color='#3C3C3B')
    ax.set_xlabel('Tagesstunde [h]', fontsize=9)
    ax.set_ylabel('Kuhllast [W]', fontsize=9)
    ax.grid(True, alpha=0.3, linestyle=':')
    ax.legend(loc='upper left', fontsize=8)
    ax.set_xlim(0, 23)
    
    buf = io.BytesIO()
    plt.savefig(buf, format='png', bbox_inches='tight', dpi=150)
    plt.close(fig)
    return buf.getvalue()


# ==========================================
# 7. PDF REPORT: KUNDENVERSION
# ==========================================



def build_transfer_report(proj, kunde, bearbeiter, firma, room_results, g_sums,
                          samsung_recommendations, selected_hw, total_kw, selected_hw_ag):
    """Erstellt JSON-Datensatz für coolMATCH-Übergabe"""
    import json
    data = {
        "projekt": proj,
        "kunde": kunde,
        "bearbeiter": bearbeiter,
        "firma": firma,
        "datum": datetime.now().strftime("%d.%m.%Y %H:%M"),
        "version": APP_VERSION,
        "zonen": [],
        "peaks": {
            "vdi_neu": int(np.max(g_sums["VDI_N"])),
            "vdi_alt": int(np.max(g_sums["VDI_A"])),
            "recknagel": int(np.max(g_sums["RECK"])),
            "praktiker": int(np.max(g_sums["PRAK"])),
            "kaltluftsee": int(np.max(g_sums["KLTS"])),
            "ki_hybrid": int(np.max(g_sums["KI"])),
        },
        "installation": {
            "gesamt_kw": round(total_kw, 2),
            "anzahl_zonen": len([kw for kw in selected_hw if kw > 0]),
        }
    }
    for idx, r in enumerate(room_results):
        ig_kw = selected_hw[idx] if idx < len(selected_hw) else 0
        ag_info = selected_hw_ag[idx] if idx < len(selected_hw_ag) else ("—", 0, "N.V.")
        sr = samsung_recommendations[idx] if idx < len(samsung_recommendations) else {}
        zone = {
            "name": r["ZONE"],
            "peak_vdi_neu": r["VDI NEU"],
            "peak_vdi_alt": r["VDI ALT"],
            "peak_recknagel": r["RECKNAGEL"],
            "peak_praktiker": r["PRAKTIKER"],
            "peak_kaltluftsee": r.get("KALTLUFTSEE", 0),
            "peak_ki_hybrid": r.get("KI HYBRID", 0),
            "ig_kw": ig_kw,
            "ag_typ": ag_info[0] if isinstance(ag_info, (list, tuple)) else "—",
            "ag_kw": ag_info[1] if isinstance(ag_info, (list, tuple)) and len(ag_info) > 1 else 0,
            "ag_artnr": ag_info[2] if isinstance(ag_info, (list, tuple)) and len(ag_info) > 2 else "N.V.",
            "samsung_empfehlung": sr.get("primary", {}).get("art_nr", "—") if sr else "—",
        }
        data["zonen"].append(zone)
    return json.dumps(data, ensure_ascii=False, indent=2)


def generate_kunden_pdf(proj, kunde, bearbeiter, firma, room_results, g_sums,
                         individual_profiles, samsung_recommendations,
                         selected_hw, total_installed_kw, selected_hw_ag=None,
                         room_inputs=None, partner_firma="", selected_ig_artnr=None):
    if selected_hw_ag is None: selected_hw_ag = []
    if selected_ig_artnr is None: selected_ig_artnr = ['—'] * 5
    if room_inputs is None:    room_inputs = [{} for _ in range(5)]
    zone_names = [r.get('ZONE', f'Zone {i+1}') for i, r in enumerate(room_results)]

    buf = _io.BytesIO()
    hf  = lambda c, d: _hf_normal(c, d, partner_firma, firma)
    doc = SimpleDocTemplate(buf, pagesize=A4,
                             topMargin=45*mm, bottomMargin=22*mm,
                             leftMargin=15*mm, rightMargin=15*mm)
    story = []

    # Deckblatt
    _make_cover(story, proj, kunde, bearbeiter, firma,
                partner_firma, 'kunde', g_sums, selected_hw)

    # Eingabedaten
    _eingabe_tabelle(story, room_inputs, zone_names)

    # Executive Summary
    story += _section_hdr('Executive Summary', 'Analyse & Empfehlung')
    peak_vdi = int(np.max(g_sums['VDI_N']))
    peak_ki  = int(np.max(g_sums['KI']))
    einspar  = round((peak_vdi - peak_ki) / peak_vdi * 100) if peak_vdi > 0 else 0
    summary  = (
        f"Für Projekt «{proj}» (Auftraggeber: {kunde}) wurde eine Kühllastanalyse "
        f"nach 6 Berechnungsverfahren durchgeführt. Simultanspitze VDI 6007: "
        f"{fmt_number(peak_vdi)} W ({peak_vdi/1000:.1f} kW). Das KI-Hybrid-Modell mit "
        f"Pre-Cooling reduziert auf {fmt_number(peak_ki)} W — Einsparung {einspar}%. "
        f"Gesamtinstallation: {total_installed_kw:.1f} kW Samsung Wind-Free."
    )
    story.append(Paragraph(summary, _S['body']))
    story.append(Spacer(1, 4*mm))

    # Ergebnis-Matrix (OHNE Preise)
    story += _section_hdr('Kühllast-Ergebnisse', 'Alle 6 Methoden — Simultanspitzenwerte [W]')
    hdr = ['Zone', 'VDI 6007', 'VDI 2078 Alt', 'Recknagel', 'Praktiker', 'Kaltl.see', 'KI-Hybrid']
    rows = [hdr]
    for r in room_results:
        rows.append([r['ZONE'], fmt_number(r['VDI NEU']), fmt_number(r['VDI ALT']),
                     fmt_number(r['RECKNAGEL']), fmt_number(r['PRAKTIKER']),
                     fmt_number(r.get('KALTLUFTSEE',0)), fmt_number(r.get('KI HYBRID',0))])
    rows.append(['SIMULTAN-PEAK',
                 fmt_number(int(np.max(g_sums['VDI_N']))), fmt_number(int(np.max(g_sums['VDI_A']))),
                 fmt_number(int(np.max(g_sums['RECK']))),  fmt_number(int(np.max(g_sums['PRAK']))),
                 fmt_number(int(np.max(g_sums['KLTS']))),  fmt_number(int(np.max(g_sums['KI'])))])
    t = Table(rows, colWidths=[28*mm,24*mm,24*mm,24*mm,24*mm,24*mm,24*mm], repeatRows=1)
    t.setStyle(_tbl_style_fn())
    story += [t, Spacer(1, 5*mm)]

    # Geräteauswahl (kein Preis im Kundenbericht)
    _geraete_tabelle(story, room_results, selected_hw, selected_hw_ag, zone_names, 
                     show_prices=False, show_artnr=False, selected_ig_artnr=selected_ig_artnr)

    # Alle 6 Einzelzonen-Diagramme
    story.append(PageBreak())
    story += _section_hdr('Simultan-Diagramme', 'Alle 6 Berechnungsverfahren — Einzelzonen')
    diagramme = [
        ('VDI 6007 Neu',   'vdi_n', 'VDI_N'),
        ('VDI 2078 Alt',   'vdi_a', 'VDI_A'),
        ('Praktiker',      'prak',  'PRAK'),
        ('Recknagel',      'reck',  'RECK'),
        ('Kaltluftsee',    'klts',  'KLTS'),
        ('KI-Hybrid',      'ki',    'KI'),
    ]
    for i, (title, mode_key, sum_key) in enumerate(diagramme):
        if i > 0 and i % 2 == 0:
            story.append(PageBreak())
        story.append(Paragraph(title, _S['h2']))
        img_b = make_pdf_chart(individual_profiles, g_sums[sum_key], title, mode_key)
        story += _chart(img_b, width=165*mm)

    # Disclaimer Footer-Seite
    story.append(PageBreak())
    story += _section_hdr('Rechtlicher Hinweis & Haftungsausschluss')
    story.append(Paragraph(_COPYRIGHT, _S['body']))

    doc.build(story, onFirstPage=_hf_cover, onLaterPages=hf)
    return buf.getvalue()


def generate_uebergabe_pdf(proj, kunde, bearbeiter, firma, room_results, g_sums,
                            individual_profiles, samsung_recommendations,
                            selected_hw, total_installed_kw, selected_hw_ag=None,
                            room_inputs=None, partner_firma="", selected_ig_artnr=None,
                            liefertermin="—"):
    if selected_hw_ag is None: selected_hw_ag = []
    if selected_ig_artnr is None: selected_ig_artnr = ['—'] * 5
    if room_inputs is None:    room_inputs = [{} for _ in range(5)]
    zone_names = [r.get('ZONE', f'Zone {i+1}') for i, r in enumerate(room_results)]

    buf = _io.BytesIO()
    hf  = lambda c, d: _hf_normal(c, d, partner_firma, firma)
    doc = SimpleDocTemplate(buf, pagesize=A4,
                             topMargin=45*mm, bottomMargin=22*mm,
                             leftMargin=15*mm, rightMargin=15*mm)
    story = []

    # Deckblatt
    _make_cover(story, proj, kunde, bearbeiter, firma,
                partner_firma, 'uebergabe', g_sums, selected_hw, liefertermin)

    # Eingabedaten
    _eingabe_tabelle(story, room_inputs, zone_names)

    # Vollständige Ergebnismatrix MIT Preisen
    story += _section_hdr('Vollständige Ergebnismatrix', 'Alle 6 Methoden [W]')
    hdr = ['Zone', 'VDI Neu', 'VDI Alt', 'Recknagel', 'Praktiker', 'Kaltl.', 'KI-Hyb.']
    rows = [hdr]
    for r in room_results:
        rows.append([r['ZONE'], fmt_number(r['VDI NEU']), fmt_number(r['VDI ALT']),
                     fmt_number(r['RECKNAGEL']), fmt_number(r['PRAKTIKER']),
                     fmt_number(r.get('KALTLUFTSEE',0)), fmt_number(r.get('KI HYBRID',0))])
    rows.append(['SIMULTAN-PEAK',
                 fmt_number(int(np.max(g_sums['VDI_N']))), fmt_number(int(np.max(g_sums['VDI_A']))),
                 fmt_number(int(np.max(g_sums['RECK']))),  fmt_number(int(np.max(g_sums['PRAK']))),
                 fmt_number(int(np.max(g_sums['KLTS']))),  fmt_number(int(np.max(g_sums['KI'])))])
    t = Table(rows, colWidths=[28*mm,24*mm,24*mm,24*mm,24*mm,24*mm,24*mm], repeatRows=1)
    t.setStyle(_tbl_style_fn())
    story += [t, Spacer(1, 5*mm)]

    # Geräteauswahl MIT Preisen (Technikübergabe)
    _geraete_tabelle(story, room_results, selected_hw, selected_hw_ag, zone_names,
                     show_prices=True, show_artnr=True, selected_ig_artnr=selected_ig_artnr)

    # Alle 6 Einzelzonen-Diagramme
    story.append(PageBreak())
    story += _section_hdr('Simultan-Diagramme', 'Alle 6 Berechnungsverfahren — Einzelzonen')
    diagramme = [
        ('VDI 6007 Neu',   'vdi_n', 'VDI_N'),
        ('VDI 2078 Alt',   'vdi_a', 'VDI_A'),
        ('Praktiker',      'prak',  'PRAK'),
        ('Recknagel',      'reck',  'RECK'),
        ('Kaltluftsee',    'klts',  'KLTS'),
        ('KI-Hybrid',      'ki',    'KI'),
    ]
    for i, (title, mode_key, sum_key) in enumerate(diagramme):
        if i > 0 and i % 2 == 0:
            story.append(PageBreak())
        story.append(Paragraph(title, _S['h2']))
        img_b = make_pdf_chart(individual_profiles, g_sums[sum_key], title, mode_key)
        story += _chart(img_b, width=165*mm)

    # Methodenvergleich
    story.append(PageBreak())
    story += _section_hdr('Methodenvergleich', 'Alle Methoden überlagert')
    story += _chart(make_comparison_chart(g_sums), width=165*mm)

    # Haftungsausschluss
    story.append(PageBreak())
    story += _section_hdr('Rechtlicher Hinweis & Haftungsausschluss')
    story.append(Paragraph(_COPYRIGHT, _S['body']))

    doc.build(story, onFirstPage=_hf_cover, onLaterPages=hf)
    return buf.getvalue()


def generate_word_report(proj, kunde, bearbeiter, firma, room_results, g_sums,
                          selected_hw, total_installed_kw, selected_hw_ag=None,
                          room_inputs=None, partner_firma="", selected_ig_artnr=None):
    """Word-Dokument mit python-docx — vollständiger Bericht"""
    if selected_hw_ag is None: selected_hw_ag = []
    if selected_ig_artnr is None: selected_ig_artnr = ['—'] * 5
    if room_inputs is None:    room_inputs = [{} for _ in range(5)]
    try:
        from docx import Document as DocxDoc
    except ImportError:
        st.error("⚠️ python-docx nicht installiert. Bitte installieren: pip install python-docx")
        raise ImportError("python-docx fehlt — pip install python-docx")
    from docx.shared import Pt, RGBColor, Cm, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = DocxDoc()
    for sec in doc.sections:
        sec.top_margin    = Mm(20)
        sec.bottom_margin = Mm(20)
        sec.left_margin   = Mm(20)
        sec.right_margin  = Mm(20)

    def _h(text, level=1, color=(54,169,225)):
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
            run.font.name = 'Arial'
        return p

    def _p(text, bold=False, size=10, italic=False):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = bold; r.italic = italic
        r.font.size = Pt(size); r.font.name = 'Arial'
        return p

    def _shade_cell(cell, hex_fill):
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), hex_fill)
        shd.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shd)

    def _tbl(headers, rows, col_widths_cm):
        t = doc.add_table(rows=1+len(rows), cols=len(headers))
        t.style = 'Table Grid'
        hr = t.rows[0]
        for ci, h in enumerate(headers):
            cell = hr.cells[ci]
            cell.width = Cm(col_widths_cm[ci])
            _shade_cell(cell, '3C3C3B')
            r = cell.paragraphs[0].add_run(h)
            r.bold = True; r.font.color.rgb = RGBColor(255,255,255)
            r.font.size = Pt(8); r.font.name = 'Arial'
        for ri, row in enumerate(rows):
            dr = t.rows[ri+1]
            fill = 'F4F4F4' if ri % 2 == 0 else 'FFFFFF'
            for ci, val in enumerate(row):
                cell = dr.cells[ci]
                cell.width = Cm(col_widths_cm[ci])
                _shade_cell(cell, fill)
                r = cell.paragraphs[0].add_run(str(val))
                r.font.size = Pt(8); r.font.name = 'Arial'
        doc.add_paragraph()

    # Titelseite
    doc.add_paragraph()
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run('°coolMATH Pro — Kühllastanalyse')
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(54,169,225); r.font.name='Arial'
    doc.add_paragraph()
    for lbl, val in [('Projekt', proj), ('Kunde', kunde), ('Bearbeiter', bearbeiter),
                      ('Firma', partner_firma or firma), ('Datum', datetime.now().strftime('%d.%m.%Y')),
                      ('Version', APP_VERSION)]:
        p = doc.add_paragraph()
        r1 = p.add_run(f'{lbl}: '); r1.bold=True; r1.font.name='Arial'; r1.font.size=Pt(11)
        r2 = p.add_run(val); r2.font.name='Arial'; r2.font.size=Pt(11)
    doc.add_page_break()

    # Eingabedaten
    _h('Eingabedaten', level=1)
    zone_names = [r.get('ZONE', f'Zone {zi+1}') for zi, r in enumerate(room_results)]
    hdr_e = ['Parameter'] + zone_names
    params_e = [('Fläche [m²]','flaeche'), ('Höhe [m]','hoehe'), ('Personen','personen'),
                ('Fenster [m²]','fenster'), ('Orientierung','orientierung')]
    rows_e = []
    for lbl, key in params_e:
        row = [lbl]
        for zi in range(5):
            ri = room_inputs[zi] if zi < len(room_inputs) else {}
            row.append(str(ri.get(key,'—')) if isinstance(ri,dict) else '—')
        rows_e.append(row)
    _tbl(hdr_e, rows_e, [3.5,2.5,2.5,2.5,2.5,2.5])

    # Ergebnismatrix
    _h('Ergebnis-Matrix', level=1)
    hdr_r = ['Zone','VDI Neu','VDI Alt','Recknagel','Praktiker','Kaltl.see','KI-Hybrid']
    rows_r = []
    for r in room_results:
        rows_r.append([r['ZONE'], fmt_number(r['VDI NEU']), fmt_number(r['VDI ALT']),
                       fmt_number(r['RECKNAGEL']), fmt_number(r['PRAKTIKER']),
                       fmt_number(r.get('KALTLUFTSEE',0)), fmt_number(r.get('KI HYBRID',0))])
    rows_r.append(['SIMULTAN-PEAK',
                   fmt_number(int(np.max(g_sums['VDI_N']))), fmt_number(int(np.max(g_sums['VDI_A']))),
                   fmt_number(int(np.max(g_sums['RECK']))), fmt_number(int(np.max(g_sums['PRAK']))),
                   fmt_number(int(np.max(g_sums['KLTS']))), fmt_number(int(np.max(g_sums['KI'])))])
    _tbl(hdr_r, rows_r, [3.2,2.3,2.3,2.3,2.3,2.3,2.3])

    # Innengeräte
    _h('Innengeräte', level=1)
    hdr_ig = ['Zone', 'Leistung', 'Artikelnummer', 'Listenpreis']
    rows_ig = []
    for zi in range(5):
        ig_kw = selected_hw[zi] if zi < len(selected_hw) else 0
        ig_artnr = selected_ig_artnr[zi] if zi < len(selected_ig_artnr) else '—'
        rows_ig.append([
            zone_names[zi],
            f'{ig_kw:.1f} kW' if ig_kw else 'N.V.',
            ig_artnr,
            '(s. Angebot)'
        ])
    _tbl(hdr_ig, rows_ig, [3.0, 2.5, 6.0, 3.5])
    
    # Außengeräte
    _h('Außengeräte', level=1)
    hdr_ag = ['Zone', 'Typ', 'Leistung', 'Artikelnummer', 'Listenpreis']
    rows_ag = []
    for zi in range(5):
        ag_inf = selected_hw_ag[zi] if zi < len(selected_hw_ag) else ('—', 0, 'N.V.')
        ag_typ = ag_inf[0] if isinstance(ag_inf, (list, tuple)) else '—'
        ag_kw = ag_inf[1] if isinstance(ag_inf, (list, tuple)) and len(ag_inf) > 1 else 0
        ag_artnr = ag_inf[2] if isinstance(ag_inf, (list, tuple)) and len(ag_inf) > 2 else 'N.V.'
        rows_ag.append([
            zone_names[zi],
            ag_typ,
            f'{ag_kw:.1f} kW' if ag_kw else 'N.V.',
            ag_artnr,
            '(s. Angebot)'
        ])
    _tbl(hdr_ag, rows_ag, [2.8, 2.0, 2.5, 5.5, 3.0])
    _p(f'Gesamt: {sum(selected_hw):.1f} kW', bold=True, size=11)

    # Copyright + Haftung
    doc.add_page_break()
    _h('Rechtlicher Hinweis', level=1)
    _p(_COPYRIGHT, size=9, italic=True)

    buf = _io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_excel_anfrage(proj, kunde, bearbeiter, firma, selected_hw, selected_hw_ag, 
                           zone_names, selected_ig_artnr=None, liefertermin="—"):
    """
    Generiert Excel-Anfrage für °coolsulting
    ALLES IN EINEM SHEET mit übersichtlicher Struktur
    """
    if selected_ig_artnr is None:
        selected_ig_artnr = ['—'] * 5
    
    # Sammle alle Daten in einer Liste (Zeilen)
    rows = []
    
    # ===== PROJEKTINFORMATIONEN =====
    rows.append(['PROJEKTINFORMATIONEN', '', '', '', ''])
    rows.append(['Projekt', proj, '', '', ''])
    rows.append(['Kunde', kunde, '', '', ''])
    rows.append(['Bearbeiter', bearbeiter, '', '', ''])
    rows.append(['Firma', firma, '', '', ''])
    rows.append(['Datum', datetime.now().strftime('%d.%m.%Y'), '', '', ''])
    rows.append(['Liefertermin', liefertermin, '', '', ''])
    rows.append(['', '', '', '', ''])  # Leerzeile
    
    # ===== INNENGERÄTE =====
    rows.append(['INNENGERÄTE', '', '', '', ''])
    rows.append(['Zone', 'Leistung [kW]', 'Artikelnummer', 'Menge', ''])
    
    for zi in range(5):
        if zi < len(selected_hw) and selected_hw[zi] > 0:
            rows.append([
                zone_names[zi] if zi < len(zone_names) else f'Zone {zi+1}',
                selected_hw[zi],
                selected_ig_artnr[zi] if zi < len(selected_ig_artnr) else '—',
                1,
                ''
            ])
    
    rows.append(['', '', '', '', ''])  # Leerzeile
    
    # ===== AUSSENGERÄTE =====
    rows.append(['AUSSENGERÄTE', '', '', '', ''])
    rows.append(['Zone', 'Typ', 'Leistung [kW]', 'Artikelnummer', 'Menge'])
    
    for zi in range(5):
        if zi < len(selected_hw_ag):
            ag_inf = selected_hw_ag[zi]
            if isinstance(ag_inf, (list, tuple)) and len(ag_inf) >= 3:
                ag_typ = ag_inf[0]
                ag_kw = ag_inf[1]
                ag_artnr = ag_inf[2]
                
                if ag_kw > 0 and ag_artnr != 'N.V.':
                    rows.append([
                        zone_names[zi] if zi < len(zone_names) else f'Zone {zi+1}',
                        ag_typ,
                        ag_kw,
                        ag_artnr,
                        1
                    ])
    
    # Erstelle DataFrame und Excel
    df = pd.DataFrame(rows)
    
    output = _io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='Anfrage', index=False, header=False)
        
        # Formatierung
        workbook = writer.book
        worksheet = writer.sheets['Anfrage']
        
        # Spaltenbreiten anpassen
        worksheet.column_dimensions['A'].width = 20
        worksheet.column_dimensions['B'].width = 18
        worksheet.column_dimensions['C'].width = 25
        worksheet.column_dimensions['D'].width = 25
        worksheet.column_dimensions['E'].width = 10
        
        # Überschriften fett machen (PROJEKTINFORMATIONEN, INNENGERÄTE, AUSSENGERÄTE)
        from openpyxl.styles import Font, Alignment
        bold_font = Font(bold=True, size=12)
        header_font = Font(bold=True, size=10)
        
        for row in worksheet.iter_rows():
            cell_value = str(row[0].value) if row[0].value else ''
            if cell_value in ['PROJEKTINFORMATIONEN', 'INNENGERÄTE', 'AUSSENGERÄTE']:
                row[0].font = bold_font
            elif cell_value in ['Zone', 'Projekt', 'Kunde', 'Bearbeiter', 'Firma', 'Datum', 'Liefertermin']:
                row[0].font = header_font
    
    output.seek(0)
    return output.getvalue()


def main():
    setup_page()
    db_init()

    # --- LOGIN DISABLED ---
    # auth_ok, auth_user = check_login()
    # if not auth_ok:
    #     return
    auth_role     = "admin"
    auth_username = "coolsulting"
    partner_firma = "°coolsulting"  # Wird später aus Input überschrieben

    col_hdr, col_logo = st.columns([4, 1])
    with col_hdr:
        st.markdown(
            '<div style="padding-top:12px; padding-left:4px; line-height:1;">'
            '<div style="display:inline-block; line-height:0.88;">'
            '<span class="cool-text">°cool</span>'
            '<span class="math-text">MATH</span>'
            '</div>'
            '<br>'
            '<span class="version-badge" style="margin-top:6px; display:inline-block;">'
            'PRO v42 &mdash; 6-METHODEN SIMULATION'
            '</span>'
            '</div>',
            unsafe_allow_html=True
        )
    with col_logo:
        logo_opts = [
            "Coolsulting_Logo_ohneHG_outlines_weiss.png",
            "Coolsulting_Logo_ohneHG_weiss.png",
        ]
        for logo in logo_opts:
            if os.path.exists(logo):
                st.image(logo, width=200)
                break
    
    st.write("---")
    
    # --- PROJEKT KONFIGURATION ---
    # Geladene Projektdaten auslesen
    loaded = st.session_state.get('loaded_project', None)
    if loaded:
        default_proj = loaded.get('projekt', 'Headquarter')
        default_kunde = loaded.get('kunde', 'Beispiel AG')
        default_bearb = loaded.get('bearbeiter', 'Michael Schaepers')
        st.info(f"📂 Geladenes Projekt: {default_proj}")
    else:
        default_proj = 'Headquarter'
        default_kunde = 'Beispiel AG'
        default_bearb = 'Michael Schaepers'
    
    with st.expander("⚙️ PROJEKT-KONFIGURATION & BEARBEITER", expanded=True):
        px1, px2, px3, px4 = st.columns(4)
        proj_name  = px1.text_input("PROJEKT", default_proj)
        kunde_name = px2.text_input("KUNDE", default_kunde)
        bearbeiter = px3.text_input("BEARBEITER", default_bearb)
        firma      = px4.text_input("FIRMA", "°coolsulting")
        
        # Liefertermin-Feld
        px5, px6 = st.columns([1, 3])
        liefertermin = px5.date_input(
            "📅 LIEFERTERMIN",
            value=None,
            help="Optionaler Liefertermin für Technikübergabe-PDF"
        )
        if liefertermin:
            liefertermin_str = liefertermin.strftime("%d.%m.%Y")
        else:
            liefertermin_str = "—"
    
    # Partner-Firma für PDFs (aus Input übernehmen)
    partner_firma = firma
    
    # --- GEBÄUDE-PARAMETER ---
    st.markdown('<div class="section-header">🏢 Gebäude-Parameter</div>', unsafe_allow_html=True)
    gp1, gp2, gp3 = st.columns(3)
    bau_std = gp1.selectbox("GEBAEUDE-STANDARD", 
                             ["Altbau", "Bestand", "Neubau (GEG)", "Passivhaus"], index=0)
    bau_m   = gp2.selectbox("GEBAEUDE-MASSE", 
                             ["Schwer (Beton/Stein)", "Mittel (Ziegel/Holz-Beton)", "Leicht (Holz/Trockenbau)"], 
                             index=1)
    raumhoehe = gp3.number_input("RAUMHOEHE [m]", 2.0, 6.0, 2.5, step=0.1)
    
    # --- ZONEN KONFIGURATION ---
    st.markdown('<div class="section-header">🏠 Zonen-Konfiguration (bis 5 Räume)</div>', 
                unsafe_allow_html=True)
    
    tabs = st.tabs([f"ZONE {i+1}" for i in range(5)])
    def_ori = ["SUED", "SUED", "SUED", "SUED", "SUED"]
    
    # Ergebnis-Container
    g_sums = {k: np.zeros(24) for k in ["VDI_N", "VDI_A", "PRAK", "RECK", "KLTS", "KI"]}
    individual_profiles = []
    room_results        = []
    room_inputs_list    = []
    samsung_recs        = []
    
    for i, tab in enumerate(tabs):
        with tab:
            st.markdown('<div class="card">', unsafe_allow_html=True)
            
            rc1, rc2, rc3, rc4 = st.columns(4)
            r_name  = rc1.text_input("Bezeichnung", f"Raum {i+1}", key=f"rn{i}")
            area    = rc2.number_input("Flaeche [m²]", 5.0, 500.0, 50.0 if i==0 else 20.0, key=f"ar{i}")
            win     = rc3.number_input("Fenster [m²]", 0.0, 150.0, 2.4, key=f"wi{i}")
            orient  = rc4.selectbox("Ausrichtung", list(SOLAR_DB.keys()),
                                     index=list(SOLAR_DB.keys()).index(def_ori[i]), key=f"or{i}")
            
            rc5, rc6, rc7, rc8 = st.columns(4)
            glass   = rc5.selectbox("Glas", ["Einfach", "Doppel", "Dreifach", "Sonnenschutz"], 
                                     index=0, key=f"gl{i}")
            shade   = rc6.selectbox("Sonnenschutz", 
                                     ["Keine", "Vorhang (Innen)", "Raffstore (Aussen)", "Rollladen"], 
                                     index=1, key=f"sh{i}")
            pers    = rc7.slider("Personen", 0, 15, 2, key=f"pe{i}")
            tech    = rc8.number_input("Technik [W]", 0.0, 10000.0, 200.0, key=f"te{i}")
            
            st.markdown('</div>', unsafe_allow_html=True)
            
            # ---- BERECHNUNGEN ----
            u, g, fc = get_phys_constants(bau_std, glass, shade)
            
            c_reck   = calc_recknagel(area, orient, bau_std, glass, shade, pers, tech, win)
            c_vdi_a  = calc_vdi_alt(c_reck)
            c_vdi_n  = calc_vdi_neu(area, orient, bau_std, glass, shade, pers, tech, win, bau_m)
            c_prak   = calc_praktiker(area, orient, bau_std, glass, shade, pers, tech)
            c_klts   = calc_kaltluftsee(area, orient, bau_std, glass, shade, pers, tech, win, bau_m, raumhoehe)
            c_ki     = calc_ki_hybrid(area, orient, bau_std, glass, shade, pers, tech, win, bau_m)
            
            # Summierung
            g_sums["RECK"]  += c_reck
            g_sums["VDI_A"] += c_vdi_a
            g_sums["VDI_N"] += c_vdi_n
            g_sums["PRAK"]  += c_prak
            g_sums["KLTS"]  += c_klts
            g_sums["KI"]    += c_ki
            
            individual_profiles.append({
                "name":  r_name,
                "reck":  c_reck,
                "vdi_a": c_vdi_a,
                "vdi_n": c_vdi_n,
                "prak":  c_prak,
                "klts":  c_klts,
                "ki":    c_ki,
            })
            
            # Samsung Empfehlung (auf Basis VDI Neu)
            peak_vdi = int(np.max(c_vdi_n))
            primary, alt = find_samsung_device(peak_vdi)
            samsung_recs.append({"zone": r_name, "primary": primary, "alt": alt, "peak_w": peak_vdi})
            
            # Eingabedaten erfassen
            room_inputs_list.append({
                "name":        r_name,
                "flaeche":     area,
                "hoehe":       raumhoehe,
                "personen":    pers,
                "fenster":     win,
                "orientierung": orient,
                "nutzung":     glass,
                "u_wert":      u,
            })
            room_results.append({
                "ZONE":       r_name,
                "VDI NEU":    peak_vdi,
                "VDI ALT":    int(np.max(c_vdi_a)),
                "RECKNAGEL":  int(np.max(c_reck)),
                "PRAKTIKER":  int(np.max(c_prak)),
                "KALTLUFTSEE":int(np.max(c_klts)),
                "KI HYBRID":  int(np.max(c_ki)),
            })
    
    # ==========================================
    # ERGEBNIS-MATRIX
    # ==========================================
    st.markdown('<div class="matrix-wrapper">', unsafe_allow_html=True)
    st.markdown('<div class="matrix-title">📊 Ergebnis-Matrix [Watt] — 6 Methoden</div>', 
                unsafe_allow_html=True)
    
    df_res = pd.DataFrame(room_results)
    totals = {
        "ZONE":         "GEBAEUDE SIMULTAN-PEAK",
        "VDI NEU":      int(np.max(g_sums["VDI_N"])),
        "VDI ALT":      int(np.max(g_sums["VDI_A"])),
        "RECKNAGEL":    int(np.max(g_sums["RECK"])),
        "PRAKTIKER":    int(np.max(g_sums["PRAK"])),
        "KALTLUFTSEE":  int(np.max(g_sums["KLTS"])),
        "KI HYBRID":    int(np.max(g_sums["KI"])),
    }
    df_res = pd.concat([df_res, pd.DataFrame([totals])], ignore_index=True)
    
    tbl = "<table class='styled-table'><thead><tr>"
    col_map = {
        "ZONE": "Zone",
        "VDI NEU": "VDI 6007 Neu",
        "VDI ALT": "VDI 2078 Alt",
        "RECKNAGEL": "Recknagel",
        "PRAKTIKER": "Praktiker",
        "KALTLUFTSEE": "Kaltluftsee",
        "KI HYBRID": "KI-Hybrid"
    }
    for col in df_res.columns:
        tbl += f"<th>{col_map.get(col, col)}</th>"
    tbl += "</tr></thead><tbody>"
    
    for _, row in df_res.iterrows():
        is_total = "SIMULTAN" in str(row["ZONE"])
        cls = " class='total-row'" if is_total else ""
        tbl += f"<tr{cls}>"
        for val in row:
            if isinstance(val, (int, float)) and not isinstance(val, bool):
                tbl += f"<td>{val:,}</td>"
            else:
                tbl += f"<td>{val}</td>"
        tbl += "</tr>"
    tbl += "</tbody></table>"
    
    st.markdown(tbl, unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)
    
    # ==========================================
    # SAMSUNG WIND-FREE EMPFEHLUNGEN
    # ==========================================
    st.markdown(
        '<div class="section-header">❄️ Samsung Wind-Free Wandgeräte — Geräteauslegung</div>',
        unsafe_allow_html=True
    )

    # --- Geräte je Methode berechnen ---
    method_peaks = {}
    for i in range(5):
        r = room_results[i]
        method_peaks[i] = {
            "VDI NEU":     r["VDI NEU"],
            "VDI ALT":     r["VDI ALT"],
            "RECKNAGEL":   r["RECKNAGEL"],
            "PRAKTIKER":   r["PRAKTIKER"],
            "KALTLUFTSEE": r["KALTLUFTSEE"],
            "KI HYBRID":   r["KI HYBRID"],
        }

    # Safety factor
    METHOD_SAFETY = {
        "VDI NEU":     1.10,
        "VDI ALT":     1.10,
        "RECKNAGEL":   1.10,
        "PRAKTIKER":   1.10,
        "KALTLUFTSEE": 1.10,
        "KI HYBRID":   1.10,
    }

    # Serien-Auswahl PER SPALTE (über der Vergleichstabelle)
    _serien_namen = list(SAMSUNG_SERIEN.keys())
    _wand_serien  = [s for s in _serien_namen if s not in
                     ("Mini-Kassette 620x620","1-Weg-Kassette","Kanaleinbau","Standtruhe")]
    _alle_serien  = _serien_namen  # inkl. Kassette etc.

    st.markdown(
        "<div style='font-size:10px;font-weight:700;color:rgba(255,255,255,0.6);"
        "text-transform:uppercase;letter-spacing:1px;margin-bottom:4px;'>"
        "SAMSUNG SERIE je Zone (für Vergleichstabelle + Empfehlung)</div>",
        unsafe_allow_html=True
    )
    serie_cols = st.columns([2.2, 1, 1, 1, 1, 1])
    serie_cols[0].markdown(
        "<div style='font-size:10px;color:#aaa;padding-top:8px;'>Methode</div>",
        unsafe_allow_html=True
    )
    zone_serien = []
    zone_names  = [room_results[i]["ZONE"] for i in range(5)]
    for ci in range(5):
        s = serie_cols[ci+1].selectbox(
            zone_names[ci],
            _alle_serien,
            index=_alle_serien.index(SAMSUNG_DEFAULT_SERIE),
            key=f"serie_col{ci}",
            label_visibility="visible",
            format_func=lambda x: SERIE_SHORT.get(x, x)
        )
        zone_serien.append(s)

    # selected_serie = erste Zone (für device_label Fallback)
    selected_serie = zone_serien[0]

    def device_label(peak_w, safety=1.10, serie=None):
        """Gerätekurzbezeichnung aus gewählter Serie"""
        if serie is None:
            serie = selected_serie
        req = (peak_w * safety) / 1000.0
        db  = SAMSUNG_SERIEN.get(serie, SAMSUNG_SERIEN[SAMSUNG_DEFAULT_SERIE])
        for kw in sorted(db.keys()):
            if kw >= req:
                d = db[kw]
                return kw, f"{kw:.1f}kW", d["art_nr"], d["preis"]
        kw = sorted(db.keys())[-1]
        d  = db[kw]
        return kw, f"{kw:.1f}kW", d["art_nr"], d["preis"]
        r = room_results[i]
        method_peaks[i] = {
            "VDI NEU":     r["VDI NEU"],
            "VDI ALT":     r["VDI ALT"],
            "RECKNAGEL":   r["RECKNAGEL"],
            "PRAKTIKER":   r["PRAKTIKER"],
            "KALTLUFTSEE": r["KALTLUFTSEE"],
            "KI HYBRID":   r["KI HYBRID"],
        }

    # Safety factor: uniform 10% for all methods per user spec
    METHOD_SAFETY = {
        "VDI NEU":     1.10,
        "VDI ALT":     1.10,
        "RECKNAGEL":   1.10,
        "PRAKTIKER":   1.10,
        "KALTLUFTSEE": 1.10,
        "KI HYBRID":   1.10,
    }

    def device_label(peak_w, safety=1.10, serie=None):
        """Gibt Gerätekurzbezeichnung für Peak zurück (aus gewählter Serie)"""
        if serie is None:
            serie = selected_serie
        req = (peak_w * safety) / 1000.0
        db = SAMSUNG_SERIEN.get(serie, SAMSUNG_SERIEN[SAMSUNG_DEFAULT_SERIE])
        for kw in sorted(db.keys()):
            if kw >= req:
                d = db[kw]
                return kw, f"{kw:.1f}kW", d["art_nr"], d["preis"]
        kw = sorted(db.keys())[-1]
        d = db[kw]
        return kw, f"{kw:.1f}kW", d["art_nr"], d["preis"]

    # Farben je Methode
    METHOD_COLORS = {
        "VDI NEU":     ("#1a6fa8", "#e8f4fc"),
        "VDI ALT":     ("#b7770d", "#fef9e7"),
        "RECKNAGEL":   ("#3c3c3b", "#f4f4f4"),
        "PRAKTIKER":   ("#c0392b", "#fdf2f2"),
        "KALTLUFTSEE": ("#7d3c98", "#f9f0ff"),
        "KI HYBRID":   ("#1a7a5e", "#e8faf4"),
    }
    METHOD_LABELS = {
        "VDI NEU":     "VDI 6007 Neu",
        "VDI ALT":     "VDI 2078 Alt",
        "RECKNAGEL":   "Recknagel",
        "PRAKTIKER":   "Praktiker ★ OFFIZ. EMPF.",
        "KALTLUFTSEE": "Kaltluftsee",
        "KI HYBRID":   "KI-Hybrid",
    }

    zone_names = [room_results[i]["ZONE"] for i in range(5)]

    # --- Tabelle: eine Zeile je Methode ---
    st.markdown("""
    <div class="matrix-wrapper" style="padding:20px 25px;">
    <div style="font-size:13px; font-weight:700; color:#3C3C3B; text-transform:uppercase;
                letter-spacing:2px; margin-bottom:16px; border-bottom:2px solid #36A9E1;
                padding-bottom:8px;">
        Geraetevorschlag je Berechnungsmethode (Samsung Wind-Free Standard | +10% Norm-Zuschlag)
    </div>
    """, unsafe_allow_html=True)

    # Header-Zeile
    hdr_cols = st.columns([2.2, 1, 1, 1, 1, 1])
    hdr_cols[0].markdown(
        "<div style='font-size:10px;font-weight:700;color:#3C3C3B;"
        "text-transform:uppercase;letter-spacing:1px;padding:4px 0;'>Methode</div>",
        unsafe_allow_html=True
    )
    for ci, zn in enumerate(zone_names):
        hdr_cols[ci+1].markdown(
            f"<div style='font-size:10px;font-weight:700;color:#3C3C3B;"
            f"text-transform:uppercase;letter-spacing:1px;text-align:center;"
            f"padding:4px 0;'>{zn}</div>",
            unsafe_allow_html=True
        )

    for mkey, mlabel in METHOD_LABELS.items():
        dark_color, light_color = METHOD_COLORS[mkey]
        is_official = mkey == "PRAKTIKER"

        border = f"2px solid {dark_color}" if is_official else f"1px solid {dark_color}40"
        bg     = light_color
        shadow = "box-shadow:0 2px 8px rgba(0,0,0,0.12);" if is_official else ""

        row_cols = st.columns([2.2, 1, 1, 1, 1, 1])

        # Methoden-Label
        star = " ⭐" if is_official else ""
        row_cols[0].markdown(
            f"<div style='background:{bg};border:{border};border-radius:8px;"
            f"padding:8px 10px;{shadow}margin:2px 0;'>"
            f"<div style='font-size:11px;font-weight:700;color:{dark_color};"
            f"text-transform:uppercase;letter-spacing:0.5px;'>{mlabel}{star}</div>"
            f"</div>",
            unsafe_allow_html=True
        )

        # Gerät je Zone — mit jeweiliger Zonen-Serie
        method_safety_val = METHOD_SAFETY.get(mkey, 1.10)
        for ci in range(5):
            peak_w = method_peaks[ci][mkey]
            kw, short, art_nr, preis = device_label(peak_w, safety=method_safety_val,
                                                     serie=zone_serien[ci])
            row_cols[ci+1].markdown(
                f"<div style='background:{bg};border:{border};border-radius:8px;"
                f"padding:6px 8px;{shadow}margin:2px 0;text-align:center;'>"
                f"<div style='font-size:12px;font-weight:700;color:{dark_color};'>{short}</div>"
                f"<div style='font-size:10px;color:#666;margin-top:2px;'>"
                f"{kw:.1f} kW | {peak_w:,} W</div>"
                f"<div style='font-size:9px;color:#999;'>{art_nr}</div>"
                f"</div>",
                unsafe_allow_html=True
            )

    st.markdown("</div>", unsafe_allow_html=True)

    # --- VDI 2078 ALT EMPFEHLUNG (automatisch, Wind-Free Standard) ---
    st.markdown("""
    <div style="margin-top:24px; margin-bottom:8px;">
        <span style="font-size:12px;font-weight:700;color:rgba(255,255,255,0.8);
                     text-transform:uppercase;letter-spacing:2px;">
            VDI 2078 Alt Empfehlung (automatisch | Wind-Free Standard)
        </span>
    </div>""", unsafe_allow_html=True)

    green_cols = st.columns(5)
    for i, gcol in enumerate(green_cols):
        peak_vdi_alt = method_peaks[i]["VDI ALT"]
        kw_rec, _, art_rec, preis_rec = device_label(
            peak_vdi_alt, safety=1.10, serie="Wind-Free Standard"
        )
        gcol.markdown(f"""
            <div style="background:linear-gradient(135deg,#1565c0,#1e88e5);
                        border-radius:12px;padding:14px 12px;color:white;margin-bottom:4px;
                        box-shadow:0 2px 8px rgba(21,101,192,0.4);">
                <div style="font-size:9px;font-weight:800;opacity:0.85;
                            letter-spacing:1px;text-transform:uppercase;">⭐ EMPFEHLUNG VDI 2078 Alt</div>
                <div style="font-size:15px;font-weight:700;margin:5px 0 3px 0;
                            letter-spacing:-0.3px;">WF Standard {kw_rec:.1f} kW</div>
                <div style="font-size:10px;opacity:0.92;line-height:1.7;">
                    📦 {art_rec}<br>
                    ❄️ {kw_rec:.1f} kW &nbsp;|&nbsp; SEER ~6.2<br>
                    💶 {preis_rec:.0f} EUR LP
                </div>
            </div>""", unsafe_allow_html=True)

    # ==========================================
    # EDITIERBARE FINALE GERÄTEAUSWAHL
    # ==========================================
    st.markdown("""
    <div style="margin-top:16px; margin-bottom:8px;">
        <span style="font-size:14px;font-weight:700;color:white;text-transform:uppercase;
                     letter-spacing:2px;">✏️ Finale Geräteauswahl</span>
        <span style="font-size:11px;color:rgba(255,255,255,0.7);margin-left:10px;">
            (editierbar — Innen- &amp; Außengerät | Serie pro Spalte wählbar)
        </span>
    </div>
    """, unsafe_allow_html=True)

    # ==========================================
    # RAC-AG Datenbank: je IG-Serie passende Außengeräte
    # ==========================================
    # FJM Multi-Außengeräte (für alle FJM IG-Typen: Kassette, Kanal, Standtruhe, WF-Serien)
    FJM_AG = {
        4.0:  {"art_nr": "AJ040TXJ2KG/EU", "bez": "FJM Multi AG  4,0 kW",   "preis": 2347},
        5.0:  {"art_nr": "AJ050TXJ2KG/EU", "bez": "FJM Multi AG  5,0 kW",   "preis": 2706},
        5.2:  {"art_nr": "AJ052TXJ3KG/EU", "bez": "FJM Multi AG  5,2 kW",   "preis": 3061},
        6.8:  {"art_nr": "AJ068TXJ3KG/EU", "bez": "FJM Multi AG  6,8 kW",   "preis": 3548},
        8.0:  {"art_nr": "AJ080TXJ4KG/EU", "bez": "FJM Multi AG  8,0 kW",   "preis": 4494},
        10.0: {"art_nr": "AJ100TXJ5KG/EU", "bez": "FJM Multi AG 10,0 kW",   "preis": 5533},
    }

    # RAC Single-Split-AGs je IG-Serie (IG-Artnr Präfix → passende AG-Liste)
    RAC_AG_BY_SERIE = {
        "Airise Living": [
            (2.5, "AR50F09C1BHX/EU",  "RAC AG Airise Living 2,5 kW",  984),
            (3.5, "AR50F12C1BHX/EU",  "RAC AG Airise Living 3,5 kW", 1019),
            (5.0, "AR50F18C1BHX/EU",  "RAC AG Airise Living 5,0 kW", 1564),
            (6.5, "AR50F24C1BHX/EU",  "RAC AG Airise Living 6,5 kW", 2056),
        ],
        "Wind-Free Standard": [
            (2.5, "AR60F09C1AWX/EU",  "RAC AG WF Standard 2,5 kW",   1362),
            (3.5, "AR60F12C1AWX/EU",  "RAC AG WF Standard 3,5 kW",   1540),  # ca.
            (5.0, "AR60F18C1AWX/EU",  "RAC AG WF Standard 5,0 kW",   1900),
        ],
        "Wind-Free Exklusiv": [
            (2.0, "AR70F07C1AWX/EU",  "RAC AG WF Exklusiv 2,0 kW",   1448),
            (2.5, "AR70F09C1AWX/EU",  "RAC AG WF Exklusiv 2,5 kW",   1446),
            (3.5, "AR70F12C1AWX/EU",  "RAC AG WF Exklusiv 3,5 kW",   1540),
            (4.3, "AR70F15C1AWX/EU",  "RAC AG WF Exklusiv 4,3 kW",   2135),
            (5.0, "AR70F18C1AWX/EU",  "RAC AG WF Exklusiv 5,0 kW",   2106),
            (6.5, "AR70F24C1AWX/EU",  "RAC AG WF Exklusiv 6,5 kW",   2761),
        ],
        "Wind-Free Exklusiv Black": [
            (2.0, "AR70F07C1AWX/EU",  "RAC AG WF Exklusiv 2,0 kW",   1448),
            (2.5, "AR70F09C1AWX/EU",  "RAC AG WF Exklusiv 2,5 kW",   1446),
            (3.5, "AR70F12C1AWX/EU",  "RAC AG WF Exklusiv 3,5 kW",   1540),
        ],
        "Wind-Free Exklusiv-Premiere": [
            (2.0, "AR70H07C1AWX/EU",  "RAC AG WF Exkl.-Prem. 2,0 kW", 1604),
            (2.5, "AR70H09C1AWX/EU",  "RAC AG WF Exkl.-Prem. 2,5 kW", 1688),
            (3.5, "AR70H12C1AWX/EU",  "RAC AG WF Exkl.-Prem. 3,5 kW", 1872),
            (4.3, "AR70H15C1AWX/EU",  "RAC AG WF Exkl.-Prem. 4,3 kW", 2360),
            (5.0, "AR70H18C1AWX/EU",  "RAC AG WF Exkl.-Prem. 5,0 kW", 2708),
            (6.5, "AR70H24C1AWX/EU",  "RAC AG WF Exkl.-Prem. 6,5 kW", 3556),
        ],
        "Wind-Free Exklusiv-Premiere Black": [
            (2.0, "AR70H07C1AWX/EU",  "RAC AG WF Exkl.-Prem. 2,0 kW", 1604),
            (2.5, "AR70H09C1AWX/EU",  "RAC AG WF Exkl.-Prem. 2,5 kW", 1688),
            (3.5, "AR70H12C1AWX/EU",  "RAC AG WF Exkl.-Prem. 3,5 kW", 1872),
        ],
        "Wind-Free Elite": [
            (2.5, "AR70F09CAAWKX/EU", "RAC AG WF Elite 2,5 kW",       1752),
            (3.5, "AR70F12CAAWKX/EU", "RAC AG WF Elite 3,5 kW",       1944),
        ],
        "Wind-Free Elite-Premiere Plus": [
            (2.5, "AR70H09CAAWX/EU",  "RAC AG WF Elite-Prem.Plus 2,5 kW", 1824),
            (3.5, "AR70H12CAAWX/EU",  "RAC AG WF Elite-Prem.Plus 3,5 kW", 2024),
        ],
        "Wind-Free Elite-Premiere Plus Black": [
            (2.5, "AR70H09CAAWX/EU",  "RAC AG WF Elite-Prem.Plus 2,5 kW", 1824),
            (3.5, "AR70H12CAAWX/EU",  "RAC AG WF Elite-Prem.Plus 3,5 kW", 2024),
        ],
    }
    # FJM IG-Typen → immer FJM Multi AG
    FJM_IG_SERIEN = {"Mini-Kassette 620x620", "1-Weg-Kassette", "Kanaleinbau", "Standtruhe"}
    # FJM Wandgeräte-Serien → können RAC AG oder FJM AG bekommen (Umschalter)
    FJM_WAND_SERIEN = {"Wind-Free Standard","Wind-Free Exklusiv","Wind-Free Exklusiv Black",
                       "Wind-Free Exklusiv-Premiere","Wind-Free Exklusiv-Premiere Black",
                       "Wind-Free Elite","Wind-Free Elite-Premiere Plus",
                       "Wind-Free Elite-Premiere Plus Black","Airise Living"}

    # IG-Optionen aufbauen
    IG_OPTIONS = [(0.0, "N.V.", "— nicht vorhanden —")]
    for sname, sdata in SAMSUNG_SERIEN.items():
        for kw in sorted(sdata.keys()):
            d = sdata[kw]
            label = f"{d['art_nr']}  |  {kw:.1f} kW  |  {sname}  |  {d['preis']:.0f} EUR"
            IG_OPTIONS.append((kw, sname, label))
    IG_KEYS   = list(range(len(IG_OPTIONS)))
    # Kurze Labels für das Dropdown (WF statt Wind-Free, Serie verkürzt)
    def _ig_short_label(lbl):
        if lbl == "— nicht vorhanden —": return lbl
        parts = [p.strip() for p in lbl.split("|")]
        if len(parts) >= 3:
            artnr = parts[0]
            kw_p  = parts[1]
            serie_short = SERIE_SHORT.get(parts[2].strip(), parts[2].strip())
            preis = parts[-1]
            return f"{artnr}  |  {kw_p}  |  {serie_short}  |  {preis}"
        return lbl
    IG_LABELS = {k: _ig_short_label(IG_OPTIONS[k][2]) for k in IG_KEYS}

    selected_hw    = []
    selected_hw_ag = []
    selected_ig_artnr = []  # Neu: IG Art.-Nr. speichern
    final_cols = st.columns(5)

    for i, col in enumerate(final_cols):
        with col:
            r_name    = zone_names[i]
            prak_peak = method_peaks[i]["PRAKTIKER"]
            vdi_peak  = method_peaks[i]["VDI NEU"]
            z_serie   = zone_serien[i]

            # Default IG: Praktiker-Empfehlung in Zonen-Serie
            prak_kw, _, _, _ = device_label(prak_peak, safety=1.10, serie=z_serie)
            def_ig_idx = 0
            for ig_idx, (kw, sname, _) in enumerate(IG_OPTIONS):
                if sname == z_serie and kw == prak_kw:
                    def_ig_idx = ig_idx
                    break

            # Info-Box - alle 6 Methoden berechnen
            vdi_kw,   _, _, _ = device_label(vdi_peak, safety=1.10, serie=z_serie)
            vdi_a_kw, _, _, _ = device_label(method_peaks[i]["VDI ALT"], safety=1.10, serie=z_serie)
            reck_kw,  _, _, _ = device_label(method_peaks[i]["RECKNAGEL"], safety=1.10, serie=z_serie)
            klts_kw,  _, _, _ = device_label(method_peaks[i]["KALTLUFTSEE"], safety=1.10, serie=z_serie)
            ki_kw,    _, _, _ = device_label(method_peaks[i]["KI HYBRID"], safety=1.10, serie=z_serie)
            st.markdown(
                f"<div style='background:rgba(255,255,255,0.12);border:1px solid "
                f"rgba(255,255,255,0.3);border-radius:10px;padding:10px;margin-bottom:6px;'>"
                f"<div style='font-size:11px;font-weight:700;color:white;"
                f"text-transform:uppercase;'>{r_name}</div>"
                f"<div style='font-size:9px;color:rgba(255,255,255,0.75);margin-top:4px;line-height:1.4;'>"
                f"VDI Neu: {vdi_kw:.1f} kW &nbsp;|&nbsp; "
                f"VDI Alt: {reck_kw:.1f} kW<br/>"
                f"Recknagel: {vdi_a_kw:.1f} kW &nbsp;|&nbsp; "
                f"Praktiker ★: {prak_kw:.1f} kW<br/>"
                f"Kaltluftsee: {klts_kw:.1f} kW &nbsp;|&nbsp; "
                f"KI-Hybrid: {ki_kw:.1f} kW"
                f"</div></div>",
                unsafe_allow_html=True
            )

            # ❄️ Innengerät
            st.markdown("<div style='font-size:9px;color:rgba(255,255,255,0.6);"
                        "margin-bottom:2px;'>❄️ INNENGERÄT</div>", unsafe_allow_html=True)
            ig_val = st.selectbox(
                f"IG {r_name}", IG_KEYS, index=def_ig_idx, key=f"hw{i}",
                format_func=lambda x: IG_LABELS[x],
                label_visibility="collapsed"
            )
            ig_kw    = IG_OPTIONS[ig_val][0]
            ig_serie = IG_OPTIONS[ig_val][1]
            selected_hw.append(ig_kw)
            
            # IG Art.-Nr. extrahieren
            ig_artnr = '—'
            if ig_kw > 0 and ig_serie in SAMSUNG_SERIEN and ig_kw in SAMSUNG_SERIEN[ig_serie]:
                ig_artnr = SAMSUNG_SERIEN[ig_serie][ig_kw]['art_nr']
            selected_ig_artnr.append(ig_artnr)

            # 🔀 RAC / FJM Umschalter (nur bei Wandgeräten; Kassette/Kanal/Truhe → immer FJM)
            is_fjm_ig  = ig_serie in FJM_IG_SERIEN
            if is_fjm_ig:
                ag_modus = "FJM Multi"
            else:
                ag_modus = st.radio(
                    f"AG-Typ {r_name}",
                    ["RAC Single-Split", "FJM Multi"],
                    index=1, key=f"agm{i}",
                    horizontal=True,
                    label_visibility="collapsed"
                )

            # 🌡️ Außengerät — gefiltert nach Modus
            st.markdown("<div style='font-size:9px;color:rgba(255,255,255,0.6);"
                        "margin-top:4px;margin-bottom:2px;'>🌡️ AUSSENGERÄT</div>",
                        unsafe_allow_html=True)

            if ag_modus == "FJM Multi":
                fjm_keys_raw = list(FJM_AG.keys())
                # N.V. vorne
                fjm_keys   = ["NV"] + fjm_keys_raw
                fjm_labels = {"NV": "— nicht vorhanden —"}
                fjm_labels.update({k: f"{FJM_AG[k]['art_nr']}  |  {FJM_AG[k]['bez']}  |  {FJM_AG[k]['preis']:,} EUR"
                                   for k in fjm_keys_raw})
                # Default: Zone 1 = AJ100, Zone 2-5 = N.V. (auch wenn IG > 0)
                def_fjm = 0  # Default = N.V.
                if i == 0:
                    # Zone 1: Default = AJ100 (10.0kW index 6) wenn kein IG, sonst passendes AG
                    if ig_kw == 0:
                        def_fjm = 6  # AJ100
                    else:
                        # IG ausgewählt: kleinstes AG >= ig_kw
                        for fi, fk in enumerate(fjm_keys_raw):
                            if fk >= ig_kw:
                                def_fjm = fi + 1
                                break
                # Zone 2-5: immer N.V. (def_fjm bleibt 0)
                ag_sel = st.selectbox(
                    f"AG {r_name}", fjm_keys, index=def_fjm, key=f"ag{i}",
                    format_func=lambda x, m=fjm_labels: m[x],
                    label_visibility="collapsed"
                )
                if ag_sel == "NV":
                    selected_hw_ag.append(("FJM", 0, "N.V."))
                else:
                    selected_hw_ag.append(("FJM", ag_sel, FJM_AG[ag_sel]["art_nr"]))
            else:
                # RAC: passendes AG zur IG-Serie
                rac_list_raw = RAC_AG_BY_SERIE.get(ig_serie,
                               RAC_AG_BY_SERIE.get("Wind-Free Exklusiv", []))
                if not rac_list_raw:
                    rac_list_raw = [(2.5, "—", "Kein passendes RAC AG", 0)]
                # N.V. vorne einfügen
                rac_list   = [(0, "NV", "— nicht vorhanden —", 0)] + list(rac_list_raw)
                rac_keys   = list(range(len(rac_list)))
                rac_labels = {0: "— nicht vorhanden —"}
                rac_labels.update({
                    k+1: f"{rac_list_raw[k][1]}  |  {rac_list_raw[k][2]}  |  {rac_list_raw[k][3]:,} EUR"
                    for k in range(len(rac_list_raw))
                })
                # Default: kleinstes RAC AG >= ig_kw (oder 0 = N.V. wenn ig_kw=0)
                def_rac = 0
                if ig_kw > 0:
                    for ri, (rkw, _, _, _) in enumerate(rac_list_raw):
                        if rkw >= ig_kw:
                            def_rac = ri + 1  # +1 wegen N.V. vorne
                            break
                ag_sel = st.selectbox(
                    f"AG {r_name}", rac_keys, index=def_rac, key=f"ag{i}",
                    format_func=lambda x, m=rac_labels: m[x],
                    label_visibility="collapsed"
                )
                if ag_sel == 0:
                    selected_hw_ag.append(("RAC", 0, "N.V."))
                else:
                    entry = rac_list_raw[ag_sel - 1]
                    selected_hw_ag.append(("RAC", entry[0], entry[1]))

    # Gesamtleistung + Preis (IG + AG)
    total_kw = sum(selected_hw)
    _ig_preis_map = {}
    for kw, sname, label in IG_OPTIONS:
        if kw > 0 and "|" in label:
            try:
                p = label.split("|")[-1].strip().replace(" EUR","").replace(".","").replace(",",".")
                _ig_preis_map[(kw, sname)] = float(p)
            except Exception:
                pass
    total_preis_ig = 0.0
    for idx2, kw in enumerate(selected_hw):
        if kw > 0:
            ig_idx = next((j for j,(k,s,_) in enumerate(IG_OPTIONS)
                           if k == kw and s == zone_serien[idx2]), None)
            if ig_idx is not None:
                try:
                    p = IG_OPTIONS[ig_idx][2].split("|")[-1].strip().replace(" EUR","").replace(".","").replace(",",".")
                    total_preis_ig += float(p)
                except Exception:
                    pass

    # AG-Preise addieren
    total_preis_ag = 0.0
    for ag_inf in selected_hw_ag:
        if isinstance(ag_inf, (list, tuple)) and len(ag_inf) > 1:
            ag_typ  = ag_inf[0]
            ag_kw   = ag_inf[1]
            if ag_kw and ag_kw != 0 and ag_typ != "—":
                try:
                    if ag_typ == "FJM" and ag_kw in FJM_AG:
                        total_preis_ag += FJM_AG[ag_kw]["preis"]
                    elif ag_typ == "RAC":
                        # Suche in RAC_AG_BY_SERIE
                        for serie_list in RAC_AG_BY_SERIE.values():
                            for rkw, rartnr, rbez, rpreis in serie_list:
                                if abs(rkw - float(ag_kw)) < 0.01:
                                    total_preis_ag += rpreis
                                    break
                except Exception:
                    pass

    total_preis = total_preis_ig + total_preis_ag

    st.markdown(f"""
    <div class="card" style="margin-top:16px;">
        <div style="display:flex;justify-content:space-between;align-items:center;">
            <div>
                <div style="font-size:11px;font-weight:700;color:#888;
                            text-transform:uppercase;letter-spacing:1px;">
                    Installierte Gesamtleistung
                </div>
                <div style="font-size:32px;font-weight:700;color:{CI_BLUE};
                            line-height:1.1;">{total_kw:.1f} kW</div>
                <div style="font-size:11px;color:#aaa;margin-top:2px;">
                    Samsung Innengeräte — Finale Auswahl
                </div>
            </div>
            <div style="text-align:right;">
                <div style="font-size:11px;font-weight:700;color:#888;
                            text-transform:uppercase;letter-spacing:1px;">
                    Listenpreis Geräte (netto)
                </div>
                <div style="font-size:28px;font-weight:700;color:{CI_GRAY};
                            line-height:1.1;">{format(int(total_preis), ",.0f").replace(",", ".")} EUR</div>
                <div style="font-size:10px;color:#bbb;margin-top:2px;">
                    IG: {format(int(total_preis_ig), ",.0f").replace(",", ".")} EUR &nbsp;|&nbsp; AG: {format(int(total_preis_ag), ",.0f").replace(",", ".")} EUR &nbsp;|&nbsp; zzgl. MwSt.
                </div>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)


    # ==========================================
    # GRÜNE KARTEN — FINALE GERÄTEAUSWAHL
    # ==========================================
    st.markdown("""
    <div style="margin-top:20px;margin-bottom:8px;">
        <span style="font-size:12px;font-weight:700;color:rgba(255,255,255,0.8);
                     text-transform:uppercase;letter-spacing:2px;">
            ✅ Finale Geräte — Auswahl Zusammenfassung
        </span>
    </div>""", unsafe_allow_html=True)

    finale_cards = st.columns(5)
    for i, fcol in enumerate(finale_cards):
        ig_kw_final = selected_hw[i]
        ig_artnr_final = "—"
        ig_preis_final = 0.0
        ig_serie_final = zone_serien[i]
        for kw_o, sn_o, lbl_o in IG_OPTIONS:
            if kw_o == ig_kw_final and sn_o == zone_serien[i]:
                parts = [p.strip() for p in lbl_o.split("|")]
                ig_artnr_final = parts[0] if parts else "—"
                ig_serie_final = parts[2] if len(parts) > 2 else sn_o
                try:
                    ig_preis_final = float(
                        parts[-1].replace(" EUR","").replace(".","").replace(",",".")
                    )
                except Exception:
                    pass
                break
        ag_info  = selected_hw_ag[i] if i < len(selected_hw_ag) else ("—", 0, "—")
        ag_typ   = ag_info[0] if isinstance(ag_info, tuple) else "—"
        ag_artnr = ag_info[2] if isinstance(ag_info, tuple) and len(ag_info)>2 else "—"

        fcol.markdown(f"""
            <div style="background:linear-gradient(135deg,#1b5e20,#388e3c);
                        border-radius:12px;padding:14px 12px;color:white;
                        box-shadow:0 3px 10px rgba(27,94,32,0.5);margin-bottom:4px;">
                <div style="font-size:9px;font-weight:800;opacity:0.8;letter-spacing:1px;
                            text-transform:uppercase;margin-bottom:4px;">
                    ✅ FINALE — {zone_names[i].upper()}</div>
                <div style="font-size:15px;font-weight:700;margin:3px 0;">
                    {"N.V." if ig_kw_final == 0 else f"{ig_kw_final:.1f} kW"}</div>
                <div style="font-size:9px;opacity:0.92;line-height:1.8;">
                    ❄️ IG: {ig_artnr_final}<br>
                    📋 {ig_serie_final}<br>
                    🌡️ AG ({ag_typ}): {ag_artnr}<br>
                    💶 {ig_preis_final:.0f} EUR LP (IG)
                </div>
            </div>""", unsafe_allow_html=True)

    # ==========================================
    # VERGLEICHS-DIAGRAMME
    # ==========================================
    st.markdown('<div class="section-header">📈 Simultan-Trendkurven — Alle Methoden</div>',
                unsafe_allow_html=True)

    _layout_dark = dict(
        template="plotly_white",
        paper_bgcolor='rgba(0,0,0,0)',
        plot_bgcolor='rgba(255,255,255,0.05)',
        font=dict(color="white", family="Arial"),
        height=520,
        legend=dict(orientation="h", y=-0.18, x=0,
                    bgcolor='rgba(0,0,0,0)', font=dict(color="white", size=11)),
        xaxis=dict(title="Tagesstunde [h]", gridcolor='rgba(255,255,255,0.1)', color='white'),
        yaxis=dict(title="Kühllast [W]",    gridcolor='rgba(255,255,255,0.1)', color='white'),
        margin=dict(l=60, r=20, t=20, b=90)
    )
    _layout_light = dict(
        template="plotly_white",
        paper_bgcolor='white',
        plot_bgcolor='#fafafa',
        height=400,
        legend=dict(orientation="h", y=-0.22, bgcolor='rgba(0,0,0,0)'),
        xaxis=dict(title="Stunde [h]"),
        yaxis=dict(title="Kühllast [W]"),
        margin=dict(l=60, r=20, t=50, b=90)
    )

    fig_master = go.Figure()
    for name, data, color, lw, dash in [
        ("PRAKTIKER (Heuristik)", g_sums["PRAK"],  "#E74C3C", 3,   "dot"),
        ("VDI 6007 NEU",          g_sums["VDI_N"], "white",   5,   "solid"),
        ("VDI 2078 ALT",          g_sums["VDI_A"], "#F39C12", 2.5, "dash"),
        ("RECKNAGEL",             g_sums["RECK"],  CI_GRAY,   2,   "longdash"),
        ("KALTLUFTSEE",           g_sums["KLTS"],  "#9B59B6", 2.5, "dashdot"),
        ("KI-HYBRID",             g_sums["KI"],    "#1ABC9C", 3,   "solid"),
    ]:
        fig_master.add_trace(go.Scatter(
            x=HOURS, y=data, name=name,
            line=dict(width=lw, color=color, dash=dash)
        ))
    fig_master.update_layout(**_layout_dark)
    st.plotly_chart(fig_master, use_container_width=True)

    def plot_zones(mode_key, title, total_key):
        fig = go.Figure()
        zone_colors = [CI_BLUE, "#E74C3C", "#2ECC71", "#F39C12", "#9B59B6"]
        for idx2, p in enumerate(individual_profiles):
            fig.add_trace(go.Scatter(
                x=HOURS, y=p[mode_key], name=p["name"],
                line=dict(width=2, color=zone_colors[idx2 % len(zone_colors)], dash='dash'),
                opacity=0.8
            ))
        fig.add_trace(go.Scatter(
            x=HOURS, y=g_sums[total_key], name="GESAMT SIMULTAN",
            line=dict(width=5, color="#3C3C3B")
        ))
        layout = dict(_layout_light)
        layout["title"] = dict(text=title,
                               font=dict(color=CI_GRAY, size=14, family='Arial Black'))
        fig.update_layout(**layout)
        return fig

    st.markdown("<div style='font-size:11px;font-weight:700;color:rgba(255,255,255,0.5);"
                "text-transform:uppercase;letter-spacing:1px;margin:12px 0 4px 0;'>"
                "Einzelzonen-Diagramme</div>", unsafe_allow_html=True)

    c1, c2 = st.columns(2)
    with c1:
        st.plotly_chart(plot_zones("vdi_n", "VDI 6007 Neu — Einzelzonen", "VDI_N"),
                        use_container_width=True)
    with c2:
        st.plotly_chart(plot_zones("vdi_a", "VDI 2078 Alt — Einzelzonen", "VDI_A"),
                        use_container_width=True)
    c3, c4 = st.columns(2)
    with c3:
        st.plotly_chart(plot_zones("prak", "Praktiker — Einzelzonen", "PRAK"),
                        use_container_width=True)
    with c4:
        st.plotly_chart(plot_zones("reck", "Recknagel — Einzelzonen", "RECK"),
                        use_container_width=True)
    c5, c6 = st.columns(2)
    with c5:
        st.plotly_chart(plot_zones("klts", "Kaltluftsee / Quelllüftung", "KLTS"),
                        use_container_width=True)
    with c6:
        st.plotly_chart(plot_zones("ki", "KI-Hybrid (Peak-Shaving)", "KI"),
                        use_container_width=True)

    # ==========================================
    # EXPORT SEKTION
    # ==========================================
    st.markdown('<div class="section-header">📤 Export & Berichte</div>', unsafe_allow_html=True)
    
    exp1, exp2, exp3 = st.columns(3)
    
    # 1. Übergabebericht JSON
    with exp1:
        st.markdown("""
        <div class="card-blue">
            <div style="font-size:11px; font-weight:700; opacity:0.8; margin-bottom:8px">
                🔄 ÜBERGABEBERICHT (coolMATCH)
            </div>
            <div style="font-size:12px">
                Strukturierter JSON-Datensatz für die Übergabe an coolMATCH Kalkulator.
                Enthält alle Zonen, Kühllastwerte und Geräteempfehlungen.
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        if st.button("📋 JSON ÜBERGABEBERICHT", use_container_width=True):
            transfer = build_transfer_report(
                proj_name, kunde_name, bearbeiter, firma,
                room_results, g_sums, samsung_recs, selected_hw, total_kw, selected_hw_ag
            )
            json_str = json.dumps(transfer, indent=2, ensure_ascii=False)
            st.download_button(
                "⬇️ DOWNLOAD JSON",
                data=json_str.encode('utf-8'),
                file_name=f"coolMATH_Transfer_{proj_name}_{datetime.now().strftime('%Y%m%d')}.json",
                mime="application/json",
                key="json_dl"
            )
    
    # 2. Technischer Übergabe-PDF
    with exp2:
        st.markdown("""
        <div class="card-blue">
            <div style="font-size:11px; font-weight:700; opacity:0.8; margin-bottom:8px">
                🔧 TECHNIKÜBERGABE PDF
            </div>
            <div style="font-size:12px">
                Vollständiger technischer Bericht mit allen 6 Methoden, Diagrammen 
                und Hardware-Abgleich für interne Übergabe.
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        if st.button("🔧 TECHNIKÜBERGABE GENERIEREN", use_container_width=True):
            with st.spinner("⚙️ PDF wird erstellt..."):
                try:
                    pdf_bytes = generate_uebergabe_pdf(
                        proj_name, kunde_name, bearbeiter, firma,
                        room_results, g_sums, individual_profiles,
                        samsung_recs, selected_hw, total_kw,
                        selected_hw_ag=selected_hw_ag,
                        room_inputs=room_inputs_list,
                        partner_firma=partner_firma,
                        selected_ig_artnr=selected_ig_artnr,
                        liefertermin=liefertermin_str
                    )
                    st.download_button(
                        "⬇️ DOWNLOAD TECHNIKÜBERGABE",
                        data=pdf_bytes,
                        file_name=f"coolMATH_Uebergabe_{proj_name}_{datetime.now().strftime('%Y%m%d')}.pdf",
                        mime="application/pdf",
                        key="pdf_tech_dl"
                    )
                    st.success("✅ Technikübergabe bereit!")
                except Exception as e:
                    st.error(f"Fehler: {e}")
    
    # 3. Kundenbericht PDF
    with exp3:
        st.markdown("""
        <div class="card-blue">
            <div style="font-size:11px; font-weight:700; opacity:0.8; margin-bottom:8px">
                👔 KUNDENBERICHT PDF
            </div>
            <div style="font-size:12px">
                Professioneller Kundenbericht mit Executive Summary, Methodenvergleich, 
                Geräteempfehlung und Methodik-Erläuterung.
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        if st.button("📄 KUNDENBERICHT GENERIEREN", use_container_width=True):
            with st.spinner("📑 Kundenbericht wird erstellt..."):
                try:
                    pdf_bytes = generate_kunden_pdf(
                        proj_name, kunde_name, bearbeiter, firma,
                        room_results, g_sums, individual_profiles,
                        samsung_recs, selected_hw, total_kw,
                        selected_hw_ag=selected_hw_ag,
                        room_inputs=room_inputs_list,
                        partner_firma=partner_firma,
                        selected_ig_artnr=selected_ig_artnr
                    )
                    st.download_button(
                        "⬇️ DOWNLOAD KUNDENBERICHT",
                        data=pdf_bytes,
                        file_name=f"coolMATH_Kundenbericht_{proj_name}_{datetime.now().strftime('%Y%m%d')}.pdf",
                        mime="application/pdf",
                        key="pdf_kd_dl"
                    )
                    st.success("✅ Kundenbericht bereit!")
                except Exception as e:
                    st.error(f"Fehler: {e}")
    
    # ==========================================
    # WORD-EXPORT + DB-SPEICHERUNG + MONDAY + EXCEL
    # ==========================================
    st.markdown('<div class="section-header">📁 Word-Export | 💾 Projekt Speichern | 📤 Monday.com | 📊 Excel-Anfrage</div>',
                unsafe_allow_html=True)
    wrd1, wrd2, wrd3, wrd4 = st.columns(4)

    with wrd1:
        st.markdown("""<div class="card-blue"><div style="font-size:11px;font-weight:700;opacity:0.8;margin-bottom:8px">
        📄 WORD-BERICHT</div><div style="font-size:12px">Vollständiger Bericht als .docx — editierbar, 
        mit Eingabedaten, Ergebnis-Matrix, Geräteauswahl.</div></div>""", unsafe_allow_html=True)
        
        if not is_docx_available():
            st.warning("⚠️ `python-docx` nicht installiert. Bitte installieren:\n```\npip install python-docx\n```")
        elif st.button("📝 WORD GENERIEREN", use_container_width=True):
            with st.spinner("Word wird erstellt..."):
                try:
                    word_bytes = generate_word_report(
                        proj_name, kunde_name, bearbeiter, firma,
                        room_results, g_sums, selected_hw, total_kw,
                        selected_hw_ag=selected_hw_ag,
                        room_inputs=room_inputs_list,
                        partner_firma=partner_firma,
                        selected_ig_artnr=selected_ig_artnr
                    )
                    st.download_button("⬇️ WORD HERUNTERLADEN", word_bytes,
                        file_name=f"coolMATH_{proj_name}_{datetime.now().strftime('%Y%m%d')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="word_dl")
                    st.success("✅ Word-Dokument bereit!")
                except Exception as e:
                    st.error(f"Fehler: {e}")

    with wrd2:
        st.markdown("""<div class="card-blue"><div style="font-size:11px;font-weight:700;opacity:0.8;margin-bottom:8px">
        💾 PROJEKT SPEICHERN</div><div style="font-size:12px">Projekt in Datenbank speichern — 
        wiederabrufbar über Projektarchiv. Nur für Ihre Firma sichtbar.</div></div>""", unsafe_allow_html=True)
        if st.button("💾 PROJEKT SPEICHERN", use_container_width=True):
            with st.spinner("Speichere..."):
                pid = db_save_project(
                    partner_firma, auth_username, proj_name, kunde_name, bearbeiter,
                    room_inputs_list, room_results, g_sums, selected_hw, selected_hw_ag
                )
                if pid:
                    st.success(f"✅ Gespeichert! Projekt-ID: `{pid}`")
                    st.caption("Abrufbar über das Projektarchiv.")
                else:
                    st.error("Speicherung fehlgeschlagen.")

    with wrd3:
        st.markdown("""<div class="card-blue"><div style="font-size:11px;font-weight:700;opacity:0.8;margin-bottom:8px">
        📤 MONDAY.COM</div><div style="font-size:12px">Projekt + PDF automatisch zu Monday.com 
        übertragen (Board-Konfiguration in st.secrets).</div></div>""", unsafe_allow_html=True)
        if st.button("📤 MONDAY UPLOAD", use_container_width=True):
            if "monday_obj" not in st.session_state:
                st.session_state.monday_obj = MondayIntegration()
            mon = st.session_state.monday_obj
            if not mon.is_configured():
                st.warning("⚠️ Monday.com nicht konfiguriert. Bitte API-Token + Board-ID in st.secrets['monday'] eintragen.")
            else:
                with st.spinner("📤 Upload läuft..."):
                    try:
                        # Sicherstellen dass samsung_recs nicht None ist
                        _samsung_recs = samsung_recs if samsung_recs else []
                        # PDF für Anhang
                        _pdf_mon = generate_uebergabe_pdf(
                            proj_name, kunde_name, bearbeiter, firma,
                            room_results, g_sums, individual_profiles,
                            _samsung_recs, selected_hw, total_kw,
                            selected_hw_ag=selected_hw_ag,
                            room_inputs=room_inputs_list, partner_firma=partner_firma,
                            selected_ig_artnr=selected_ig_artnr
                        )
                        geraete_str = " | ".join([
                            f"Z{zi+1}: {selected_hw[zi]:.1f}kW" for zi in range(5) if selected_hw[zi]>0
                        ])
                        ok, item_id = mon.save_to_monday({
                            "projekt":    proj_name,
                            "kunde":      kunde_name,
                            "bearbeiter": bearbeiter,
                            "peak_kw":    round(float(np.max(g_sums["VDI_N"]))/1000, 2),
                            "geraete":    geraete_str,
                        }, _pdf_mon, f"coolMATH_{proj_name}.pdf")
                        if ok:
                            st.success(f"✅ Monday.com Item: {item_id}")
                        else:
                            st.error(f"Fehler: {item_id}")
                    except Exception as e:
                        st.error(f"Monday-Fehler: {e}")
    
    with wrd4:
        st.markdown("""<div class="card-blue"><div style="font-size:11px;font-weight:700;opacity:0.8;margin-bottom:8px">
        📊 EXCEL-ANFRAGE</div><div style="font-size:12px">Geräte-Anfrage als Excel für °coolsulting — 
        Innen- und Außengeräte mit Artikelnummern.</div></div>""", unsafe_allow_html=True)
        if st.button("📊 EXCEL GENERIEREN", use_container_width=True):
            with st.spinner("Excel wird erstellt..."):
                try:
                    excel_bytes = generate_excel_anfrage(
                        proj_name, kunde_name, bearbeiter, firma,
                        selected_hw, selected_hw_ag, zone_names,
                        selected_ig_artnr=selected_ig_artnr,
                        liefertermin=liefertermin_str
                    )
                    st.download_button(
                        "⬇️ EXCEL HERUNTERLADEN",
                        data=excel_bytes,
                        file_name=f"coolMATH_Anfrage_{proj_name}_{datetime.now().strftime('%Y%m%d')}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        key="excel_dl"
                    )
                    st.success("✅ Excel-Anfrage bereit!")
                except Exception as e:
                    st.error(f"Fehler: {e}")

    # ==========================================
    # PROJEKTARCHIV
    # ==========================================
    with st.expander("📂 Projektarchiv — gespeicherte Projekte"):
        projekte = db_load_projects(partner_firma, auth_role)
        if projekte:
            if auth_role == "admin":
                st.caption("👑 Admin-Ansicht: alle Projekte aller Firmen")
            else:
                st.caption(f"🔒 Nur Projekte von: {partner_firma}")
            
            # DataFrame mit View-Controls
            proj_df = pd.DataFrame(projekte,
                columns=["ID","Firma","Projekt","Kunde","Bearbeiter","Datum"]).astype(str)
            st.dataframe(proj_df, use_container_width=True, hide_index=True)
            
            st.divider()
            
            # Projekt auswählen + Laden
            col1, col2 = st.columns([4, 1])
            with col1:
                projekt_optionen = {f"{p[2]} | {p[3]} | {p[5]}": p[0] 
                                   for p in projekte}
                selected = st.selectbox(
                    "Projekt zum Laden auswählen:",
                    options=list(projekt_optionen.keys()),
                    label_visibility="collapsed"
                )
            with col2:
                if st.button("📥 LADEN", type="primary", use_container_width=True):
                    proj_id = projekt_optionen[selected]
                    row = db_load_project(proj_id)
                    if row:
                        try:
                            import json as _json
                            room_data_json = row[8] if len(row) > 8 else '{}'
                            results_json   = row[9] if len(row) > 9 else '{}'
                            devices_json   = row[10] if len(row) > 10 else '{}'
                            
                            room_data = _json.loads(room_data_json) if room_data_json else {}
                            results   = _json.loads(results_json) if results_json else {}
                            devices   = _json.loads(devices_json) if devices_json else {}
                            
                            st.session_state['loaded_project'] = {
                                'projekt': row[4],
                                'kunde': row[5],
                                'bearbeiter': row[6],
                                'room_inputs': room_data.get('room_inputs', []),
                                'room_results': results.get('room_results', []),
                                'peaks': results.get('peaks', {}),
                                'selected_hw': devices.get('selected_hw', [0,0,0,0,0]),
                                'selected_hw_ag': devices.get('selected_hw_ag', []),
                            }
                            
                            st.success(f"✅ Projekt '{row[4]}' geladen!")
                            st.rerun()
                        except Exception as e:
                            st.error(f"Fehler: {e}")
                    else:
                        st.error("Laden fehlgeschlagen")
        else:
            st.info("Noch keine gespeicherten Projekte.")

    # --- Footer ---
    st.write("---")
    st.markdown(f"""
    <div style="text-align:center; color:rgba(255,255,255,0.5); font-size:11px; padding:10px">
        © {datetime.now().year} °coolsulting — Michael Schäpers | coolMATH Pro {APP_VERSION}<br>
        Alle Rechte vorbehalten. Haftungsausschluss: Ergebnisse dienen als Planungshilfe, kein Ersatz für Fachplanerprüfung.<br>
        6-Methoden Kühllastsimulation | Samsung Wind-Free Integration | 
        VDI 6007 | VDI 2078 | Recknagel | Kaltluftsee | KI-Hybrid
    </div>
    """, unsafe_allow_html=True)


if __name__ == '__main__':
    main()