# ==========================================
# DATEI:    coolNEIGHBOR_v6.py
# VERSION:  6.0 FLAGSHIP EDITION
# DATUM:    2026-03-29
# AUTOR:    deg coolsulting . Michael Schaepers
#
# NEU IN V6.0:
#  . AT / DE / HR Laendernormen (OeAL 3, TA Laerm, HRN ISO 1996)
#  . Vollstaendige Oktavband-DB (63-8000 Hz) aus Samsung TDB RAC R32 NASA 2025
#  . NR-Kurven-Visualisierung (rechnerisch, ISO 1996-1) via Plotly
#  . Beurteilungspegel Lr = Lp + KT + KI (TA Laerm / OeAL 3 konform)
#  . ISO 9613-2 Beugungsrechnung (Vektorgeometrie L/U-foermig, Fresnel)
#  . Kaskaden-Addition: 10*log10(n) fuer n gleiche Einheiten
#  . Aerodynamik-Pruefung Luftkanal (>3 m/s -> Pegelerhohung)
#  . Logarithmische Mehrquellen-Addition
#  . Abstandsverdopplungs-Schnellrechner
#  . PDF + Word Export mit Haftungsausschluss
#  . deg coolsulting Branding: #36A9E1, POE Vetica UI, TT Commons
# ==========================================

import streamlit as st
import math
import io
import os
from datetime import datetime
from docx import Document
from docx.shared import RGBColor

# ReportLab fuer professionellen PDF-Export
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import mm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                 TableStyle, HRFlowable, PageBreak)
from reportlab.graphics.shapes import Drawing, Rect, Line, String
from reportlab.graphics import renderPDF
from reportlab.pdfgen import canvas as rl_canvas

import plotly.graph_objects as go

# BRANDING
C_BLUE  = "#36A9E1"
C_GREY  = "#3C3C3B"
C_RED   = "#E13636"
C_GREEN = "#36E16A"
BASE_DIR   = os.path.dirname(os.path.abspath(__file__))
LOGO_FILE  = "Coolsulting_Logo_ohneHG_outlines_blau.png"
FONT_HEAD  = "POE Vetica UI.ttf"
FONT_BODY  = "TT Commons DemiBold.otf"

# NORMEN & GRENZWERTE
LAENDER = [
    "AT - Oesterreich (OeAL 3 / OENORM S 5021)",
    "DE - Deutschland (TA Laerm)",
    "HR - Kroatien (HRN ISO 1996-2)",
]

GRENZWERTE = {
    "AT - Oesterreich (OeAL 3 / OENORM S 5021)": {
        "Kat 1 - Erholungs-/Kurgebiet":        {"Tag": 45, "Nacht": 30},
        "Kat 2 - Reines Wohngebiet":            {"Tag": 45, "Nacht": 35},
        "Kat 3 - Allgemeines Wohngebiet":       {"Tag": 50, "Nacht": 40},
        "Kat 4 - Kern-/Mischgebiet":            {"Tag": 55, "Nacht": 45},
        "Kat 5 - Gewerbegebiet":                {"Tag": 60, "Nacht": 50},
        "Kat 6 - Industriegebiet":              {"Tag": 65, "Nacht": 55},
    },
    "DE - Deutschland (TA Laerm)": {
        "Kurgebiete, Krankenhaeuser":           {"Tag": 45, "Nacht": 35},
        "Reine Wohngebiete":                    {"Tag": 50, "Nacht": 35},
        "Allgemeine Wohn- u. Kleinsiedlungsgebiete": {"Tag": 55, "Nacht": 40},
        "Kern-, Dorf- und Mischgebiete":        {"Tag": 60, "Nacht": 45},
        "Gewerbegebiete":                       {"Tag": 65, "Nacht": 50},
        "Industriegebiete":                     {"Tag": 70, "Nacht": 70},
    },
    "HR - Kroatien (HRN ISO 1996-2)": {
        "Tiha zona / Ruhige Zone":              {"Tag": 45, "Nacht": 35},
        "Stambena zona / Wohngebiet":           {"Tag": 55, "Nacht": 45},
        "Mjesovita zona / Mischgebiet":         {"Tag": 60, "Nacht": 50},
        "Poslovna zona / Gewerbegebiet":        {"Tag": 65, "Nacht": 55},
        "Industrijska zona / Industriegebiet":  {"Tag": 70, "Nacht": 65},
    },
}

# SAMSUNG GERAETEDATENBANK
# Quellen:
#   RAC Single Split: Samsung TDB RAC R32 NASA Single Split 2025, Ver 1.0, Dez 2024
#   FJM Multi Split:  Samsung TDB Free Joint Multi R32 Europe 2025, Ver 1.0
# Lw = Schallleistungspegel dB(A) Ausseneinheit, Cooling
# Lp_oktav = [63, 125, 250, 500, 1000, 2000, 4000, 8000 Hz] aus NR-Kurvenanalyse
SAMSUNG_DB = {
    # ── RAC SINGLE SPLIT ─────────────────────────────────────────────────
    "Samsung RAC Elite (AR9500T)": {
        "AR70F09CAAWXEU (2.5 kW)": {
            "Lw": 59, "Lp_1m": 45, "Luft_m3h": 1800, "Gewicht_kg": 30.7,
            "Lp_oktav": [46, 44, 40, 36, 32, 27, 20, 12],
        },
        "AR70F12CAAWXEU (3.5 kW)": {
            "Lw": 62, "Lp_1m": 46, "Luft_m3h": 1800, "Gewicht_kg": 30.7,
            "Lp_oktav": [48, 46, 42, 38, 34, 29, 22, 14],
        },
    },
    "Samsung RAC Avant (AR7500T)": {
        "AR70F07C1AWXEU (2.0 kW)": {
            "Lw": 59, "Lp_1m": 45, "Luft_m3h": 1740, "Gewicht_kg": 30.7,
            "Lp_oktav": [46, 44, 40, 36, 32, 27, 20, 12],
        },
        "AR70F09C1AWXEU (2.5 kW)": {
            "Lw": 59, "Lp_1m": 45, "Luft_m3h": 1740, "Gewicht_kg": 30.7,
            "Lp_oktav": [46, 44, 40, 36, 32, 27, 20, 12],
        },
        "AR70F12C1AWXEU (3.5 kW)": {
            "Lw": 62, "Lp_1m": 46, "Luft_m3h": 1740, "Gewicht_kg": 30.7,
            "Lp_oktav": [48, 46, 42, 38, 34, 29, 22, 14],
        },
        "AR70F15C1AWXEU (4.2 kW)": {
            "Lw": 65, "Lp_1m": 48, "Luft_m3h": 2820, "Gewicht_kg": 36.8,
            "Lp_oktav": [51, 49, 45, 41, 37, 32, 25, 17],
        },
        "AR70F18C1AWXEU (5.0 kW)": {
            "Lw": 65, "Lp_1m": 51, "Luft_m3h": 2820, "Gewicht_kg": 36.8,
            "Lp_oktav": [53, 51, 47, 43, 39, 34, 27, 19],
        },
        "AR70F24C1AWXEU (6.8 kW)": {
            "Lw": 68, "Lp_1m": 54, "Luft_m3h": 3420, "Gewicht_kg": 38.6,
            "Lp_oktav": [56, 54, 50, 46, 42, 37, 30, 22],
        },
    },
    "Samsung RAC Comfort (AR6000T)": {
        "AR60F09C1AWXEU (2.5 kW)": {
            "Lw": 63, "Lp_1m": 45, "Luft_m3h": 2400, "Gewicht_kg": 24.0,
            "Lp_oktav": [46, 44, 40, 37, 33, 28, 21, 13],
        },
        "AR60F12C1AWXEU (3.5 kW)": {
            "Lw": 63, "Lp_1m": 46, "Luft_m3h": 2400, "Gewicht_kg": 24.0,
            "Lp_oktav": [48, 46, 42, 38, 34, 29, 22, 14],
        },
        "AR60F18C1AWXEU (5.0 kW)": {
            "Lw": 65, "Lp_1m": 51, "Luft_m3h": 3000, "Gewicht_kg": 36.8,
            "Lp_oktav": [53, 51, 47, 43, 39, 34, 27, 19],
        },
        "AR60F24C1AWXEU (6.8 kW)": {
            "Lw": 68, "Lp_1m": 54, "Luft_m3h": 3600, "Gewicht_kg": 38.6,
            "Lp_oktav": [56, 54, 50, 46, 42, 37, 30, 22],
        },
    },
    # AR50 = Standard-RAC (NICHT Wind-Free! Wind-Free ist eine Inneneinheitstechnologie)
    "Samsung RAC Standard (AR5000T)": {
        "AR50F07C1AHXEU (2.0 kW)": {
            "Lw": 60, "Lp_1m": 45, "Luft_m3h": 1800, "Gewicht_kg": 24.0,
            "Lp_oktav": [46, 44, 40, 37, 33, 28, 21, 13],
        },
        "AR50F09C1AHXEU (2.5 kW)": {
            "Lw": 63, "Lp_1m": 45, "Luft_m3h": 1800, "Gewicht_kg": 24.0,
            "Lp_oktav": [46, 44, 40, 37, 33, 28, 21, 13],
        },
        "AR50F12C1AHXEU (3.5 kW)": {
            "Lw": 63, "Lp_1m": 46, "Luft_m3h": 1800, "Gewicht_kg": 24.4,
            "Lp_oktav": [48, 46, 42, 38, 34, 29, 22, 14],
        },
        "AR50F15C1AHXEU (4.2 kW)": {
            "Lw": 65, "Lp_1m": 48, "Luft_m3h": 2400, "Gewicht_kg": 30.7,
            "Lp_oktav": [51, 49, 45, 41, 37, 32, 25, 17],
        },
        "AR50F18C1AHXEU (5.0 kW)": {
            "Lw": 65, "Lp_1m": 51, "Luft_m3h": 2400, "Gewicht_kg": 36.8,
            "Lp_oktav": [53, 51, 47, 43, 39, 34, 27, 19],
        },
        "AR50F24C1AHXEU (6.8 kW)": {
            "Lw": 68, "Lp_1m": 54, "Luft_m3h": 3000, "Gewicht_kg": 38.6,
            "Lp_oktav": [56, 54, 50, 46, 42, 37, 30, 22],
        },
    },
    # ── FJM MULTI-SPLIT AUSSENEINHEITEN ──────────────────────────────────
    # Quelle: Samsung TDB Free Joint Multi R32 Europe 2025
    # Lw Ausseneinheit Cooling (aus Spezifikations-Tabelle)
    "Samsung FJM Multi-Split (AJ-Ausseneinheit)": {
        "AJ040TXJ2KG/EU  (4.0 kW / 2 Innen)": {
            "Lw": 61, "Lp_1m": 45, "Luft_m3h": 1860, "Gewicht_kg": 33.0,
            "Lp_oktav": [48, 46, 42, 38, 34, 29, 22, 14],
        },
        "AJ050TXJ2KG/EU  (5.0 kW / 2 Innen)": {
            "Lw": 61, "Lp_1m": 45, "Luft_m3h": 2100, "Gewicht_kg": 37.0,
            "Lp_oktav": [48, 46, 42, 38, 34, 29, 22, 14],
        },
        "AJ052TXJ3KG/EU  (5.2 kW / 3 Innen)": {
            "Lw": 61, "Lp_1m": 45, "Luft_m3h": 2100, "Gewicht_kg": 44.0,
            "Lp_oktav": [48, 46, 42, 38, 34, 29, 22, 14],
        },
        "AJ068TXJ3KG/EU  (6.8 kW / 3 Innen)": {
            "Lw": 64, "Lp_1m": 48, "Luft_m3h": 2520, "Gewicht_kg": 57.5,
            "Lp_oktav": [51, 49, 45, 41, 37, 32, 25, 17],
        },
        "AJ080TXJ4KG/EU  (8.0 kW / 4 Innen)": {
            "Lw": 66, "Lp_1m": 50, "Luft_m3h": 3480, "Gewicht_kg": 70.0,
            "Lp_oktav": [53, 51, 47, 43, 39, 34, 27, 19],
        },
        "AJ100TXJ5KG/EU  (10.0 kW / 5 Innen)": {
            "Lw": 70, "Lp_1m": 54, "Luft_m3h": 4020, "Gewicht_kg": 76.5,
            "Lp_oktav": [57, 55, 51, 47, 43, 38, 31, 23],
        },
    },
    "Fremdgeraet / Freie Eingabe": {
        "Manuelle Eingabe": {
            "Lw": 65.0, "Lp_1m": 50.0, "Luft_m3h": 2800, "Gewicht_kg": 40.0,
            "Lp_oktav": [52, 50, 46, 42, 38, 33, 26, 18],
        },
    },
}

# NR-Kurven Referenzwerte (ISO 1996-1), Frequenzen: [63,125,250,500,1k,2k,4k,8k]
NR_KURVEN = {
    15: [51.8, 35.5, 22.1, 12.5,  5.4,  0.1, -3.5, -6.2],
    20: [55.4, 39.4, 26.5, 17.1, 10.1,  4.9,  1.4, -1.2],
    25: [59.1, 43.4, 30.9, 21.7, 14.9,  9.7,  6.3,  3.8],
    30: [62.7, 47.3, 35.3, 26.3, 19.7, 14.5, 11.2,  8.8],
    35: [66.4, 51.2, 39.7, 30.9, 24.5, 19.3, 16.2, 13.8],
    40: [70.0, 55.1, 44.1, 35.5, 29.3, 24.1, 21.1, 18.8],
    45: [73.7, 59.1, 48.5, 40.1, 34.1, 28.9, 26.0, 23.8],
    50: [77.3, 63.0, 52.9, 44.7, 38.9, 33.7, 31.0, 28.8],
    55: [81.0, 66.9, 57.3, 49.3, 43.7, 38.5, 36.0, 33.8],
    60: [84.6, 70.9, 61.7, 53.9, 48.5, 43.3, 41.0, 38.8],
    65: [88.2, 74.8, 66.1, 58.5, 53.3, 48.1, 46.0, 43.8],
}
OKTAV_BANDS = [63, 125, 250, 500, 1000, 2000, 4000, 8000]

DISCLAIMER_TEXT = (
    "WICHTIGER HINWEIS: Kein Sachverstaendigengutachten!\n\n"
    "Bei diesem Dokument handelt es sich ausschliesslich um eine planungstechnische "
    "Schallimmissions-Prognoseberechnung im Rahmen der Anlagenplanung.\n\n"
    "1. Charakter und Zweck\n"
    "Die vorliegende Berechnung ist eine planungstechnische Prognose auf Basis der "
    "Berechnungsverfahren nach ISO 9613-1/2, TA Laerm, OeAL 3, HRN ISO 1996 sowie "
    "den Herstellerangaben (Samsung TDB RAC R32 NASA, Ver. 1.0, Dezember 2024).\n\n"
    "2. Haftungsausschluss\n"
    "Der Ersteller uebernimmt keine Haftung fuer: Abweichungen zwischen Prognose und "
    "Messung; Veraenderungen nach Berichtserstellung; Entscheidungen behoerdlicher "
    "Stellen; Ansprueche Dritter.\n\n"
    "3. Empfehlung\n"
    "Fuer eine behoerdlich anerkannte Schallmessung ist ein akkreditiertes Akustikinstitut "
    "oder gerichtlich beeideter Sachverstaendiger zu beauftragen. "
    "Messgeraete: Klasse 1 nach DIN EN 61672 (+/-0.7 dB, 20-12500 Hz)."
)

# ──────────────────────────────────────────────────
# PHYSIK-ENGINE
# ──────────────────────────────────────────────────

def log_add(levels):
    if not levels:
        return 0.0
    return 10 * math.log10(sum(10 ** (l / 10) for l in levels))

def kaskaden_zuschlag(n):
    if n <= 1:
        return 0.0
    return round(10 * math.log10(n), 1)

def atm_daempfung(d_m, temp_c, rel_hum):
    alpha = max(1.9 * (1 + (10 - temp_c) * 0.02) * (1 + (70 - rel_hum) * 0.01), 0.5)
    return round((alpha * d_m) / 1000.0, 3)

def beugungsdaempfung(umweg_delta, barriere_typ, f_hz=500.0):
    if umweg_delta <= 0:
        return 0.0
    n = 2.0 * umweg_delta / (340.0 / f_hz)
    if n <= 0:
        return 0.0
    dz = 10 * math.log10(3 + 20 * n)
    limits = {"Massiv & Luftdicht": 25.0, "Leichtbauweise dicht": 15.0,
              "Teiloffen / Lamellen": 5.0, "Luftdurchlässig (Gitter)": 0.0}
    lim = next((v for k, v in limits.items() if k in barriere_typ), 25.0)
    return round(min(dz, lim), 1)

def aero_zuschlag(luft_m3h, qs_m2):
    if qs_m2 <= 0:
        return 0.0, 0.0
    v = luft_m3h / (3600.0 * qs_m2)
    zs = 0.0 if v <= 3.0 else round(((v - 3.0) / 0.03) * 0.25, 1)
    return round(v, 2), zs

def berechne_gesamt(lw, anzahl, luft, qs, d_geh, d_koerper, k_t, k_i,
                    d_total, topologie, b1_mat, b2_mat, b1_lenkt, b2_lenkt,
                    q, w_dist, temp, hum, raumzuschlag, d_direkt):
    """
    ISO 9613-2 Schallausbreitungsrechnung.

    d_total  = entrollter Schallweg entlang der Schallbahn [m]
               (= Luftlinie bei direktem Weg, laenger bei Beugung)
    d_direkt = kuerzte Luftlinie Quelle -> Immissionsort [m]
               (immer <= d_total; bei direktem Weg = d_total)
    umweg    = d_total - d_direkt = Mehrweg durch Beugung an Kanten [m]
               WICHTIG: umweg steigt mit groesserer Beugungsbahn,
               groesserer umweg = staerkere Fresnel-Beugungsdaempfung.
               d_direkt kleiner bei gleichem d_total = groesserer Umweg = mehr Daempfung.
               Das ist physikalisch KORREKT (tieferer Schallschatten hinter der Kante).
    """
    d_kas = kaskaden_zuschlag(anzahl)
    v_luft, d_aero = aero_zuschlag(luft, qs)
    lw_eff = lw + d_kas + d_aero - d_geh - d_koerper
    q_db = {1: 0, 2: 3, 4: 6, 8: 9}.get(q, 3)
    # Wandreflexion: nur bei sehr kleinem Wandabstand
    w_plus = 3.0 if w_dist < 0.5 else (1.5 if w_dist < 1.0 else 0.0)
    a_atm = atm_daempfung(d_total, temp, hum)
    # Geometrische Ausbreitung auf d_total (entrollter Weg = massgebend fuer Pegel)
    a_div = 10 * math.log10(4 * math.pi * max(d_total, 0.001) ** 2)
    lp_frei = lw_eff + q_db - a_div + w_plus - a_atm + raumzuschlag
    # Beugungsberechnung: umweg = entrollter Weg - direkte Luftlinie
    # d_direkt muss <= d_total sein (physikalische Bedingung)
    d_direkt_safe = min(d_direkt, d_total)
    umweg = max(d_total - d_direkt_safe, 0.0)
    dz1 = 0.0
    dz2 = 0.0
    d_uml = 0.0
    if umweg > 0 and "Direkt" not in topologie:
        # Bei L-foermig: gesamter Umweg an einer Kante
        # Bei U-foermig: aufgeteilt auf zwei Kanten
        if "U-foermig" in topologie or "U-förmig" in topologie:
            dz1 = beugungsdaempfung(umweg * 0.6, b1_mat)
            dz2 = beugungsdaempfung(umweg * 0.4, b2_mat)
        else:
            dz1 = beugungsdaempfung(umweg, b1_mat)
    if b1_lenkt and "Direkt" not in topologie:
        d_uml += 2.0
    if b2_lenkt and ("U-foermig" in topologie or "U-förmig" in topologie):
        d_uml += 1.5
    lp = round(lp_frei - dz1 - dz2 - d_uml, 1)
    lr = round(lp + k_t + k_i, 1)
    return {
        "lw_eff": round(lw_eff, 1), "d_kas": d_kas, "d_aero": d_aero,
        "v_luft": v_luft, "q_db": q_db, "w_plus": w_plus,
        "a_atm": a_atm, "a_div": round(a_div, 1), "lp_frei": round(lp_frei, 1),
        "dz1": dz1, "dz2": dz2, "d_uml": d_uml, "umweg": round(umweg, 2),
        "lp": lp, "lr": lr,
    }

def oktav_am_io(lp_oktav_src, d_total, q, d_geh, anzahl):
    q_db = {1: 0, 2: 3, 4: 6, 8: 9}.get(q, 3)
    d_div = 10 * math.log10(4 * math.pi * max(d_total, 0.001) ** 2)
    d_kas = kaskaden_zuschlag(anzahl)
    return [round(lp + q_db - d_div + d_kas - d_geh, 1) for lp in lp_oktav_src]

def nr_klasse(lp_io):
    for nr in sorted(NR_KURVEN.keys()):
        if all(lp_io[i] <= NR_KURVEN[nr][i] for i in range(8)):
            return nr
    return None

# ──────────────────────────────────────────────────
# DIAGRAMM
# ──────────────────────────────────────────────────

def nr_plot(lp_io, titel):
    fig = go.Figure()
    for nr in [20, 25, 30, 35, 40, 45, 50, 55]:
        vals = NR_KURVEN[nr]
        fig.add_trace(go.Scatter(
            x=OKTAV_BANDS, y=vals, mode="lines",
            line=dict(color="#444466", width=1, dash="dot"),
            name=f"NR {nr}", showlegend=False,
            hovertemplate=f"NR {nr}: %{{y:.1f}} dB<extra></extra>"))
        fig.add_annotation(x=8000, y=vals[-1], text=f"NR {nr}", showarrow=False,
                           font=dict(size=9, color="#8888AA"), xanchor="left")
    fig.add_trace(go.Scatter(
        x=OKTAV_BANDS, y=lp_io, mode="lines+markers",
        name="Prognose (Immissionsort)",
        line=dict(color=C_BLUE, width=3),
        marker=dict(size=9, color=C_BLUE, line=dict(color="white", width=1.5)),
        hovertemplate="%{x} Hz: <b>%{y:.1f} dB</b><extra></extra>"))
    fig.update_layout(
        title=dict(text=titel, font=dict(color=C_BLUE, size=13)),
        xaxis=dict(type="log", tickvals=OKTAV_BANDS,
                   ticktext=["63", "125", "250", "500", "1k", "2k", "4k", "8k"],
                   title="Oktavband-Mittenfrequenz [Hz]", gridcolor="#2A2A3A"),
        yaxis=dict(title="Schalldruckpegel [dB]", gridcolor="#2A2A3A", range=[0, 75]),
        plot_bgcolor="#111122", paper_bgcolor="#111122",
        font=dict(color="#DDDDDD"), height=340,
        margin=dict(l=60, r=90, t=50, b=50))
    return fig

# ──────────────────────────────────────────────────
# PDF EXPORT
# ──────────────────────────────────────────────────

def sanitize(txt):
    txt = str(txt)
    pairs = [
        ("\u2013","-"),("\u2014","-"),("\u2015","-"),("\u2012","-"),("\u2212","-"),
        ("\u00b0","Grad"),("\u00b3","3"),("\u00b2","2"),
        ("\u00e4","ae"),("\u00f6","oe"),("\u00fc","ue"),
        ("\u00c4","Ae"),("\u00d6","Oe"),("\u00dc","Ue"),
        ("\u00df","ss"),("\u202f"," "),("\u00a0"," "),
        ("\u00b7","."),("\u2019","\x27"),
    ]
    for o,n in pairs:
        txt = txt.replace(o,n)
    txt = txt.encode("latin-1",errors="replace").decode("latin-1")
    return txt

def create_pdf(daten, projekt, ersteller, erg=None, grenzwert=None,
               lp_io=None, lp_oktav_src=None, konform=None, lr=None, lp_val=None,
               modell="", land="", widmung="", auftraggeber=""):
    """Professioneller PDF-Bericht mit ReportLab - deg coolsulting Branding"""

    # --- Farben ---
    BLUE    = colors.HexColor("#36A9E1")
    GREY    = colors.HexColor("#3C3C3B")
    LGREY   = colors.HexColor("#F4F8FB")
    GREEN   = colors.HexColor("#1DB954")
    RED     = colors.HexColor("#E13636")
    WHITE   = colors.white
    BLACK   = colors.HexColor("#1A1A1A")
    MIDGREY = colors.HexColor("#888888")

    buf = io.BytesIO()
    W, H = A4  # 595 x 842 pt

    # --- Canvas-basierter Aufbau (volle Kontrolle) ---
    c = rl_canvas.Canvas(buf, pagesize=A4)

    def draw_header(c, page_num):
        # Blauer Header-Balken
        c.setFillColor(BLUE)
        c.rect(0, H - 28*mm, W, 28*mm, fill=1, stroke=0)
        # Weisser Titeltext
        c.setFillColor(WHITE)
        c.setFont("Helvetica-Bold", 14)
        c.drawString(12*mm, H - 12*mm, "coolNEIGHBOR  |  Schallimmissions-Prognose")
        c.setFont("Helvetica", 9)
        c.drawString(12*mm, H - 20*mm, "coolsulting e.U.  |  Michael Schaepers")
        # Seitenzahl rechts
        c.setFont("Helvetica", 8)
        c.drawRightString(W - 12*mm, H - 20*mm, f"Seite {page_num}")

    def draw_footer(c):
        c.setFillColor(MIDGREY)
        c.rect(0, 0, W, 10*mm, fill=1, stroke=0)
        c.setFillColor(WHITE)
        c.setFont("Helvetica-Oblique", 7)
        c.drawCentredString(W/2, 3.5*mm,
            "Planungstechnische Prognose - kein Sachverstaendigengutachten | ISO 9613-1/2 | OeAL 3 | TA Laerm | HRN ISO 1996")

    def section_title(c, y, text):
        c.setFillColor(BLUE)
        c.rect(12*mm, y - 1*mm, W - 24*mm, 7*mm, fill=1, stroke=0)
        c.setFillColor(WHITE)
        c.setFont("Helvetica-Bold", 10)
        c.drawString(15*mm, y + 1.5*mm, text)
        return y - 10*mm

    def data_row(c, y, label, value, highlight=False, good=False, bad=False):
        if highlight:
            bg = LGREY
            c.setFillColor(bg)
            c.rect(12*mm, y - 1*mm, W - 24*mm, 6.5*mm, fill=1, stroke=0)
        c.setFillColor(GREY)
        c.setFont("Helvetica", 9)
        c.drawString(15*mm, y + 1*mm, str(label))
        if good:
            c.setFillColor(GREEN)
        elif bad:
            c.setFillColor(RED)
        else:
            c.setFillColor(BLACK)
        c.setFont("Helvetica-Bold" if (good or bad) else "Helvetica", 9)
        c.drawString(95*mm, y + 1*mm, str(value))
        return y - 6*mm

    def draw_status_box(c, y, lr_val, gw, ok):
        box_h = 14*mm
        col = GREEN if ok else RED
        c.setFillColor(col)
        c.setStrokeColor(col)
        c.roundRect(12*mm, y - box_h, W - 24*mm, box_h, 3*mm, fill=1, stroke=0)
        c.setFillColor(WHITE)
        c.setFont("Helvetica-Bold", 14)
        status = "KONFORM" if ok else "UEBERSCHRITTEN"
        delta = round(gw - lr_val, 1)
        txt = f"{status}  |  Lr = {lr_val} dB(A)  |  Grenzwert {gw} dB(A)"
        if ok:
            txt += f"  |  Reserve: +{delta} dB(A)"
        else:
            txt += f"  |  Ueberschreitung: {abs(delta)} dB(A)"
        c.drawCentredString(W/2, y - box_h/2 - 2*mm, txt)
        return y - box_h - 5*mm

    def draw_bar_chart(c, y_top, lp_io_vals, grenzwert_val, x_left=12*mm, chart_w=None, chart_h=55*mm):
        if chart_w is None:
            chart_w = W - 24*mm
        bands = [63, 125, 250, 500, 1000, 2000, 4000, 8000]
        n = len(bands)
        bar_w = (chart_w - 10*mm) / n
        max_val = max(max(lp_io_vals) + 5, grenzwert_val + 10, 60)

        # Hintergrund
        c.setFillColor(LGREY)
        c.rect(x_left, y_top - chart_h, chart_w, chart_h, fill=1, stroke=0)

        # Grenzwert-Linie
        gw_y = y_top - chart_h + (grenzwert_val / max_val) * chart_h
        c.setStrokeColor(RED)
        c.setDash(4, 3)
        c.line(x_left, gw_y, x_left + chart_w, gw_y)
        c.setDash()
        c.setFillColor(RED)
        c.setFont("Helvetica-Bold", 7)
        c.drawString(x_left + chart_w - 22*mm, gw_y + 1*mm, f"Grenzwert {grenzwert_val} dB")

        # Balken
        for i, (hz, val) in enumerate(zip(bands, lp_io_vals)):
            bar_x = x_left + 5*mm + i * bar_w
            bar_h_px = max((val / max_val) * chart_h, 1)
            col = GREEN if val <= grenzwert_val else RED
            c.setFillColor(col)
            c.setStrokeColor(WHITE)
            c.rect(bar_x + 1*mm, y_top - chart_h, bar_w - 2*mm, bar_h_px, fill=1, stroke=1)
            # Wert oben
            c.setFillColor(BLACK)
            c.setFont("Helvetica-Bold", 7)
            c.drawCentredString(bar_x + bar_w/2, y_top - chart_h + bar_h_px + 1*mm, str(val))
            # Hz unten
            c.setFont("Helvetica", 6.5)
            hz_label = f"{hz}Hz" if hz < 1000 else f"{hz//1000}kHz"
            c.drawCentredString(bar_x + bar_w/2, y_top - chart_h - 4*mm, hz_label)

        # Y-Achse Beschriftung
        c.setFillColor(MIDGREY)
        c.setFont("Helvetica", 7)
        for tick in [0, 20, 40, 60]:
            tick_y = y_top - chart_h + (tick / max_val) * chart_h
            c.drawRightString(x_left - 1*mm, tick_y - 1*mm, str(tick))
            c.setStrokeColor(colors.HexColor("#CCCCCC"))
            c.line(x_left, tick_y, x_left + chart_w, tick_y)

        # Achsentitel
        c.setFillColor(GREY)
        c.setFont("Helvetica-Oblique", 7)
        c.drawCentredString(x_left + chart_w/2, y_top - chart_h - 8*mm, "Oktavband-Mittenfrequenz [Hz]")

        return y_top - chart_h - 12*mm

    def draw_summary_boxes(c, y, metrics):
        """4 Kennwert-Boxen nebeneinander"""
        n = len(metrics)
        box_w = (W - 24*mm) / n
        box_h = 18*mm
        for i, (label, value, unit, good) in enumerate(metrics):
            bx = 12*mm + i * box_w
            bg = colors.HexColor("#E8F7FD") if good else colors.HexColor("#FDE8E8")
            border = GREEN if good else RED
            c.setFillColor(bg)
            c.setStrokeColor(border)
            c.roundRect(bx + 1*mm, y - box_h, box_w - 2*mm, box_h, 2*mm, fill=1, stroke=1)
            c.setFillColor(MIDGREY)
            c.setFont("Helvetica", 7)
            c.drawCentredString(bx + box_w/2, y - 4*mm, label)
            c.setFillColor(BLUE)
            c.setFont("Helvetica-Bold", 13)
            c.drawCentredString(bx + box_w/2, y - 11*mm, f"{value}")
            c.setFillColor(GREY)
            c.setFont("Helvetica", 7)
            c.drawCentredString(bx + box_w/2, y - 15.5*mm, unit)
        return y - box_h - 6*mm

    # ══════════════════════════════════════
    # SEITE 1 – DECKBLATT & ERGEBNIS
    # ══════════════════════════════════════
    draw_header(c, 1)
    draw_footer(c)

    y = H - 35*mm

    # Projekttitel
    c.setFillColor(GREY)
    c.setFont("Helvetica-Bold", 18)
    c.drawString(12*mm, y, sanitize(projekt))
    y -= 7*mm

    c.setFillColor(MIDGREY)
    c.setFont("Helvetica", 9)
    c.drawString(12*mm, y,
        f"Erstellt: {datetime.now().strftime('%d.%m.%Y %H:%M')}  |  "
        f"Ersteller: {sanitize(ersteller)}  |  Auftraggeber: {sanitize(auftraggeber) or '-'}")
    y -= 5*mm

    c.setFillColor(BLUE)
    c.rect(12*mm, y, W - 24*mm, 0.4*mm, fill=1, stroke=0)
    y -= 8*mm

    # STATUS BOX
    if lr is not None and grenzwert is not None and konform is not None:
        y = draw_status_box(c, y, lr, grenzwert, konform)
        y -= 3*mm

    # KENNWERTE-BOXEN
    if erg is not None and lr is not None and lp_val is not None and grenzwert is not None:
        delta = round(grenzwert - lr, 1)
        metrics = [
            ("Schalldruckpegel Lp", f"{lp_val}", "dB(A)", lp_val <= grenzwert),
            ("Beurteilungspegel Lr", f"{lr}", "dB(A)", konform),
            ("Grenzwert", f"{grenzwert}", "dB(A)", True),
            ("Reserve / Ueberschr.", f"{delta:+.1f}", "dB(A)", konform),
        ]
        y = draw_summary_boxes(c, y, metrics)
        y -= 3*mm

    # 1. OBJEKTBESCHREIBUNG
    y = section_title(c, y, "1. OBJEKTBESCHREIBUNG & AUSGANGSLAGE")
    desc_lines = [
        f"Standort / Projekt:   {sanitize(projekt)}",
        f"Auftraggeber:         {sanitize(auftraggeber) if auftraggeber else '-'}",
        f"Norm / Grundlage:     OENORM ISO 9613-2  |  {sanitize(land)}",
        f"Gebietswidmung:       {sanitize(widmung)}",
        f"Beurteilungszeitraum: {sanitize(daten.get('Beurteilungszeitraum', 'Nacht'))}",
        f"Immissionsrichtwert:  {grenzwert} dB(A)  (massgebender Nacht-Grenzwert)",
    ]
    for i, line in enumerate(desc_lines):
        y = data_row(c, y, line.split(':')[0].strip() + ":", ':'.join(line.split(':')[1:]).strip(), highlight=(i%2==0))
    y -= 4*mm

    # Intro-Text wie Musterbericht
    intro = (f"Das Aussengeraet ({sanitize(modell)}) wird am geplanten Aufstellort installiert. "
             f"Die vorliegende Berechnung ermittelt den zu erwartenden Schalldruckpegel am "
             f"Immissionsort (Nachbarfenster) unter Beruecksichtigung der baulichen Gegebenheiten "
             f"sowie der technischen Schutzmassnahmen nach OENORM ISO 9613-2.")
    c.setFillColor(GREY)
    c.setFont("Helvetica-Oblique", 8.5)
    words = intro.split()
    cur_line = ""
    for word in words:
        test = cur_line + " " + word if cur_line else word
        if c.stringWidth(test, "Helvetica-Oblique", 8.5) < (W - 30*mm):
            cur_line = test
        else:
            if cur_line:
                c.drawString(15*mm, y, cur_line)
                y -= 4.5*mm
            cur_line = word
    if cur_line:
        c.drawString(15*mm, y, cur_line)
        y -= 7*mm

    # 2. TECHNISCHE KENNDATEN
    y = section_title(c, y, "2. TECHNISCHE KENNDATEN DER SCHALLQUELLE")
    if erg:
        rows_g = [
            ("Geraetetyp / Modell",    sanitize(modell),                                    False),
            ("Schallleistungspegel Lw",f"{daten.get('Lw','-')} dB(A)",                      True),
            ("Kaskaden-Zuschlag",      f"+{erg.get('d_kas',0)} dB(A)  ({daten.get('Anzahl',1)} Einheiten)", False),
            ("Schallschutzkapselung",  f"{daten.get('Kapselung','-')}",                      True),
            ("Koerperschall-Sockel",   f"{daten.get('Koerperschall','-')}",                  False),
            ("Effektiver Lw gesamt",   f"{erg.get('lw_eff','-')} dB(A)",                     True),
        ]
        for i,(lbl,val,_) in enumerate(rows_g):
            is_eff = "Effektiver" in lbl
            y = data_row(c, y, lbl, val, highlight=(i%2==0), good=is_eff)
    y -= 3*mm

    # 3. AUFSTELLGEOMETRIE
    y = section_title(c, y, "3. AUFSTELLGEOMETRIE & RICHTFAKTOR")
    if erg:
        rows_a = [
            ("Richtfaktor Q (Aufstellbedingung)", f"{daten.get('Richtfaktor Q','-')}", False),
            ("Einbausituation / Diffusfeld",       f"{daten.get('Einbausituation','-')}", True),
            ("Wandabstand",                        f"{daten.get('Wandabstand', '-')}", False),
            ("Schallweg gesamt (entrollt)",        f"{daten.get('Gesamtschallweg','-')}", True),
            ("Direkte Luftlinie Quelle-Fenster",   f"{daten.get('Direkte Luftlinie','-')}", False),
            ("Schallumweg (Beugungsdelta)",         f"+{erg.get('umweg',0)} m", True),
            ("Topologie",                          f"{daten.get('Topologie','-')}", False),
        ]
        for i,(lbl,val,hi) in enumerate(rows_a):
            y = data_row(c, y, lbl, val, highlight=(i%2==0))

    c.showPage()

    # ══════════════════════════════════════
    # SEITE 2 – OKTAVSPEKTRUM & PROTOKOLL
    # ══════════════════════════════════════
    draw_header(c, 2)
    draw_footer(c)
    y = H - 35*mm

    # Diagramm-Titel
    y = section_title(c, y, "OKTAVSPEKTRUM AM IMMISSIONSORT (ISO 1996-1 / NR-Kurven)")
    y -= 2*mm

    if lp_io and grenzwert:
        y = draw_bar_chart(c, y, lp_io, grenzwert, chart_h=60*mm)
        y -= 5*mm

        # Legende
        c.setFillColor(GREEN); c.rect(12*mm, y + 2*mm, 8*mm, 4*mm, fill=1, stroke=0)
        c.setFillColor(GREY); c.setFont("Helvetica", 8)
        c.drawString(22*mm, y + 2.5*mm, "Oktavband unter Grenzwert")
        c.setFillColor(RED); c.rect(75*mm, y + 2*mm, 8*mm, 4*mm, fill=1, stroke=0)
        c.drawString(85*mm, y + 2.5*mm, "Oktavband ueber Grenzwert")
        y -= 8*mm

        # Tabelle Oktavwerte
        y = section_title(c, y, "OKTAVBAND-PEGEL AM IMMISSIONSORT [dB]")
        bands = [63, 125, 250, 500, 1000, 2000, 4000, 8000]
        col_w = (W - 24*mm) / 8
        # Header
        c.setFillColor(BLUE)
        for i, hz in enumerate(bands):
            bx = 12*mm + i * col_w
            c.rect(bx, y - 6*mm, col_w, 6*mm, fill=1, stroke=0)
            c.setFillColor(WHITE)
            c.setFont("Helvetica-Bold", 8)
            lbl = f"{hz}Hz" if hz < 1000 else f"{hz//1000}kHz"
            c.drawCentredString(bx + col_w/2, y - 4*mm, lbl)
            c.setFillColor(BLUE)
        y -= 6*mm
        # Werte
        for i, (hz, val) in enumerate(zip(bands, lp_io)):
            bx = 12*mm + i * col_w
            ok = val <= grenzwert
            c.setFillColor(colors.HexColor("#E8F7FD") if ok else colors.HexColor("#FDE8E8"))
            c.rect(bx, y - 6.5*mm, col_w, 6.5*mm, fill=1, stroke=1)
            c.setFillColor(GREEN if ok else RED)
            c.setFont("Helvetica-Bold", 9)
            c.drawCentredString(bx + col_w/2, y - 4.8*mm, str(val))
        y -= 10*mm

    # 4. BERECHNUNGSSCHRITTE (Musterbericht-Stil)
    y = section_title(c, y, "4. BERECHNUNG DER SCHALLAUSBREITUNG (ISO 9613-2)")
    if erg:
        # 4.1 Ausgangspegel
        c.setFillColor(BLUE); c.setFont("Helvetica-Bold", 8.5)
        c.drawString(15*mm, y, "4.1  Ausgangspegel (Schallleistung)")
        y -= 5.5*mm
        steps_1 = [
            ("Schallleistungspegel Lw (Herstellerangabe):",   f"{daten.get('Lw','-')} dB(A)"),
            ("Reduktion durch Kapselung:",                     f"{daten.get('Kapselung','-')}"),
            ("Koerperschall-Entkopplung (Sockel):",            f"{daten.get('Koerperschall','-')}"),
            ("Kaskaden-Zuschlag (n Einheiten):",               f"+{erg.get('d_kas',0)} dB(A)"),
            ("Aerodynamischer Zuschlag (Kanal):",              f"+{erg.get('d_aero',0)} dB(A)" if erg.get('d_aero',0) > 0 else "- (freiblasend)"),
            ("  => Effektiver Ausgangspegel Lw:",              f"{erg.get('lw_eff','-')} dB(A)"),
        ]
        for i,(lbl,val) in enumerate(steps_1):
            is_sub = lbl.strip().startswith("=>")
            bg = colors.HexColor("#D0F0FF") if is_sub else (LGREY if i%2==0 else WHITE)
            c.setFillColor(bg)
            c.rect(12*mm, y - 5*mm, W - 24*mm, 5*mm, fill=1, stroke=0)
            c.setFillColor(GREY); c.setFont("Helvetica-Bold" if is_sub else "Helvetica", 8.5)
            c.drawString(18*mm, y - 3.5*mm, lbl)
            c.setFillColor(BLUE if is_sub else BLACK)
            c.setFont("Helvetica-Bold" if is_sub else "Helvetica", 8.5)
            c.drawString(130*mm, y - 3.5*mm, val)
            y -= 5*mm
        y -= 2*mm

        # 4.2 Abstandsdaempfung
        c.setFillColor(BLUE); c.setFont("Helvetica-Bold", 8.5)
        c.drawString(15*mm, y, "4.2  Abstandsdaempfung (geometrische Ausbreitung)")
        y -= 5.5*mm
        q_num = daten.get('Richtfaktor Q','Q=2').split()[0].replace('Q=','')
        steps_2 = [
            ("Richtfaktor Q (Aufstellbedingung):",     f"{daten.get('Richtfaktor Q','-')}"),
            ("Gesamtschallweg r:",                      f"{daten.get('Gesamtschallweg','-')}"),
            ("Geometr. Daempfung Adiv = 10*log(4*pi*r2):", f"-{erg.get('a_div',0)} dB(A)"),
            ("Wandreflexion:",                          f"+{erg.get('w_plus',0)} dB(A)"),
            ("Einbausituation / Raumzuschlag:",         f"+{erg.get('raumzuschlag',0)} dB" if erg.get('raumzuschlag',0) > 0 else "- (Freifeld)"),
            ("Atmosphaerische Daempfung (ISO 9613-1):", f"-{erg.get('a_atm',0)} dB(A)"),
        ]
        for i,(lbl,val) in enumerate(steps_2):
            c.setFillColor(LGREY if i%2==0 else WHITE)
            c.rect(12*mm, y - 5*mm, W - 24*mm, 5*mm, fill=1, stroke=0)
            c.setFillColor(GREY); c.setFont("Helvetica", 8.5)
            c.drawString(18*mm, y - 3.5*mm, lbl)
            c.setFillColor(BLACK); c.setFont("Helvetica", 8.5)
            c.drawString(130*mm, y - 3.5*mm, val)
            y -= 5*mm
        y -= 2*mm

        # 4.3 Beugungsverlust
        if erg.get('dz1',0) > 0 or erg.get('dz2',0) > 0:
            c.setFillColor(BLUE); c.setFont("Helvetica-Bold", 8.5)
            c.drawString(15*mm, y, "4.3  Beugungsdaempfung an Kanten (ISO 9613-2 / Fresnel)")
            y -= 5.5*mm
            steps_3 = [
                ("Schallumweg delta (Beugungspfad):", f"+{erg.get('umweg',0)} m"),
                ("Beugungsdaempfung Kante 1 (Fresnel):", f"-{erg.get('dz1',0)} dB(A)"),
                ("Beugungsdaempfung Kante 2:", f"-{erg.get('dz2',0)} dB(A)" if erg.get('dz2',0) > 0 else "-"),
                ("Aerodyn. Ablenkung:", f"-{erg.get('d_uml',0)} dB(A)" if erg.get('d_uml',0) > 0 else "-"),
                ("  => Beugungsverlust gesamt:", f"-{round(erg.get('dz1',0)+erg.get('dz2',0)+erg.get('d_uml',0),1)} dB(A)"),
            ]
            for i,(lbl,val) in enumerate(steps_3):
                is_sub = lbl.strip().startswith("=>")
                bg = colors.HexColor("#D0F0FF") if is_sub else (LGREY if i%2==0 else WHITE)
                c.setFillColor(bg)
                c.rect(12*mm, y - 5*mm, W - 24*mm, 5*mm, fill=1, stroke=0)
                c.setFillColor(GREY); c.setFont("Helvetica-Bold" if is_sub else "Helvetica", 8.5)
                c.drawString(18*mm, y - 3.5*mm, lbl)
                c.setFillColor(BLUE if is_sub else BLACK)
                c.setFont("Helvetica-Bold" if is_sub else "Helvetica", 8.5)
                c.drawString(130*mm, y - 3.5*mm, val)
                y -= 5*mm
            y -= 2*mm

        # 4.4 ERGEBNISTABELLE (wie Musterbericht Tab. "Parameter / Wert")
        y = section_title(c, y, "4. ERGEBNIS DER PROGNOSE")
        result_rows = [
            ("Effektiver Ausgangspegel (inkl. aller Korrekturen)", f"{erg.get('lw_eff','-')} dB(A)", False),
            ("Abzug: Distanz & Aufstellungsgeometrie",            f"-{erg.get('a_div',0)} dB(A)", False),
            ("Abzug: Beugungsverlust an Kante(n)",                f"-{round(erg.get('dz1',0)+erg.get('dz2',0)+erg.get('d_uml',0),1)} dB(A)", False),
            ("Abzug: Atmosphaerische Daempfung",                  f"-{erg.get('a_atm',0)} dB(A)", False),
            ("Tonhaltigkeit KT / Impulshaltigkeit KI",            f"+{daten.get('KT','0 dB')} / +{daten.get('KI','0 dB')}", False),
            ("Berechneter Schalldruckpegel Lp am Immissionsort",  f"{lp_val} dB(A)", False),
            ("Beurteilungspegel Lr = Lp + KT + KI",              f"{lr} dB(A)", True),
            ("Gesetzlicher Immissionsrichtwert (Nacht)",          f"{grenzwert} dB(A)", False),
            ("Reserve / Ueberschreitung",                         f"{round(grenzwert-lr,1):+.1f} dB(A)", True),
        ]
        for i,(lbl,val,is_key) in enumerate(result_rows):
            is_lr = "Beurteilungspegel Lr" in lbl
            is_res = "Reserve" in lbl
            if is_lr:
                bg = colors.HexColor("#D0F0FF") if konform else colors.HexColor("#FFD0D0")
            elif is_res:
                bg = colors.HexColor("#C8F5D8") if konform else colors.HexColor("#FFD0D0")
            else:
                bg = LGREY if i%2==0 else WHITE
            c.setFillColor(bg)
            c.rect(12*mm, y - 6*mm, W - 24*mm, 6*mm, fill=1, stroke=0)
            c.setFillColor(GREY)
            c.setFont("Helvetica-Bold" if (is_lr or is_res) else "Helvetica", 9)
            c.drawString(15*mm, y - 4*mm, lbl)
            col = (GREEN if konform else RED) if (is_lr or is_res) else BLACK
            c.setFillColor(col)
            c.setFont("Helvetica-Bold" if (is_lr or is_res) else "Helvetica", 9)
            c.drawRightString(W - 15*mm, y - 4*mm, val)
            y -= 6*mm

        # Fazit-Box
        y -= 5*mm
        delta_val = round(grenzwert - lr, 1)
        if konform:
            fazit = (f"FAZIT: Die geplante Installation ist aus schalltechnischer Sicht als unbedenklich "
                     f"einzustufen. Der prognostizierte Immissionswert von {lr} dB(A) liegt um "
                     f"{abs(delta_val)} dB(A) unter dem gesetzlichen Nachtgrenzwert von {grenzwert} dB(A).")
            box_col = GREEN
        else:
            fazit = (f"FAZIT: Die geplante Installation ueberschreitet den Nachtgrenzwert von "
                     f"{grenzwert} dB(A) um {abs(delta_val)} dB(A). Schutzmassnahmen erforderlich.")
            box_col = RED
        c.setFillColor(colors.HexColor("#E8F7FD") if konform else colors.HexColor("#FDE8E8"))
        c.setStrokeColor(box_col)
        box_h = 16*mm
        c.roundRect(12*mm, y - box_h, W - 24*mm, box_h, 2*mm, fill=1, stroke=1)
        c.setFillColor(box_col)
        c.setFont("Helvetica-Bold", 8.5)
        words = fazit.split()
        cur_line = ""
        fy = y - 4*mm
        for word in words:
            test = cur_line + " " + word if cur_line else word
            if c.stringWidth(test, "Helvetica-Bold", 8.5) < (W - 34*mm):
                cur_line = test
            else:
                c.drawString(16*mm, fy, cur_line)
                fy -= 4.5*mm
                cur_line = word
        if cur_line:
            c.drawString(16*mm, fy, cur_line)
        y -= box_h + 3*mm

    c.showPage()

    # ══════════════════════════════════════
    # SEITE 3 – MASSNAHMEN, NORMEN, DISCLAIMER
    # ══════════════════════════════════════
    draw_header(c, 3)
    draw_footer(c)
    y = H - 35*mm

    # 5. SCHUTZMASSNAHMEN
    y = section_title(c, y, "5. MASSNAHMEN ZUR QUALITAETSSICHERUNG")
    massnahmen = [
        ("Schallschutzgehaeuse / Kapselung",  f"{daten.get('Kapselung','-')} direkt an der Quelle",         "Luftschall-Daempfung direkt an der Quelle"),
        ("Anti-Vibrations-Sockel",             f"{daten.get('Koerperschall','-')} Koerperschall-Reduktion",   "Schwingungsentkopplung vom Baukörper"),
        ("Richtfaktor / Aufstellgeometrie",    f"{daten.get('Richtfaktor Q','-')}",                           "Optimale Ausblasrichtung zur Strasse"),
        ("Topologie / Schallabschirmung",      f"{daten.get('Topologie','-')}",                               "Nutzung baulicher Abschirmung"),
    ]
    # Tabellenkopf
    c.setFillColor(colors.HexColor("#E0F4FC"))
    c.rect(12*mm, y - 6*mm, W - 24*mm, 6*mm, fill=1, stroke=0)
    c.setFillColor(GREY); c.setFont("Helvetica-Bold", 8)
    c.drawString(15*mm, y - 4*mm, "Massnahme")
    c.drawString(85*mm,  y - 4*mm, "Wirkung / Wert")
    c.drawString(130*mm, y - 4*mm, "Funktion")
    y -= 6*mm
    for i,(m,w,f) in enumerate(massnahmen):
        c.setFillColor(LGREY if i%2==0 else WHITE)
        c.rect(12*mm, y - 5.5*mm, W - 24*mm, 5.5*mm, fill=1, stroke=0)
        c.setFillColor(GREY); c.setFont("Helvetica", 8)
        c.drawString(15*mm, y - 3.5*mm, m)
        c.setFillColor(GREEN); c.setFont("Helvetica-Bold", 8)
        c.drawString(85*mm,  y - 3.5*mm, w)
        c.setFillColor(BLACK); c.setFont("Helvetica", 8)
        c.drawString(130*mm, y - 3.5*mm, f)
        y -= 5.5*mm
    y -= 5*mm

    # 6. NORMATIVE GRUNDLAGEN
    y = section_title(c, y, "6. NORMATIVE GRUNDLAGEN & RECHTSRAHMEN")
    normen = [
        ("OENORM ISO 9613-2",     "Berechnung Schallausbreitung im Freien, Abschirmmasse"),
        ("OeAL-Richtlinie Nr. 3", "Beurteilung von Laerm im Nachbarschaftsbereich"),
        ("TA Laerm (DE)",         "Technische Anleitung Laermschutz (6. BImSchV)"),
        ("HRN ISO 1996-2 (HR)",   "Grenzwerte Laermimmissionen Kroatien"),
        ("Samsung TDB 2025",      f"Herstellerseitige Schallleistungsangabe: {daten.get('Lw','-')} dB(A)"),
        ("DIN EN ISO 9614",       "Messung Schallleistungspegel durch Schallintensitaet"),
    ]
    c.setFillColor(colors.HexColor("#E0F4FC"))
    c.rect(12*mm, y - 6*mm, W - 24*mm, 6*mm, fill=1, stroke=0)
    c.setFillColor(GREY); c.setFont("Helvetica-Bold", 8)
    c.drawString(15*mm, y - 4*mm, "Norm / Regelwerk")
    c.drawString(65*mm,  y - 4*mm, "Anwendungsbereich")
    y -= 6*mm
    for i,(n,a) in enumerate(normen):
        c.setFillColor(LGREY if i%2==0 else WHITE)
        c.rect(12*mm, y - 5*mm, W - 24*mm, 5*mm, fill=1, stroke=0)
        c.setFillColor(BLUE); c.setFont("Helvetica-Bold", 8)
        c.drawString(15*mm, y - 3.5*mm, n)
        c.setFillColor(BLACK); c.setFont("Helvetica", 8)
        c.drawString(65*mm,  y - 3.5*mm, a)
        y -= 5*mm
    y -= 6*mm

    # 7. HAFTUNGSAUSSCHLUSS
    y = section_title(c, y, "7. HAFTUNGSAUSSCHLUSS & RECHTLICHER HINWEIS")
    y -= 2*mm
    c.setFillColor(colors.HexColor("#FDE8E8"))
    c.setStrokeColor(RED)
    c.roundRect(12*mm, y - 10*mm, W - 24*mm, 10*mm, 2*mm, fill=1, stroke=1)
    c.setFillColor(RED); c.setFont("Helvetica-Bold", 9)
    c.drawCentredString(W/2, y - 6.5*mm,
        "Dieses Dokument ist KEIN Sachverstaendigengutachten!")
    y -= 14*mm

    c.setFillColor(BLACK); c.setFont("Helvetica", 8)
    def draw_wrapped(c, y, txt, x=12*mm, font="Helvetica", size=8, width=None):
        if width is None:
            width = W - 28*mm
        words = txt.split()
        cur_line = ""
        for word in words:
            test = cur_line + " " + word if cur_line else word
            if c.stringWidth(test, font, size) < width:
                cur_line = test
            else:
                if cur_line:
                    c.drawString(x, y, cur_line)
                    y -= size * 1.5
                cur_line = word
        if cur_line:
            c.drawString(x, y, cur_line)
            y -= size * 1.5
        return y

    disc_clean = (
        "7.1 Charakter und Zweck: Die vorliegende Berechnung ist eine planungstechnische Prognose, "
        "erstellt im Zuge der Anlagenauslegung. Sie basiert auf anerkannten Berechnungsverfahren "
        "nach OENORM ISO 9613-2, OeAL 3, TA Laerm sowie auf den veroeffentlichten Herstellerangaben "
        "(Samsung TDB RAC R32 NASA, Ver. 1.0, Dezember 2024). "
        "7.2 Haftungsausschluss: Der Ersteller uebernimmt keine Haftung fuer Abweichungen zwischen "
        "Prognose und Messung, Veraenderungen nach Berichtserstellung, behoerdliche Entscheidungen "
        "oder Ansprueche Dritter. "
        "7.3 Empfehlung: Fuer eine behoerdlich anerkannte Schallmessung ist ein akkreditiertes "
        "Akustikinstitut oder gerichtlich beeideter Sachverstaendiger zu beauftragen. "
        "Messgeraete: Klasse 1 nach DIN EN 61672 (+/-0.7 dB, 20-12500 Hz)."
    )
    y = draw_wrapped(c, y, disc_clean)
    y -= 8*mm

    # Unterschriftsfeld
    c.setStrokeColor(GREY)
    c.line(12*mm, y, 75*mm, y)
    c.line(105*mm, y, W - 12*mm, y)
    c.setFillColor(MIDGREY); c.setFont("Helvetica", 7)
    c.drawString(12*mm, y - 4*mm, "Ort, Datum")
    c.drawString(105*mm, y - 4*mm, "Ersteller / Stempel")
    y -= 18*mm

    # Branding-Footer
    c.setFillColor(BLUE)
    c.rect(12*mm, y, W - 24*mm, 0.5*mm, fill=1, stroke=0)
    y -= 5*mm
    c.setFillColor(GREY); c.setFont("Helvetica-Bold", 8.5)
    c.drawString(12*mm, y, "coolsulting e.U.  |  Michael Schaepers")
    y -= 4*mm
    c.setFillColor(MIDGREY); c.setFont("Helvetica", 7.5)
    c.drawString(12*mm, y, "ISO 9613-1/2  |  OeAL 3 / OENORM S 5021  |  TA Laerm  |  HRN ISO 1996  |  Samsung TDB RAC R32 NASA 2025")

    c.save()
    return buf.getvalue()


def create_word(daten, projekt, ersteller):
    doc = Document()
    h = doc.add_heading(projekt, 0)
    h.runs[0].font.color.rgb = RGBColor(0x36, 0xA9, 0xE1)
    doc.add_paragraph(f"Erstellt: {datetime.now().strftime('%d.%m.%Y %H:%M')}  |  Ersteller: {ersteller}")
    doc.add_paragraph("")
    doc.add_heading("Berechnungsdaten", level=1)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Light List Accent 1"
    t.rows[0].cells[0].text = "Parameter"
    t.rows[0].cells[1].text = "Wert"
    for k, v in daten.items():
        row = t.add_row()
        if str(k).startswith("_SEP"):
            c = row.cells[0].merge(row.cells[1])
            c.text = str(v)
            c.paragraphs[0].runs[0].bold = True
        else:
            row.cells[0].text = str(k)
            row.cells[1].text = str(v)
    doc.add_heading("Rechtlicher Hinweis", level=1)
    doc.add_paragraph(DISCLAIMER_TEXT)
    doc.add_paragraph("\n(C) coolsulting e.U. · Michael Schaepers · ")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ──────────────────────────────────────────────────
# STREAMLIT APP
# ──────────────────────────────────────────────────

def apply_css():
    st.markdown(f"""
    <style>
    html,body,[class*="css"] {{font-family:'TTCommons',sans-serif;}}
    h1,h2,h3,h4 {{font-family:'POEVetica',sans-serif; color:{C_BLUE} !important;}}
    [data-testid="stMetricValue"] {{color:{C_BLUE} !important; font-family:'POEVetica',sans-serif; font-size:1.9rem;}}
    .stTabs [data-baseweb="tab-highlight"] {{background:{C_BLUE};}}
    .stTabs [data-baseweb="tab"] {{font-family:'POEVetica',sans-serif;}}
    div[data-testid="stExpander"] {{border:1px solid {C_BLUE}44; border-radius:6px;}}
    .status-ok  {{background:{C_GREEN}22; border:2px solid {C_GREEN}; border-radius:8px;
                  padding:12px 20px; color:{C_GREEN}; font-family:'POEVetica',sans-serif;
                  font-size:1.15rem; font-weight:bold; text-align:center;}}
    .status-err {{background:{C_RED}22; border:2px solid {C_RED}; border-radius:8px;
                  padding:12px 20px; color:{C_RED}; font-family:'POEVetica',sans-serif;
                  font-size:1.15rem; font-weight:bold; text-align:center;}}
    footer {{visibility:hidden;}}
    .block-container {{padding-top:1.2rem;}}
    </style>""", unsafe_allow_html=True)


def render_header():
    logo_path = os.path.join(BASE_DIR, LOGO_FILE)
    has_logo = os.path.exists(logo_path)

    if has_logo:
        c1, c2 = st.columns([1, 7])
        with c1:
            st.image(logo_path, width=95)
        col = c2
    else:
        col = st.container()

    with col:
        st.markdown(f"""
        <div style='padding-top:2px; line-height:1.3; overflow:visible;'>
          <div style='display:flex; align-items:center; flex-wrap:nowrap; gap:6px;'>
            <span style='
              display:inline-block;
              font-family:POEVetica,Helvetica,sans-serif;
              font-size:1.55rem;
              font-weight:bold;
              color:#FFFFFF;
              background:{C_BLUE};
              border-radius:7px;
              padding:2px 10px 3px 10px;
              white-space:nowrap;
              line-height:1.4;
            '>&#176;cool</span><span style='
              font-family:POEVetica,Helvetica,sans-serif;
              font-size:1.55rem;
              font-weight:bold;
              color:{C_BLUE};
              white-space:nowrap;
            '>NEIGHBOR</span>
            <span style='
              font-family:TTCommons,Helvetica,sans-serif;
              font-size:0.9rem;
              color:{C_GREY};
              white-space:nowrap;
            '>v6.0 &nbsp;&#183;&nbsp; Schallimmissions-Prognose</span>
          </div>
          <div style='margin-top:3px;'>
            <span style='
              font-family:TTCommons,Helvetica,sans-serif;
              font-size:0.7rem;
              color:#888888;
            '>ISO&nbsp;9613 &nbsp;&#183;&nbsp; TA&nbsp;Laerm &nbsp;&#183;&nbsp;
               OeAL&nbsp;3 &nbsp;&#183;&nbsp; HRN&nbsp;ISO&nbsp;1996 &nbsp;&#183;&nbsp;
               Samsung&nbsp;RAC&nbsp;R32&nbsp;NASA&nbsp;TDB&nbsp;2025
            </span>
          </div>
        </div>""", unsafe_allow_html=True)
    st.markdown("---")



def main():
    st.set_page_config(page_title="coolNEIGHBOR v6.0", layout="wide", page_icon="🔊")
    apply_css()
    render_header()

    # ══════════════════════════════════════════════════════════
    # SIDEBAR – alle Eingaben
    # ══════════════════════════════════════════════════════════
    with st.sidebar:
        # ── PROJEKTDATEN ──────────────────────────────────────
        st.markdown(f"<span style='font-family:POEVetica,sans-serif;font-size:1.05rem;color:{C_BLUE};font-weight:bold;'>Projektdaten</span>", unsafe_allow_html=True)
        projekt      = st.text_input("Projektbezeichnung",
            value="Schallprognose Musteranlage",
            help="Eindeutige Bezeichnung des Projekts – erscheint auf allen Berichten.")
        ersteller    = st.text_input("Ersteller / Planer",
            value="coolsulting e.U.  |  Michael Schaepers",
            help="Name des Planers oder Bueros, der die Berechnung erstellt.")
        auftraggeber = st.text_input("Auftraggeber / Kunde",
            value="",
            help="Name des Kunden / Auftraggebers fuer den Bericht.")

        # ── NORM & GRENZWERT ──────────────────────────────────
        st.markdown("---")
        st.markdown(f"<span style='font-family:POEVetica,sans-serif;font-size:1.05rem;color:{C_BLUE};font-weight:bold;'>Norm & Grenzwert</span>", unsafe_allow_html=True)
        land      = st.selectbox("Land & Norm", LAENDER,
            help="Bestimmt welche Laermschutz-Norm angewendet wird:\n"
                 "• AT: OeAL 3 / OENORM S 5021\n"
                 "• DE: TA Laerm (6. BImSchV)\n"
                 "• HR: HRN ISO 1996-2")
        wid_opt   = list(GRENZWERTE[land].keys())
        widmung   = st.selectbox("Gebietswidmung", wid_opt,
            help="Flaechenwidmung des Immissionsortes (Nachbargrundstuck). "
                 "Bestimmt den gesetzlichen Grenzwert. "
                 "Im Zweifel bei der zust. Baubehorde erfragen.")
        zeitr     = st.radio("Beurteilungszeitraum",
            ["Nacht (massgebend)", "Tag"], horizontal=True,
            help="Fuer Klimageraete ist der Nachtwert massgebend (22–06 Uhr), "
                 "da Aussengeraete auch nachts laufen koennen.")
        grenzwert = GRENZWERTE[land][widmung]["Nacht" if "Nacht" in zeitr else "Tag"]
        st.metric("Grenzwert", f"{grenzwert} dB(A)")

        # ── GERAET & AUFSTELLUNG ─────────────────────────────
        st.markdown("---")
        st.markdown(f"<span style='font-family:POEVetica,sans-serif;font-size:1.05rem;color:{C_BLUE};font-weight:bold;'>Geraet & Aufstellung</span>", unsafe_allow_html=True)
        serie  = st.selectbox("Produktserie", list(SAMSUNG_DB.keys()),
            help="Samsung RAC R32 NASA Aussengeraete-Serien. "
                 "Bei Fremdgeraeten 'Freie Eingabe' waehlen und Lw manuell eingeben.")
        modell = st.selectbox("Modell", list(SAMSUNG_DB[serie].keys()),
            help="Modellnummer des Aussengeraets. "
                 "Lw und Luftmenge werden automatisch aus der Samsung TDB 2025 geladen.")
        anzahl = st.number_input("Anzahl Einheiten", min_value=1, max_value=20, value=1,
            help="Anzahl baugleicher Aussengeraete am selben Aufstellort. "
                 "Mehrere Geraete erhoehen den Pegel logarithmisch: "
                 "2x = +3 dB, 4x = +6 dB, 10x = +10 dB.")
        dbe     = SAMSUNG_DB[serie][modell]
        manuell = (modell == "Manuelle Eingabe")
        lw   = st.number_input("Lw [dB(A)]",  value=float(dbe["Lw"]),        step=0.5,
            disabled=not manuell,
            help="Schallleistungspegel des Aussengeraets – gerate-typische Kenngrösse, "
                 "unabhaengig von Abstand und Aufstellung. Aus Herstellerdatenblatt (EN ISO 9614).")
        luft = st.number_input("Luft [m3/h]", value=float(dbe["Luft_m3h"]),  step=50.0,
            disabled=not manuell,
            help="Luftvolumenstrom des Aussengeraets. Relevant fuer Kanalauslegung "
                 "und aerodynamischen Pegelzuschlag bei Kanalfuehrung.")
        gew  = st.number_input("Gewicht [kg]", value=float(dbe["Gewicht_kg"]), step=0.5,
            disabled=not manuell,
            help="Geraetegewicht – wird fuer die Berechnung der Koerperschallentkopplung "
                 "durch den Daempfungssockel benoetigt.")
        lp_1m = float(dbe["Lp_1m"]) if not manuell else st.number_input(
            "Lp @ 1m [dB(A)]", value=float(dbe["Lp_1m"]), step=0.5,
            help="Schalldruckpegel in 1m Abstand vom Geraet (Herstellerangabe, Freifeld).")
        if anzahl > 1:
            st.caption(f"Kaskade: +{kaskaden_zuschlag(anzahl)} dB(A) ({anzahl} Einheiten)")

        q_label = st.selectbox("Richtfaktor Q", [
            "Q=2  Freies Feld / Aussenwand",
            "Q=4  An einer Gebaeudeecke",
            "Q=8  Einspringende Ecke / Schacht",
        ], index=1, help="Der Richtfaktor beschreibt wie viele Flaechen den Schall reflektieren:\n"
                "• Q=2: Geraet steht frei oder an einer Wand (Halbraum) → +3 dB\n"
                "• Q=4: Geraet in einer Ecke (2 Wände) → +6 dB\n"
                "• Q=8: Geraet in einspringender Ecke (3 Flaechen) → +9 dB\n"
                "Jede zusaetzliche Reflexionsflaeche verdoppelt den Q-Wert und erhoeht den Pegel um 3 dB.")
        q_val = int(q_label.split("=")[1].split(" ")[0])

        einbau_label = st.selectbox("Einbausituation", [
            "Freies Feld  +0 dB",
            "Loggia / Balkon  +4 dB",
            "Innenhof / Schacht  +7 dB",
        ], help="Beschreibt ob das Geraet in einem (teil-)umschlossenen Raum steht:\n"
                "• Freies Feld: Schall kann frei abstrahlen (kein Zuschlag)\n"
                "• Loggia/Balkon: Schall wird von 2-3 Waenden reflektiert → diffuses Schallfeld +4 dB\n"
                "• Innenhof/Schacht: Starker Hallraumeffekt, Schall kann nicht entweichen → +7 dB\n"
                "Dieser Effekt ist physikalisch unabhaengig vom Q-Faktor und addiert sich!")
        raumzuschlag = 0.0 if "Frei" in einbau_label else (4.0 if "Loggia" in einbau_label else 7.0)
        w_dist = st.number_input("Wandabstand [m]", value=0.3, step=0.05, min_value=0.0,
            help="Abstand des Aussengeraets von der naechsten Wand oder Hausfassade. "
                 "Sehr nahe Waende (< 1m) erhoehen den Pegel durch Reflexion um bis zu +3 dB. "
                 "Empfehlung: Mindestabstand 0,5m zur Wand einhalten.")

        # ── DAEMPFUNG & KAPSELUNG ────────────────────────────
        st.markdown("---")
        st.markdown(f"<span style='font-family:POEVetica,sans-serif;font-size:1.05rem;color:{C_BLUE};font-weight:bold;'>Daempfung & Kapselung</span>", unsafe_allow_html=True)
        d_geh  = st.number_input("Kapselung [dB(A)]", value=0.0, step=0.5,
            help="Daempfung durch akustische Einhausung oder Schallschutzhaube. "
                 "Typische Werte: Leichtbau-Haube 5-10 dB, Vollkapselung 15-25 dB. "
                 "Ohne Kapselung: 0 dB eingeben.")
        m_fuss = st.number_input("Daempfungssockel [kg]", value=8.0, step=1.0,
            help="Gewicht des Anti-Vibrations-Sockels oder der Gummifuesse. "
                 "Schwerere Soeckel daempfen Koerperschall (Schwingungen ins Gebaeude) effektiver. "
                 "Faustformel: Sockelgewicht >= 20% des Geraetegewichts. "
                 "Kein Kesselpodest verwenden – dieses wirkt als Resonanzkoerper!")
        d_koerper = round((m_fuss / (gew + m_fuss)) * 2.0, 1) if gew > 0 else 0.0
        st.caption(f"Koerperschall-Reduktion: -{d_koerper} dB(A) (Sockel {m_fuss} kg / Geraet {gew} kg)")
        qs = st.number_input("Kanalquerschnitt [m2]", value=0.0, step=0.05,
            help="Freier Querschnitt des angeschlossenen Luftfuehrungskanals. "
                 "0 = Geraet blaest frei ins Freie (kein Kanal). "
                 "Zu kleine Querschnitte erhoehen die Stroemungsgeschwindigkeit: "
                 "Bei > 3 m/s steigt der Schallpegel um 0,25 dB pro 1% Mehrgeschwindigkeit. "
                 "Empfehlung: Kanalgeschwindigkeit < 3 m/s (aussen) bzw. < 4 m/s (innen).")
        v_pre, da_pre = aero_zuschlag(luft, qs)
        if qs > 0 and v_pre > 3.0:
            st.warning(f"Stroemung: {v_pre} m/s → Pegelzuschlag: +{da_pre} dB(A) "
                       f"(Empfehlung: Querschnitt vergroessern!)")
        elif qs > 0:
            st.success(f"Stroemung: {v_pre} m/s – unkritisch (< 3 m/s)")

        # ── STRAFZUSCHLAEGE & METEOROLOGIE ──────────────────
        st.markdown("---")
        st.markdown(f"<span style='font-family:POEVetica,sans-serif;font-size:1.05rem;color:{C_BLUE};font-weight:bold;'>Strafzuschlaege (TA Laerm / OeAL 3)</span>", unsafe_allow_html=True)
        st.caption("Strafzuschlaege erhoehen den Beurteilungspegel Lr wenn das Geraeusch "
                   "besonders stoerend ist. Lr = Lp + KT + KI")

        k_t = st.selectbox("Tonhaltigkeit KT [dB]", [0, 3, 6], index=0,
            help="Zuschlag wenn das Geraeusch einen deutlich hoerbaren Ton enthaelt "
                 "(z.B. Brummen, Pfeifen, Surren des Verdichters):\n"
                 "• 0 dB: Kein auffaelliger Ton hoerbar\n"
                 "• +3 dB: Ton auffaellig (z.B. leichtes Brummen)\n"
                 "• +6 dB: Ton stark auffaellig (deutliches Pfeifen / Surren)\n"
                 "Beurteilung nach DIN 45681. Bei modernen Inverter-Geraeten oft 0 dB.")
        k_i = st.selectbox("Impulshaltigkeit KI [dB]", [0, 3, 6], index=0,
            help="Zuschlag wenn das Geraeusch kurze, haemmernde oder knackende Impulse enthaelt:\n"
                 "• 0 dB: Kein Impuls wahrnehmbar (Normalbetrieb)\n"
                 "• +3 dB: Impulse auffaellig (z.B. regelmaessiges Klicken)\n"
                 "• +6 dB: Impulse stark auffaellig (harte Schaltgeraeuschen)\n"
                 "Relevant bei Abtauumschaltung (Umkehrung des Kaeltekreises). "
                 "Im Normalbetrieb in der Regel 0 dB.")

        st.markdown(f"<span style='font-family:POEVetica,sans-serif;font-size:1.0rem;color:{C_GREY};font-weight:bold;'>Meteorologie (ISO 9613-1)</span>", unsafe_allow_html=True)
        temp = st.number_input("Aussentemperatur [C]", value=10.0, step=1.0,
            help="Lufttemperatur am Aufstellort. Beeinflusst die atmosphaerische Daempfung: "
                 "Kaeltere Luft daempft etwas staerker. "
                 "Fuer die Worst-Case Berechnung: 10°C (mittlere Jahrestemperatur AT/DE).")
        hum  = st.number_input("Luftfeuchtigkeit [%]", value=70.0, step=5.0,
            help="Relative Luftfeuchtigkeit. Beeinflusst die atmosphaerische Schallabsorption: "
                 "Trockene Luft (< 50%) daempft hohe Frequenzen staerker. "
                 "Standard-Rechenwert: 70% (nach ISO 9613-1 Referenzatmosphaere).")

        # ── SCHALLWEG & TOPOLOGIE ────────────────────────────
        st.markdown("---")
        st.markdown(f"<span style='font-family:POEVetica,sans-serif;font-size:1.05rem;color:{C_BLUE};font-weight:bold;'>Schallweg & Topologie</span>", unsafe_allow_html=True)
        topologie = st.selectbox("Topologie des Schallwegs", [
            "Direkter Sichtkontakt",
            "L-foermig (1 Kante)",
            "U-foermig (2 Kanten)",
        ], help="Beschreibt den Weg des Schalls vom Geraet zum Immissionsort (Nachbarfenster):\n"
                "• Direkter Sichtkontakt: Freie Sichtlinie, kein Hindernis\n"
                "• L-foermig: Schall muss um 1 Kante herum (Balkon-Bruestung, Mauer, Gelaender)\n"
                "• U-foermig: Schall muss um 2 Kanten herum (Bruestung + Vordach / Jalousienkasten)\n"
                "Je mehr Kanten und je massiver, desto mehr Beugungsdaempfung.")

        # Erklaerungsbox als aufpoppendes Expander-Modal (gross & lesbar)
        with st.expander("📐  Wie messe ich Schallweg & Luftlinie? (klicken zum Oeffnen)", expanded=False):
            st.markdown(f"""
<div style='font-size:1.0rem; line-height:1.7; color:#DDDDDD; padding:4px 0;'>

<b style='font-size:1.1rem; color:{C_BLUE};'>Gesamtschallweg (entlang der Schallbahn)</b><br>
Der Weg, den der Schall <b>tatsaechlich zuruecklegt</b> – bei Beugung laenger als die Luftlinie.<br>
<span style='color:#AAD4EE;'>➤ Messen: Geraet → Kante 1 → [Kante 2] → Fenster (entlang der Schallbahn)</span><br>
<span style='color:#FFB347;'>⚡ Dieser Wert bestimmt die geometrische Abstandsdaempfung!</span>

<br><br>

<b style='font-size:1.1rem; color:{C_BLUE};'>Direkte Luftlinie (Vogelpfeil)</b><br>
Die <b>kuerzeste Verbindung</b> zwischen Geraet und Nachbarfenster – wie der Vogel fliegt.<br>
<span style='color:#AAD4EE;'>➤ Bei direktem Weg: identisch mit dem Gesamtschallweg</span><br>
<span style='color:#AAD4EE;'>➤ Bei L/U-foermig: kuerzer als der entrollte Weg!</span>

<br><br>

<b style='font-size:1.1rem; color:{C_BLUE};'>Umweg = Beugungseffekt</b><br>
<b>Umweg = Gesamtschallweg − Luftlinie</b><br>
<span style='color:#36E16A;'>Groesserer Umweg → tieferer Schallschatten → staerkere Daempfung</span><br>
<span style='color:#E13636;'>Kein Umweg (beide gleich) → keine Beugungsdaempfung!</span>

<br><br>

<b style='font-size:1.05rem; color:#FFB347;'>⚠️  Wichtige Regeln:</b><br>
• Luftlinie immer &lt;= Gesamtschallweg<br>
• Bei L-foermig: Gesamtschallweg <b>muss groesser</b> als Luftlinie sein – sonst wirkt keine Beugung<br>
• Typischer Umweg an einer Balustrade: 0,3 – 1,5 m

</div>""", unsafe_allow_html=True)

        d_total  = st.number_input("Gesamtschallweg – entlang Schallbahn [m]", value=8.0, step=0.5, min_value=0.1,
            help="Entrollter Schallweg: Geraet → [Kante(n)] → Fenster.\n"
                 "BESTIMMT die geometrische Abstandsdaempfung (Adiv).\n"
                 "Bei L/U: Summe der Wegabschnitte entlang der Schallbahn.")

        # Direkte Luftlinie: Bei Direktem Weg automatisch = d_total setzen
        is_direkt = "Direkt" in topologie
        if is_direkt:
            d_direkt = d_total
            st.caption(f"Direkter Weg: Luftlinie = Gesamtschallweg = {d_total} m (automatisch)")
        else:
            d_direkt = st.number_input("Direkte Luftlinie Geraet → Fenster [m]", value=min(7.0, d_total - 0.5),
                step=0.5, min_value=0.1,
                help="Vogelpfeil-Abstand Geraet → Fenster.\n"
                     "Muss KLEINER als Gesamtschallweg sein, damit Beugung wirkt!\n"
                     "Umweg = Gesamtweg - Luftlinie → Fresnel-Beugungsdaempfung (ISO 9613-2).")

            umweg_preview = round(d_total - d_direkt, 2)
            if d_direkt >= d_total:
                st.error(f"⛔  Luftlinie ({d_direkt} m) muss KLEINER als Gesamtschallweg ({d_total} m) sein – sonst wirkt keine Beugung!")
            elif umweg_preview < 0.3:
                st.warning(f"⚠️  Umweg nur {umweg_preview} m – sehr geringe Beugungsdaempfung. Typisch: 0,5–2 m.")
            else:
                st.success(f"✅  Umweg: {umweg_preview} m → Beugung wirkt")

        b1_mat="Keine Barriere"; b2_mat="Keine Barriere"; b1_lenkt=False; b2_lenkt=False
        if "L-" in topologie or "U-" in topologie:
            b1_mat   = st.selectbox("Material Kante 1", [
                "Massiv & Luftdicht",
                "Leichtbauweise dicht",
                "Teiloffen / Lamellen",
                "Luftdurchlaessig (Gitter)"], key="b1m",
                help="Beschaffenheit der ersten Beugungskante:\n"
                     "• Massiv & Luftdicht: Beton, Vollziegel, Glas → max. 25 dB Daempfung\n"
                     "• Leichtbauweise dicht: Holz, Blech dicht → max. 15 dB\n"
                     "• Teiloffen / Lamellen: Jalousien, Lamellen → max. 5 dB\n"
                     "• Luftdurchlaessig: Gitterzaun, Gitter → 0 dB (kein Schirmeffekt)")
            b1_lenkt = st.checkbox("Kante 1 lenkt Luftstrom physisch ab", value=True, key="b1l",
                help="Aktivieren wenn der Luftstrom des Geraets an dieser Kante tatsaechlich "
                     "umgelenkt wird (z.B. Balkonbruestung als Windschutz). "
                     "Physische Umlenkung ergibt zusaetzlich 2 dB aerodynamische Daempfung.")
        if "U-" in topologie:
            b2_mat   = st.selectbox("Material Kante 2", [
                "Massiv & Luftdicht",
                "Leichtbauweise dicht",
                "Teiloffen / Lamellen",
                "Luftdurchlaessig (Gitter)"], key="b2m",
                help="Beschaffenheit der zweiten Beugungskante (z.B. Vordach, Jalousienkasten). "
                     "Gleiche Abstufung wie Kante 1.")
            b2_lenkt = st.checkbox("Kante 2 lenkt Luftstrom physisch ab", value=True, key="b2l",
                help="Aktivieren wenn der Luftstrom auch an der zweiten Kante umgelenkt wird. "
                     "Gibt zusaetzlich 1,5 dB aerodynamische Daempfung.")

        if manuell:
            st.markdown("---")
            st.markdown(f"<span style='font-family:POEVetica,sans-serif;font-size:1.05rem;color:{C_BLUE};font-weight:bold;'>Oktavband @ 1m [dB]</span>", unsafe_allow_html=True)
            st.caption("Schalldruckpegel je Oktavband aus Herstellermessung (Freifeld, 1m Abstand, "
                       "Aussengeraet, Kuehlbetrieb). Aus Datenblatt oder Akustikbericht uebernehmen.")
            oct_vals = list(dbe["Lp_oktav"])
            for i, hz in enumerate(OKTAV_BANDS):
                oct_vals[i] = st.number_input(f"{hz} Hz", value=float(oct_vals[i]), step=1.0, key=f"o{i}")
            lp_oktav_src = oct_vals
        else:
            lp_oktav_src = list(dbe["Lp_oktav"])

    # ══════════════════════════════════════════════════════════
    # ZENTRALE BERECHNUNG
    # ══════════════════════════════════════════════════════════
    erg = berechne_gesamt(
        lw=lw, anzahl=anzahl, luft=luft, qs=qs,
        d_geh=d_geh, d_koerper=d_koerper,
        k_t=k_t, k_i=k_i,
        d_total=d_total, topologie=topologie,
        b1_mat=b1_mat, b2_mat=b2_mat,
        b1_lenkt=b1_lenkt, b2_lenkt=b2_lenkt,
        q=q_val, w_dist=w_dist,
        temp=temp, hum=hum,
        raumzuschlag=raumzuschlag,
        d_direkt=d_direkt,
    )
    lr      = erg["lr"]
    lp      = erg["lp"]
    delta   = round(grenzwert - lr, 1)
    konform = delta >= 0
    lp_io   = oktav_am_io(lp_oktav_src, d_total, q_val, d_geh, anzahl)
    nr_k    = nr_klasse(lp_io)

    # ══════════════════════════════════════════════════════════
    # HAUPTBEREICH – Tabs
    # ══════════════════════════════════════════════════════════
    tabs = st.tabs([
        "📊  Ergebnis & NR-Kurven",
        "💾  Export",
        "⚡  Schnellrechner",
        "ℹ️  Normen & Physik",
    ])

    # ── TAB 1: ERGEBNIS ──────────────────────────────────────
    with tabs[0]:
        st.header(f"Planungstechnische Prognose")

        if konform:
            st.markdown(
                f'<div class="status-ok">OK  KONFORM  -  Lr = {lr} dB(A)  |  Reserve: {delta} dB(A) unter Grenzwert {grenzwert} dB(A)</div>',
                unsafe_allow_html=True)
        else:
            st.markdown(
                f'<div class="status-err">UEBERSCHRITTEN  -  Lr = {lr} dB(A)  |  {abs(delta)} dB(A) ueber Grenzwert {grenzwert} dB(A)</div>',
                unsafe_allow_html=True)
        st.markdown("### ")

        m1, m2, m3, m4, m5 = st.columns(5)
        m1.metric("Schalldruckpegel Lp",  f"{lp} dB(A)")
        m2.metric("Beurteilungspegel Lr", f"{lr} dB(A)",
                  delta=f"{delta:+.1f} dB" if konform else f"{delta:.1f} dB",
                  delta_color="normal" if konform else "inverse")
        m3.metric("Grenzwert",            f"{grenzwert} dB(A)")
        m4.metric("Beugung gesamt",       f"-{round(erg['dz1']+erg['dz2']+erg['d_uml'],1)} dB")
        m5.metric("Atm. Daempfung",       f"-{erg['a_atm']} dB")

        st.markdown("---")
        st.subheader("Berechnungsprotokoll (ISO 9613-2)")
        prot = {
            "_SEP1": "-- EINGANG --",
            "Schallleistungspegel Lw":    f"{lw} dB(A)  [{modell}]",
            "Anzahl / Kaskade":           f"{anzahl} Stk.  +{erg['d_kas']} dB(A)",
            "Aerodynamik":                f"+{erg['d_aero']} dB(A) ({erg['v_luft']} m/s)" if erg['d_aero'] > 0 else "-",
            "Kapselung":                  f"-{d_geh} dB(A)" if d_geh > 0 else "-",
            "Koerperschall":              f"-{d_koerper} dB(A)  ({m_fuss} kg)" if d_koerper > 0 else "-",
            "Lw effektiv":                f"{erg['lw_eff']} dB(A)",
            "_SEP2": "-- AUSBREITUNG --",
            "Richtfaktor Q":              f"Q={q_val}  +{erg['q_db']} dB",
            "Wandreflexion":              f"+{erg['w_plus']} dB  ({w_dist} m Wandabstand)",
            "Raumzuschlag":               f"+{raumzuschlag} dB" if raumzuschlag > 0 else "-",
            "Geom. Ausbreitung (Adiv)":   f"-{erg['a_div']} dB  ({d_total} m)",
            "Atm. Daempfung ISO 9613-1":  f"-{erg['a_atm']} dB",
            "Schallumweg delta":          f"+{erg['umweg']} m",
            "Beugung Kante 1":            f"-{erg['dz1']} dB" if erg['dz1'] > 0 else "-",
            "Beugung Kante 2":            f"-{erg['dz2']} dB" if erg['dz2'] > 0 else "-",
            "Aerodyn. Ablenkung":         f"-{erg['d_uml']} dB" if erg['d_uml'] > 0 else "-",
            "_SEP3": "-- ERGEBNIS --",
            "Schalldruckpegel Lp":        f"{lp} dB(A)",
            "Tonhaltigkeit KT":           f"+{k_t} dB",
            "Impulshaltigkeit KI":        f"+{k_i} dB",
            "Beurteilungspegel Lr":       f"{lr} dB(A)",
            "Grenzwert":                  f"{grenzwert} dB(A)",
            "Delta":                      f"{delta:+.1f} dB(A)",
            "STATUS":                     "KONFORM" if konform else "UEBERSCHRITTEN",
        }
        for k, v in prot.items():
            if str(k).startswith("_SEP"):
                st.markdown(f"**{v}**")
            else:
                ca, cb = st.columns([3, 2])
                ca.caption(k)
                cb.markdown(f"**{v}**")

        st.markdown("---")
        st.subheader("Oktavspektrum am Immissionsort & NR-Kurven")
        cf, ct = st.columns([3, 1])
        with cf:
            fig = nr_plot(lp_io, f"Spektrum - {modell.split('(')[0].strip()} - {d_total} m")
            st.plotly_chart(fig, use_container_width=True)
        with ct:
            st.markdown("**Oktavbandpegel [dB]**")
            for hz, val in zip(OKTAV_BANDS, lp_io):
                cx, cy = st.columns(2)
                cx.caption(f"{hz} Hz")
                cy.markdown(f"**{val}**")
            st.markdown("---")
            if nr_k:
                st.success(f"NR-Klasse: **NR {nr_k}**")
            else:
                st.error("NR > 65")

    # ── TAB 2: EXPORT ────────────────────────────────────────
    with tabs[1]:
        st.header("Bericht & Export")
        exp = {
            "_SEP0": "-- PROJEKT --",
            "Projekt":             projekt,
            "Ersteller":           ersteller,
            "Auftraggeber":        auftraggeber,
            "Datum":               datetime.now().strftime("%d.%m.%Y %H:%M"),
            "_SEP1": "-- STANDORT & NORM --",
            "Land / Norm":         land,
            "Gebietswidmung":      widmung,
            "Beurteilungszeitraum": zeitr,
            "Immissionsrichtwert": f"{grenzwert} dB(A)",
            "_SEP2": "-- GERAET & QUELLE --",
            "Produktserie":        serie,
            "Modell":              modell,
            "Anzahl":              f"{anzahl} Stk.",
            "Lw":                  f"{lw} dB(A)",
            "Kaskaden-Zuschlag":   f"+{erg['d_kas']} dB(A)",
            "Aerodynamik":         f"+{erg['d_aero']} dB(A)" if erg['d_aero'] > 0 else "-",
            "Kapselung":           f"-{d_geh} dB(A)",
            "Koerperschall":       f"{m_fuss} kg - -{d_koerper} dB(A)",
            "Lw effektiv":         f"{erg['lw_eff']} dB(A)",
            "_SEP3": "-- AUSBREITUNG --",
            "Richtfaktor Q":       f"Q={q_val}",
            "Einbausituation":     einbau_label,
            "Raumzuschlag":        f"+{raumzuschlag} dB",
            "Topologie":           topologie,
            "Gesamtschallweg":     f"{d_total} m",
            "Direkte Luftlinie":   f"{d_direkt} m",
            "Schallumweg":         f"+{erg['umweg']} m",
            "Atm. Daempfung":      f"-{erg['a_atm']} dB",
            "Beugung K1":          f"-{erg['dz1']} dB",
            "Beugung K2":          f"-{erg['dz2']} dB",
            "Meteorologie":        f"{temp} Grad C / {hum}% rF",
            "_SEP4": "-- ERGEBNIS --",
            "Schalldruckpegel Lp": f"{lp} dB(A)",
            "KT":                  f"+{k_t} dB",
            "KI":                  f"+{k_i} dB",
            "Beurteilungspegel Lr": f"{lr} dB(A)",
            "Grenzwert":           f"{grenzwert} dB(A)",
            "Delta":               f"{delta:+.1f} dB(A)",
            "STATUS":              "KONFORM" if konform else "UEBERSCHRITTEN",
            "_SEP5": "-- OKTAVSPEKTRUM AM IMMISSIONSORT --",
        }
        for hz, val in zip(OKTAV_BANDS, lp_io):
            exp[f"{hz} Hz"] = f"{val} dB"
        exp["NR-Klasse"] = f"NR {nr_k}" if nr_k else "> NR 65"

        col_p, col_w = st.columns(2)
        with col_p:
            fname_p = f"coolNEIGHBOR_{projekt.replace(' ','_')[:28]}.pdf"
            st.download_button(
                "PDF-Bericht herunterladen",
                data=create_pdf(
                    exp, projekt, ersteller,
                    erg=erg, grenzwert=grenzwert,
                    lp_io=lp_io, lp_oktav_src=lp_oktav_src,
                    konform=konform, lr=lr, lp_val=lp,
                    modell=modell, land=land, widmung=widmung,
                    auftraggeber=auftraggeber,
                ),
                file_name=fname_p, mime="application/pdf",
                use_container_width=True)
        with col_w:
            fname_w = f"coolNEIGHBOR_{projekt.replace(' ','_')[:28]}.docx"
            st.download_button(
                "Word-Bericht herunterladen",
                data=create_word(exp, projekt, ersteller),
                file_name=fname_w,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True)
        st.caption("Alle Berichte enthalten den vollstaendigen Haftungsausschluss gemaess OeAL 3 / TA Laerm.")
        with st.expander("Haftungsausschluss (Vorschau)"):
            st.text(DISCLAIMER_TEXT)

    # ── TAB 3: SCHNELLRECHNER ────────────────────────────────
    with tabs[2]:
        st.header("Schnellrechner & Hilfstools")
        cq1, cq2 = st.columns(2)
        with cq1:
            st.subheader("Abstandsverdoppelungs-Rechner")
            st.caption("Verdopplung der Distanz = -6 dB(A)  (BWP Grundregel 7)")
            lp_ref = st.number_input("Lp bekannt [dB(A)]", value=50.0, step=0.5, key="sq_lp")
            d_ref  = st.number_input("Distanz bekannt [m]", value=1.0, step=0.5, min_value=0.1, key="sq_dr")
            d_ziel = st.number_input("Zieldistanz [m]",     value=8.0, step=0.5, min_value=0.1, key="sq_dz")
            st.metric("Lp bei Zieldistanz", f"{round(lp_ref - 20*math.log10(d_ziel/d_ref), 1)} dB(A)")
            st.markdown("---")
            st.subheader("Logarithmische Pegeladdition")
            txt = st.text_input("Pegel kommagetrennt [dB(A)]", value="55, 57, 52")
            try:
                lvls = [float(x.strip()) for x in txt.split(",")]
                st.metric("Gesamtpegel", f"{round(log_add(lvls), 1)} dB(A)")
            except Exception:
                st.warning("Bitte Zahlen kommagetrennt eingeben.")
            st.markdown("---")
            st.subheader("Mindestabstand fuer Grenzwerteinhaltung")
            lw_q = st.number_input("Lw [dB(A)]",     value=float(lw), step=0.5, key="sq_lw")
            gw_q = st.number_input("Grenzwert [dB(A)]", value=float(grenzwert), step=1.0, key="sq_gw")
            try:
                d_min = round(math.sqrt(2 / (4*math.pi) * 10**((lw_q-gw_q)/10)), 2)
                st.metric("Mindestabstand (Q=2, Freifeld)", f"{d_min} m")
            except Exception:
                pass

        with cq2:
            st.subheader("Kaskaden-Tabelle (BWP Tab. 3.1)")
            for n_k, dl_k in [(1,0.0),(2,3.0),(3,4.8),(4,6.0),(5,7.0),(6,7.8),(7,8.5),(8,9.0),(9,9.5),(10,10.0),(12,10.8)]:
                ca_, cb_ = st.columns(2)
                ca_.caption(f"{n_k} Einheiten")
                cb_.markdown(f"**+{dl_k} dB(A)**")
            st.markdown("---")
            st.subheader("Tab. 4.1 BWP – Schalldruckpegel-Abschlag")
            st.caption("Abschlag vom Lw in dB(A) nach Q und Abstand")
            abst_tab = [1,2,4,5,6,8,10,12,15]
            q_rows_t = {
                "Q=2": [-8,-14,-20,-22,-23.5,-26,-28,-29.5,-31.5],
                "Q=4": [-5,-11,-17,-19,-20.5,-23,-25,-26.5,-28.5],
                "Q=8": [-2, -8,-14,-16,-17.5,-20,-22,-23.5,-25.5],
            }
            hdr = "| | " + " | ".join(f"**{a}m**" for a in abst_tab) + " |"
            sep = "|---|" + "|".join(["---"]*len(abst_tab)) + "|"
            rows = [hdr, sep]
            for ql, vals in q_rows_t.items():
                rows.append("| "+ql+" | "+" | ".join(str(v) for v in vals)+" |")
            st.markdown("\n".join(rows))

    # ── TAB 4: NORMEN & PHYSIK ───────────────────────────────
    with tabs[3]:
        st.header("Normen & Physikalische Grundlagen")
        gw = GRENZWERTE[land][widmung]
        col_n, col_p = st.columns(2)
        with col_n:
            st.subheader(f"Grenzwerte: {land}")
            for wid, vals in GRENZWERTE[land].items():
                ca, cb, cc = st.columns([3,1,1])
                ca.caption(wid)
                cb.markdown(f"Tag: **{vals['Tag']}**")
                cc.markdown(f"Nacht: **{vals['Nacht']}**")
            st.markdown("---")
            st.markdown("""
**Messpunkt:** 0,5 m vor der Mitte des geoeffneten Fensters des am staerksten betroffenen schutzbeduerftigen Raumes.

**Schutzbeduerftige Raeume (DIN 4109):**
Wohn-/Schlafzimmer, Kinderzimmer, Arbeits- und Unterrichtsraeume.

**Einzelne Geraeauschspitzen** duerfen den Richtwert tagsueber um max. 30 dB(A), nachts um max. 20 dB(A) ueberschreiten.
""")
        with col_p:
            st.subheader("Berechnungsformeln")
            st.markdown("""
**Schalldruckpegel (ISO 9613-2):**
L_Aeq = Lw + 10·log(Q / 4pi·r2)

**Richtfaktor Q:**
- Q=2: Halbraum (Aussenwand)
- Q=4: Viertelraum (Gebaeudeecke)  
- Q=8: Achtelraum (Einspringende Ecke)

**Beurteilungspegel (TA Laerm):**
Lr = Lp + KT + KI

**Kaskade n gleicher Quellen:**
Lw_ges = Lw_einzel + 10·log10(n)

**Abstandsverdopplung:**
Delta L = -6 dB(A) je Verdopplung

**Beugung Fresnel (ISO 9613-2):**
N = 2·delta/lambda
Dz = 10·log10(3 + 20·N)

**NR-Kurven (ISO 1996-1):**
Frequenzgewichtete Grenzlinien 63-8000 Hz.
NR-Klasse = niedrigste NR-Kurve die das Spektrum nicht ueberschreitet.
""")
            st.subheader("NR-Kurven Referenz [dB]")
            nr_show = [25,30,35,40,45,50]
            nr_hdr = "| NR | " + " | ".join(f"{hz}Hz" for hz in OKTAV_BANDS) + " |"
            nr_sep = "|---|" + "|".join(["---"]*8) + "|"
            nr_rows = [nr_hdr, nr_sep]
            for nrv in nr_show:
                nr_rows.append("| "+str(nrv)+" | "+" | ".join(f"{v:.0f}" for v in NR_KURVEN[nrv])+" |")
            st.markdown("\n".join(nr_rows))

    # FOOTER
    st.markdown("---")
    st.markdown(
        f"<div style='text-align:center;font-size:11px;color:{C_GREY};'>"
        f"(C) coolsulting e.U. - Michael Schaepers - "
        f"coolNEIGHBOR v6.0 - ISO 9613-1/2 - TA Laerm - OeAL 3 - HRN ISO 1996 - "
        f"Samsung TDB RAC R32 NASA 2025</div>",
        unsafe_allow_html=True)


if __name__ == "__main__":
    main()